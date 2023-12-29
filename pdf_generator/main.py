# from docx import Document
# from docx.shared import Inches, Pt
# import requests
# from io import BytesIO

# def set_paragraph_font(paragraph, font_name, font_size, bold=False):
#     for run in paragraph.runs:
#         font = run.font
#         font.name = font_name
#         font.size = Pt(font_size)
#         # font.bold = bold

# def find_and_replace_paragraphs(paragraphs, field, replacement):
#     for paragraph in paragraphs:
#         if field in paragraph.text:
#             paragraph.text = paragraph.text.replace(field, replacement)
#             set_paragraph_font(paragraph, 'Arial', 9, True)

# def find_and_replace_tables(tables, field, replacement):
#     for table in tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 for paragraph in cell.paragraphs:
#                     find_and_replace_paragraphs([paragraph], field, replacement)

# def replace_image_with_url(doc, old_image_filename, image_url, width=None, height=None):
#     for rel in doc.part.rels.values():
#         if "image" in rel.reltype and old_image_filename in rel.target_ref:
#             doc.part.rels[rel.reltype].remove(rel)
#             response = requests.get(image_url)
#             image_data = BytesIO(response.content)
#             new_rel = doc.part.relate_to(image_data, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image")
#             doc.part.rels[new_rel.reltype].append(new_rel)

#             # Create a new paragraph with the new image
#             paragraph = doc.add_paragraph()
#             run = paragraph.add_run()
#             run.add_picture(image_data, width=width, height=height)

# # Load the document template
# doc = Document('entrance.docx')

# # Sample student_data
# student_data = {
#     'name': 'John Doe',
#     'comp_name': 'ABC Corporation',
#     'reg_no': '12345',
#     'camp_date': '2023-08-15',
#     'parent_name': 'Jane Doe',
#     'address': '123 Main Street',
#     'city': 'Anytown',
#     'district': 'County',
#     'state': 'State',
#     'email': 'john.doe@example.com',
#     'contact_no': '123-456-7890',
#     'whatsapp_no': '987-654-3210',
#     'fathers_no': '111-222-3333',
#     'mothers_no': '444-555-6666',
#     'birth_date': '2000-01-01',
#     'blood_grp': 'A+',
#     'standard': '12th',
#     'school_name': 'XYZ High School',
#     'batch':'Alpha',
#     'total_days':'7 Days',
#     'pincode':"787876",
#     'camp_name':"Summer Camp"
# }

# # Replace text fields in paragraphs
# find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD STUDENT_NAME}', student_data['name'])
# find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD COMPANY_NAME}', student_data['comp_name'])
# find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD REG_NO}', student_data['reg_no'])
# find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD CAMP_DATE}', student_data['camp_date'])
# find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD PARENT_NAME}', student_data['parent_name'])
# find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD ADDRESS}', student_data['address'])
# find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD CITY}', student_data['city'])
# find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD DISTRICT}', student_data['district'])
# find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD STATE}', student_data['state'])
# find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD EMAIL}', student_data['email'])
# find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD CONTACT_NO}', student_data['contact_no'])
# find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD WHATSAPP_NO}', student_data['whatsapp_no'])
# find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD FATHERS_NO}', student_data['fathers_no'])
# find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD MOTHERS_NO}', student_data['mothers_no'])
# find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD BIRTH_DATE}', student_data['birth_date'])
# find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD BLOOD_GRP}', student_data['blood_grp'])
# find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD STANDARD}', student_data['standard'])
# find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD SCHOOL_NAME}', student_data['school_name'])
# find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD BATCH}', student_data['batch'])
# find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD TOTAL_DAYS}', student_data['total_days'])
# find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD PINCODE}', student_data['pincode'])
# find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD CAMP_NAME}', student_data['camp_name'])

# # Replace text fields in tables
# find_and_replace_tables(doc.tables, '{MERGEFIELD STUDENT_NAME}', student_data['name'])
# find_and_replace_tables(doc.tables, '{MERGEFIELD COMPANY_NAME}', student_data['comp_name'])
# find_and_replace_tables(doc.tables, '{MERGEFIELD REG_NO}', student_data['reg_no'])
# find_and_replace_tables(doc.tables, '{MERGEFIELD CAMP_DATE}', student_data['camp_date'])
# find_and_replace_tables(doc.tables, '{MERGEFIELD PARENT_NAME}', student_data['parent_name'])
# find_and_replace_tables(doc.tables, '{MERGEFIELD ADDRESS}', student_data['address'])
# find_and_replace_tables(doc.tables, '{MERGEFIELD CITY}', student_data['city'])
# find_and_replace_tables(doc.tables, '{MERGEFIELD DISTRICT}', student_data['district'])
# find_and_replace_tables(doc.tables, '{MERGEFIELD STATE}', student_data['state'])
# find_and_replace_tables(doc.tables, '{MERGEFIELD EMAIL}', student_data['email'])
# find_and_replace_tables(doc.tables, '{MERGEFIELD CONTACT_NO}', student_data['contact_no'])
# find_and_replace_tables(doc.tables, '{MERGEFIELD WHATSAPP_NO}', student_data['whatsapp_no'])
# find_and_replace_tables(doc.tables, '{MERGEFIELD FATHERS_NO}', student_data['fathers_no'])
# find_and_replace_tables(doc.tables, '{MERGEFIELD MOTHERS_NO}', student_data['mothers_no'])
# find_and_replace_tables(doc.tables, '{MERGEFIELD BIRTH_DATE}', student_data['birth_date'])
# find_and_replace_tables(doc.tables, '{MERGEFIELD BLOOD_GRP}', student_data['blood_grp'])
# find_and_replace_tables(doc.tables, '{MERGEFIELD STANDARD}', student_data['standard'])
# find_and_replace_tables(doc.tables, '{MERGEFIELD SCHOOL_NAME}', student_data['school_name'])
# find_and_replace_tables(doc.tables, '{MERGEFIELD BATCH}', student_data['batch'])
# find_and_replace_tables(doc.tables, '{MERGEFIELD TOTAL_DAYS}', student_data['total_days'])
# find_and_replace_tables(doc.tables, '{MERGEFIELD PINCODE}', student_data['pincode'])
# find_and_replace_tables(doc.tables, '{MERGEFIELD CAMP_NAME}', student_data['camp_name'])

# # Replace specific image with a new image from URL
# replace_image_with_url(doc, 'old_image.png', 'https://example.com/new_image.png', width=Inches(2), height=Inches(2))

# # Save the modified document
# doc.save(f'{student_data["reg_no"]}_entrance_card.docx')





from docx import Document
from docx.shared import Inches, Pt
import requests
from io import BytesIO
from docx2pdf import convert

def set_paragraph_font(paragraph, font_name, font_size, bold=False):
    for run in paragraph.runs:
        font = run.font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = bold

def find_and_replace_paragraphs(paragraphs, field, replacement):
    for paragraph in paragraphs:
        if field in paragraph.text:
            paragraph.text = paragraph.text.replace(field, replacement)
            set_paragraph_font(paragraph, 'Arial', 9, False)

def find_and_replace_tables(tables, field, replacement):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    find_and_replace_paragraphs([paragraph], field, replacement)

def replace_image_with_url(doc, old_image_filename, image_url, width=None, height=None):
    for rel in doc.part.rels.values():
        if "image" in rel.reltype and old_image_filename in rel.target_ref:
            doc.part.rels[rel.reltype].remove(rel)
            response = requests.get(image_url)
            image_data = BytesIO(response.content)
            new_rel = doc.part.relate_to(image_data, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image")
            doc.part.rels[new_rel.reltype].append(new_rel)

            # Create a new paragraph with the new image
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(image_data, width=width, height=height)

# Load the document template
doc = Document('entrance.docx')

# Sample student_data
student_data = {
    'name': 'Vedant Badre',
    'comp_name': 'BNB',
    'reg_no': '7878767676',
    'camp_date': '2023-08-15',
    'parent_name': 'Jane Doe',
    'address': '123 Main Street',
    'city': 'Anytown',
    'district': 'County',
    'state': 'State',
    'email': 'john.doe@example.com',
    'contact_no': '123-456-7890',
    'whatsapp_no': '987-654-3210',
    'fathers_no': '111-222-3333',
    'mothers_no': '444-555-6666',
    'birth_date': '2000-01-01',
    'blood_grp': 'A+',
    'standard': '12th',
    'school_name': 'XYZ High School',
    'batch':'Alpha',
    'total_days':'30 Days',
    'pincode':"787876",
    'camp_name':"Summer Camp MCF Adventure Camp"
}

# Replace text fields in paragraphs
find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD STUDENT_NAME}', student_data['name'])
find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD COMPANY_NAME}', student_data['comp_name'])
find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD REG_NO}', student_data['reg_no'])
find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD CAMP_DATE}', student_data['camp_date'])
find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD PARENT_NAME}', student_data['parent_name'])
find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD ADDRESS}', student_data['address'])
find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD CITY}', student_data['city'])
find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD DISTRICT}', student_data['district'])
find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD STATE}', student_data['state'])
find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD EMAIL}', student_data['email'])
find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD CONTACT_NO}', student_data['contact_no'])
find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD WHATSAPP_NO}', student_data['whatsapp_no'])
find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD FATHERS_NO}', student_data['fathers_no'])
find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD MOTHERS_NO}', student_data['mothers_no'])
find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD BIRTH_DATE}', student_data['birth_date'])
find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD BLOOD_GRP}', student_data['blood_grp'])
find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD STANDARD}', student_data['standard'])
find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD SCHOOL_NAME}', student_data['school_name'])
find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD BATCH}', student_data['batch'])
find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD TOTAL_DAYS}', student_data['total_days'])
find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD PINCODE}', student_data['pincode'])
find_and_replace_paragraphs(doc.paragraphs, '{MERGEFIELD CAMP_NAME}', student_data['camp_name'])

# Replace text fields in tables
find_and_replace_tables(doc.tables, '{MERGEFIELD STUDENT_NAME}', student_data['name'])
find_and_replace_tables(doc.tables, '{MERGEFIELD COMPANY_NAME}', student_data['comp_name'])
find_and_replace_tables(doc.tables, '{MERGEFIELD REG_NO}', student_data['reg_no'])
find_and_replace_tables(doc.tables, '{MERGEFIELD CAMP_DATE}', student_data['camp_date'])
find_and_replace_tables(doc.tables, '{MERGEFIELD PARENT_NAME}', student_data['parent_name'])
find_and_replace_tables(doc.tables, '{MERGEFIELD ADDRESS}', student_data['address'])
find_and_replace_tables(doc.tables, '{MERGEFIELD CITY}', student_data['city'])
find_and_replace_tables(doc.tables, '{MERGEFIELD DISTRICT}', student_data['district'])
find_and_replace_tables(doc.tables, '{MERGEFIELD STATE}', student_data['state'])
find_and_replace_tables(doc.tables, '{MERGEFIELD EMAIL}', student_data['email'])
find_and_replace_tables(doc.tables, '{MERGEFIELD CONTACT_NO}', student_data['contact_no'])
find_and_replace_tables(doc.tables, '{MERGEFIELD WHATSAPP_NO}', student_data['whatsapp_no'])
find_and_replace_tables(doc.tables, '{MERGEFIELD FATHERS_NO}', student_data['fathers_no'])
find_and_replace_tables(doc.tables, '{MERGEFIELD MOTHERS_NO}', student_data['mothers_no'])
find_and_replace_tables(doc.tables, '{MERGEFIELD BIRTH_DATE}', student_data['birth_date'])
find_and_replace_tables(doc.tables, '{MERGEFIELD BLOOD_GRP}', student_data['blood_grp'])
find_and_replace_tables(doc.tables, '{MERGEFIELD STANDARD}', student_data['standard'])
find_and_replace_tables(doc.tables, '{MERGEFIELD SCHOOL_NAME}', student_data['school_name'])
find_and_replace_tables(doc.tables, '{MERGEFIELD BATCH}', student_data['batch'])
find_and_replace_tables(doc.tables, '{MERGEFIELD TOTAL_DAYS}', student_data['total_days'])
find_and_replace_tables(doc.tables, '{MERGEFIELD PINCODE}', student_data['pincode'])
find_and_replace_tables(doc.tables, '{MERGEFIELD CAMP_NAME}', student_data['camp_name'])

# Replace specific image with a new image from URL
replace_image_with_url(doc, 'old_image.png', 'https://example.com/new_image.png', width=Inches(2), height=Inches(2))

# Save the modified document in DOCX format
doc.save(f"{student_data['reg_no']}_entrance_card.docx")

# Convert DOCX to PDF
convert(f"{student_data['reg_no']}_entrance_card.docx", f"{student_data['reg_no']}_entrance_card.pdf")
