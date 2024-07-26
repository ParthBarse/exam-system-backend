from flask import Flask
from flask import request, session , make_response
from pymongo import MongoClient
from flask import Flask, request, jsonify, send_file
from flask_login import LoginManager, UserMixin, login_user, logout_user, current_user
from flask_bcrypt import Bcrypt
from flask_cors import CORS
import datetime
from datetime import datetime
import random
import json
from email.mime.text import MIMEText
import smtplib
import uuid
import re
import os
import requests
from io import BytesIO
import subprocess
from werkzeug.security import generate_password_hash, check_password_hash
import jwt
import threading
import multiprocessing
import time
import zipfile
import requests
import base64
import threading

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import subprocess

#--------------------------------------------------------------------------------

file_dir = "/home/bnbdevelopers-files/htdocs/files.bnbdevelopers.in/exam_files/"
files_url = "https://files.bnbdevelopers.in"
files_base_dir = "/home/bnbdevelopers-files/htdocs/files.bnbdevelopers.in/"
files_base_url = "https://files.bnbdevelopers.in/exam_files/"

# file_dir = "/home/mcfcamp-files/htdocs/files.mcfcamp.in/mcf_files/"
# files_url = "https://files.mcfcamp.in"
# files_base_dir = "/home/mcfcamp-files/htdocs/files.mcfcamp.in/"
# files_base_url = "https://files.mcfcamp.in/mcf_files/"

#----------------------------------------------------------------------------------



app = Flask(__name__)
CORS(app)

client = MongoClient(
    'mongodb+srv://bnbdevs:feLC7m4jiT9zrmHh@cluster0.fjnp4qu.mongodb.net/?retryWrites=true&w=majority')
app.config['MONGO_URI'] = 'mongodb+srv://bnbdevs:feLC7m4jiT9zrmHh@cluster0.fjnp4qu.mongodb.net/?retryWrites=true&w=majority'

# client = MongoClient(
#     'mongodb+srv://mcfcamp:mcf123@mcf.nyh46tl.mongodb.net/')
# app.config['MONGO_URI'] = 'mongodb+srv://mcfcamp:mcf123@mcf.nyh46tl.mongodb.net/'

app.config['SECRET_KEY'] = 'a6d217d048fdcd227661b755'
db = client['exam_system']
db2 = client['students_exam_answers']
bcrypt = Bcrypt(app)
login_manager = LoginManager(app)
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 465
app.config['MAIL_USE_SSL'] = True
app.config['MAIL_USERNAME'] = "ic2023wallet@gmail.com"
app.config['MAIL_PASSWORD'] = "irbnexpguzgxwdgx"

host = ""


# notificationFlag = True

def getNotfStat():
    settings_db = db['count_db']
    data = settings_db.find_one({"found":"2"})
    if data :
        notificationFlag=data['status']
    else:
        notificationFlag="on"
    print("Notification - ",notificationFlag)
    return notificationFlag


@app.route('/')
def hello_world():
    return 'Hello World!'


@app.route('/home')
def home():
    return 'home page'

def generate_new_receipt_no():
    count_db = db['count_db']
    c_data = count_db.find_one({"found":"1"})
    sr_no = int(c_data['sr_no'])
    count_db.update_one({"found":"1"}, {"$set": {"sr_no":int(sr_no+1)}})
    new_receipt_no = str("2024-"+str(int(sr_no+1)))
    return new_receipt_no

def set_paragraph_font(paragraph, font_name, font_size, bold):
    for run in paragraph.runs:
        font = run.font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = bold

def find_and_replace_paragraphs(paragraphs, field, replacement, specific_font=None):
    for paragraph in paragraphs:
        if field in paragraph.text:
            paragraph.text = paragraph.text.replace(field, replacement)
            if specific_font is not None:
                set_paragraph_font(paragraph, *specific_font)

def generate_certificate(doc,student_data):    
    for key, value in student_data.items():
        if key == 'NAME':
            find_and_replace_paragraphs(doc.paragraphs, f'{{MERGEFIELD {key}}}', str(value), specific_font=('Times New Roman', 34, True))
        else:
            find_and_replace_paragraphs(doc.paragraphs, f'{{MERGEFIELD {key}}}', str(value), specific_font=('Times New Roman', 14, False))
    docx_path = str(str(file_dir)+f"CERT_{student_data['seid']}.docx")
    doc.save(docx_path)
    output_path = str(str(file_dir)+f"CERT_{student_data['seid']}.pdf")
    convert_to_pdf(docx_path,output_path)

    cert_url = f"{files_base_url}CERT_{student_data['seid']}.pdf"

    students_db = db["exam_students_db"]
    students_db.update_one({"seid":student_data['seid']}, {"$set": {"cert_url":cert_url,"marks_obtained":student_data['MARKS'], "total_marks":student_data['total_marks']}})
    msg = f"""You have successfully Completed Exam : {student_data['EXAM_NAME']}.\n\nYou have scored {student_data['MARKS']} in above Exam.\nYour Exam Certificate is Attached below.\n\nThank You."""
    sub = f"Result of Exam : {student_data['EXAM_NAME']}"
    # send_email(msg, sub, data['email'])
    thread = threading.Thread(target=send_email_attachments, args=(msg, sub, student_data['email'],[output_path],))
    # generate_certificate(doc,student_data)
    thread.start()


def convert_to_pdf(docx_file, pdf_file):
    try:
        subprocess.run(['unoconv', '--output', pdf_file, '--format', 'pdf', docx_file], check=True)
        print(f"Conversion successful: {docx_file} -> {pdf_file}")
    except subprocess.CalledProcessError as e:
        print(f"Error during conversion: {e}")

def calculate_marks(correct_answers, student_answers):
    obtained_marks = 0
    total_marks = 0
    for correct in correct_answers:
        total_marks += int(correct["marks"])
        question_id = correct["question_id"]
        correct_options = set(correct["correctOptions"])
        marks = int(correct["marks"])
        
        for student in student_answers:
            if student["question_id"] == question_id:
                student_options = set(student["answers"])
                if student_options == correct_options:
                    obtained_marks += marks
    return obtained_marks, total_marks

def calculate_result(exam_id,seid):
    questions_db = db["questions_db"]
    students_ans_db = db2[seid]

    correct_answers_raw = questions_db.find({"exam_id":exam_id},{"_id":0})
    student_answers = students_ans_db.find({},{"_id":0})

    correct_answers = [
    {
        **item,
        'correctOptions': json.loads(item['correctOptions']) if 'correctOptions' in item else item['correctOptions']
    }
    for item in correct_answers_raw]

    correct_answers = list(correct_answers)
    student_answers = list(student_answers)

    obtained_marks,total_marks = calculate_marks(correct_answers, student_answers)
    print(f"Total Marks: {obtained_marks}")

    student_db = db["exam_students_db"]
    student_data = student_db.find_one({"seid":seid},{"_id":0})

    doc = Document('result.docx')
    student_data = {
        'EXAM_NO': exam_id,
        'NAME': str(student_data['first_name']+" "+student_data['last_name']),
        'EXAM_NAME': student_data['exam_name'],
        'MARKS':obtained_marks,
        'seid':seid,
        'total_marks': total_marks,
        'email':student_data['email'],
    }
    thread = threading.Thread(target=generate_certificate, args=(doc,student_data,))
    # generate_certificate(doc,student_data)
    thread.start()




#-----------------------------------------------------------------------------------------

#---------------------- System Synchronization Module ------------------------------------

file_directory = file_dir

def save_file(file, uid):
    try:
        # Get the file extension from the original filename
        original_filename = file.filename
        _, file_extension = os.path.splitext(original_filename)

        # Generate a unique filename using UUID and append the original file extension
        filename = str(uuid.uuid4()) + file_extension

        file_path = os.path.join(file_directory, uid, filename)
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        file.save(file_path)

        return f'{files_base_url}{uid}/{filename}'
    except Exception as e:
        raise e
    
def save_file2(file_data, sid, filename):
    # Create a subdirectory for the specific session id if it doesn't exist
    session_path = os.path.join(file_directory, sid)
    if not os.path.exists(session_path):
        os.makedirs(session_path)
    
    # Save the file
    file_path = os.path.join(session_path, filename)
    with open(file_path, "wb") as f:
        f.write(file_data)
    
    # Return the URL to access the file (for simplicity, we return the file path here)
    return f'{files_base_url}{sid}/{filename}'

#-------------- Supporting Functions Start ----------------



#-------------- Supporting Functions End ----------------



#------------------------------------------------------------------------------------------

def sendSMS(msg,phn):
    notifyFlag = getNotfStat()
    if notifyFlag == "off":
        phn=''
    # phn="8793015610"
    if msg and phn:
        url = "http://msg.msgclub.net/rest/services/sendSMS/sendGroupSms"
        msg_text = msg
        phn_no = phn
        querystring = {"AUTH_KEY":"2b4186d8fc21f47949e7f5e92b56390","message":msg_text,"senderId":"MCFCMP","routeId":"1","mobileNos":phn_no,"smsContentType":"english"}
        headers = {'Cache-Control': "no-cache"}
        response = requests.request("GET", url, headers=headers, params=querystring)
        print(response.text)
        return 0
    else:
        return 1
    

def send_wp(sms_content, mobile_numbers, file_paths=[]):
    notifyFlag = getNotfStat()
    if notifyFlag == "off":
        mobile_numbers=''
    # mobile_numbers="8793015610"
    if len(file_paths)>1:
            file_paths.append("THINGS_TO_BRING.pdf")
    api_url = "http://msg.msgclub.net/rest/services/sendSMS/sendGroupSms"
    auth_key = "2b4186d8fc21f47949e7f5e92b56390"
    route_id = "21"
    sender_id = "9604992000"
    sms_content_type = "english"
    payload = {
        "smsContent": sms_content,
        "routeId": route_id,
        "mobileNumbers": mobile_numbers,
        "senderId": sender_id,
        "smsContentType": sms_content_type
    }
    headers = {
        "AUTH_KEY": auth_key,
        "Content-Type": "application/json"
    }

    payload2 = {
        "smsContent": "",
        "routeId": route_id,
        "mobileNumbers": mobile_numbers,
        "senderId": sender_id,
        "smsContentType": sms_content_type
    }

    # Add file data if file_path is provided
    if file_paths:
        if len(file_paths) == 1:
            filedata_encoded = encode_file_to_base64(file_paths[0])
            if filedata_encoded:
                payload["filename"] = file_paths[0].split('/')[-1]  # Extract filename from path
                payload["filedata"] = filedata_encoded
            else:
                print(f"Error: Unable to encode {file_path} to Base64")
                return 1
            response = requests.post(api_url, json=payload, headers=headers)
            if response.status_code == 200:
                response_json = response.json()
                if 'response' in response_json:
                    print("Send WP")
                    return 0
                else:
                    return 1
            else:
                return 1
        elif len(file_paths) > 1:
            for file_path in file_paths:
                filedata_encoded = encode_file_to_base64(file_path)
                if filedata_encoded:
                    payload2["filename"] = file_path.split('/')[-1]  # Extract filename from path
                    payload2["filedata"] = filedata_encoded
                else:
                    print(f"Error: Unable to encode {file_path} to Base64")
                response = requests.post(api_url, json=payload2, headers=headers)
                if response.status_code == 200:
                    response_json = response.json()
                    if 'response' in response_json:
                        print("Send WP")
                    else:
                        print("Error")
    response = requests.post(api_url, json=payload, headers=headers)
    if response.status_code == 200:
        response_json = response.json()
        if 'response' in response_json:
            print("Send WP")
            return 0
        else:
            return 1
    else:
        return 1

def encode_file_to_base64(file_path):
    try:
        with open(file_path, "rb") as file:
            filedata = file.read()
            filedata_encoded = base64.b64encode(filedata).decode('utf-8')
            return filedata_encoded
    except Exception as e:
        print(f"Error encoding file to Base64: {str(e)}")
        return None
    

def send_email(msg, sub, mailToSend):
    notifyFlag = getNotfStat()
    if notifyFlag == "off":
        mailToSend=''
    # mailToSend = "parthbarse72@gmail.com"
    try:
        # Send the password reset link via email
        # sender_email = "mcfcamp@gmail.com"
        # smtp_server = smtplib.SMTP("smtp.gmail.com", 587)
        # smtp_server.ehlo()
        # smtp_server.starttls()
        # smtp_server.login("mcfcamp@gmail.com", "meyv ghup onbl fqhu")

        sender_email = "partbarse92@gmail.com"
        smtp_server = smtplib.SMTP("smtp.gmail.com", 587)
        smtp_server.ehlo()
        smtp_server.starttls()
        smtp_server.login("partbarse92@gmail.com", "tdmz qbky qlzc urvg")

        message_text = msg
        message = MIMEText(message_text)
        message["Subject"] = sub
        message["From"] = sender_email
        message["To"] = mailToSend

        smtp_server.sendmail(sender_email, mailToSend, message.as_string())
        print(mailToSend)
        print("Send Mail")
        smtp_server.quit()
        return 0
    except Exception as e:
        print(str(e))
        return 1
    


import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

def send_email_attachments(msg, sub, mailToSend, files=[]):
    notifyFlag = getNotfStat()
    if notifyFlag == "off":
        mailToSend=''
    # mailToSend = "parthbarse72@gmail.com"
    try:
        if len(files)>1:
            files.append("THINGS_TO_BRING.pdf")
        sender_email = "partbarse92@gmail.com"
        smtp_server = smtplib.SMTP("smtp.gmail.com", 587)
        smtp_server.ehlo()
        smtp_server.starttls()
        smtp_server.login("partbarse92@gmail.com", "tdmz qbky qlzc urvg")

        # Create a multipart message
        message = MIMEMultipart()
        message["Subject"] = sub
        message["From"] = sender_email
        message["To"] = mailToSend

        # Attach message body
        message.attach(MIMEText(msg, "plain"))

        # Attach files
        for file_path in files:
            with open(file_path, "rb") as attachment:
                part = MIMEBase("application", "octet-stream")
                part.set_payload(attachment.read())

            # Encode file in ASCII characters to send by email
            encoders.encode_base64(part)

            # Add header as key/value pair to attachment part
            part.add_header(
                "Content-Disposition",
                f"attachment; filename= {file_path}",
            )

            # Attach the attachment to the message
            message.attach(part)

        smtp_server.sendmail(sender_email, mailToSend, message.as_string())
        print(mailToSend)
        print("Send Mail")
        smtp_server.quit()
        return 0
    except Exception as e:
        print(str(e))
        return 1

# ------------------------------------------------------------------------------------------------------------

@app.route('/addExam', methods=['POST'])
def add_exam():
    try:
        data = request.form
        print("Data Recieved : ",data)
        print(data.get("exam_name"))

        # Generate a unique ID for the camp using UUID
        exam_id = str(uuid.uuid4().hex)

        exam = {
            "exam_id": exam_id,
            "exam_name": data["exam_name"].strip(),
            "exam_duration": data["exam_duration"],
            "exam_date": data["exam_date"],
            "exam_description": data["exam_description"],
            "exam_status" : data["exam_status"],
        }

        # Store the camp information in the MongoDB collection
        exams_db = db["exams_db"]
        exams_db.insert_one(exam)

        return jsonify({"message": "Exam added successfully", "exam_id": exam_id})

    except ValueError as ve:
        return jsonify({"error": str(ve)}), 400  # Bad Request

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/updateExam', methods=['PUT'])
def update_exam():
    try:
        data = request.form

        # Check if exam_id is provided
        if 'exam_id' not in data:
            raise ValueError("Missing 'exam_id' in the request.")

        # Find the exam based on exam_id
        exams_db = db["exams_db"]
        exam = exams_db.find_one({"exam_id": data['exam_id']})

        if not exam:
            return jsonify({"error": f"No exam found with exam_id: {data['exam_id']}"}), 404  # Not Found

        # Update the exam information with the received data
        for key, value in data.items():
            if key != 'exam_id':
                # If the value is provided, update the field; otherwise, keep the existing value
                if value:
                    exam[key] = value
                    if exam['exam_status'] == "on":
                        exam["exam_status"] = "Active"
                    else:
                        exam['exam_status'] = "Inactive"

        # Update the exam in the database
        exams_db.update_one({"exam_id": data['exam_id']}, {"$set": exam})

        return jsonify({"message": f"Exam with exam_id {data['exam_id']} updated successfully"})

    except ValueError as ve:
        return jsonify({"error": str(ve)}), 400  # Bad Request

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/getAllExams', methods=['GET'])
def get_all_exams():
    try:
        exams_db = db["exams_db"]
        exams = exams_db.find({}, {"_id": 0})  # Exclude the _id field from the response

        # Convert the cursor to a list of dictionaries for easier serialization
        exam_list = list(exams)

        return jsonify({"exams": exam_list})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/getAllExamsActive', methods=['GET'])
def get_all_exams_active():
    try:
        exams_db = db["exams_db"]
        exams = exams_db.find({"exam_status":"Active"}, {"_id": 0})  # Exclude the _id field from the response

        # Convert the cursor to a list of dictionaries for easier serialization
        exam_list = list(exams)

        return jsonify({"exams": exam_list})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/getAllBatches', methods=['GET'])
def get_all_batches():
    try:
        batches_db = db["batches_db"]
        batches = batches_db.find({}, {"_id": 0})  # Exclude the _id field from the response

        # Convert the cursor to a list of dictionaries for easier serialization
        batches_list = list(batches)

        return jsonify({"camps": batches_list})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/getExam', methods=['GET'])
def get_exam():
    try:
        # Get the exam_id from request parameters
        exam_id = request.args.get('exam_id')

        if not exam_id:
            return jsonify({"error": "Missing 'exam_id' parameter in the request."}), 400  # Bad Request

        # Find the exam based on exam_id
        exams_db = db["exams_db"]
        exam = exams_db.find_one({"exam_id": exam_id}, {"_id": 0})  # Exclude the _id field from the response

        if not exam:
            return jsonify({"error": f"No exam found with exam_id: {exam_id}"}), 404  # Not Found

        return jsonify({"exam": exam})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/deleteExam', methods=['DELETE'])
def delete_exam():
    try:
        # Get the exam_id from request parameters
        exam_id = request.args.get('exam_id')

        if not exam_id:
            return jsonify({"error": "Missing 'exam_id' parameter in the request."}), 400  # Bad Request

        # Find the exam based on exam_id
        exams_db = db["exams_db"]
        exam = exams_db.find_one({"exam_id": exam_id})

        if not exam:
            return jsonify({"error": f"No exam found with exam_id: {exam_id}"}), 404  # Not Found

        # Delete the exam from the database
        exams_db.delete_one({"exam_id": exam_id})

        return jsonify({"message": f"Exam with exam_id {exam_id} deleted successfully"})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/addQuestion', methods=['POST'])
def add_question():
    try:
        data = request.form
        data = dict(data)
        print(data)

        # Generate a unique ID for the batch using UUID
        question_id = str(uuid.uuid4().hex)

        data['question_id'] = question_id

        # Store the batch information in the MongoDB collection
        question_db = db["questions_db"]
        question_db.insert_one(data)

        return jsonify({"message": "Question added successfully", "question_id": question_id})

    except ValueError as ve:
        return jsonify({"error": str(ve)}), 400  # Bad Request

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error

@app.route('/updateQuestion', methods=['POST'])
def update_question():
    try:
        data = request.form

        # Check if question_id is provided
        if 'question_id' not in data:
            raise ValueError("Missing 'question_id' in the request.")

        # Find the question based on question_id
        questions_db = db["questions_db"]
        question = questions_db.find_one({"question_id": data['question_id']})

        if not question:
            return jsonify({"error": f"No question found with question_id: {data['question_id']}"}), 404  # Not Found

        # Update the question information with the received data
        for key, value in data.items():
            if key != 'question_id':
                # If the value is provided, update the field; otherwise, keep the existing value
                if value:
                    question[key] = int(value) if key == 'question_intake' else value

        # Update the question in the database
        questions_db.update_one({"question_id": data['question_id']}, {"$set": question})

        return jsonify({"message": f"question with question_id {data['question_id']} updated successfully", "camp_id":question['camp_id']})

    except ValueError as ve:
        return jsonify({"error": str(ve)}), 400  # Bad Request

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/getQuestions', methods=['GET'])
def get_questions():
    try:
        # Get the exam_id from request parameters
        exam_id = request.args.get('exam_id')

        if not exam_id:
            return jsonify({"error": "Missing 'exam_id' parameter in the request."}), 400  # Bad Request

        # Find Question based on exam_id
        question_db = db["questions_db"]
        question = question_db.find({"exam_id": exam_id}, {"_id": 0})  # Exclude the _id field from the response

        # Convert the cursor to a list of dictionaries for easier serialization
        question_list = list(question)

        return jsonify({"question": question_list})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error


@app.route('/getQuestion', methods=['GET'])
def get_question():
    try:
        # Get the question_id from request parameters
        question_id = request.args.get('question_id')

        if not question_id:
            return jsonify({"error": "Missing 'question_id' parameter in the request."}), 400  # Bad Request

        # Find the question based on question_id
        questions_db = db["questions_db"]
        question = questions_db.find_one({"question_id": question_id}, {"_id": 0})  # Exclude the _id field from the response

        if not question:
            return jsonify({"error": f"No question found with question_id: {question_id}"}), 404  # Not Found

        return jsonify({"question": question})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/deleteQuestion', methods=['DELETE'])
def delete_question():
    try:
        # Get the question_id from request parameters
        question_id = request.args.get('question_id')

        if not question_id:
            return jsonify({"error": "Missing 'question_id' parameter in the request."}), 400  # Bad Request

        # Find the question based on question_id
        questions_db = db["questions_db"]
        question = questions_db.find_one({"question_id": question_id})

        if not question:
            return jsonify({"error": f"No question found with question_id: {question_id}"}), 404  # Not Found

        # Delete the question from the database
        questions_db.delete_one({"question_id": question_id})

        return jsonify({"message": f"question with question_id {question_id} deleted successfully"})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/uploadFile', methods=['POST'])
def upload_file():
    try:
        # Check if 'file' and 'sid' parameters are present in the form data
        if 'file' not in request.files:
            return jsonify({'error': 'Missing parameters: file',"success":False}), 400

        uploaded_file = request.files['file']
        sid = "All_Files"

        # Check if the file is an allowed type (e.g., image or pdf)
        allowed_extensions = {'png', 'jpg', 'jpeg', 'gif', 'pdf'}
        if (
            '.' in uploaded_file.filename
            and uploaded_file.filename.rsplit('.', 1)[1].lower() not in allowed_extensions
        ):
            return jsonify({'error': 'Invalid file type. Only allowed: png, jpg, jpeg, gif, pdf.',"success":False}), 400

        # Save the file and get the URL
        file_url = save_file(uploaded_file, sid)

        return jsonify({'message': 'File stored successfully.', 'file_url': file_url,"success":True}), 200

    except Exception as e:
        return jsonify({'error': str(e),"success":False}), 500


@app.route('/uploadFile2', methods=['POST'])
def upload_file2():
    try:
        # Check if 'file' is in the request
        if 'file' not in request.form:
            return jsonify({'error': 'Missing parameters: file', "success": False}), 400

        file_data = request.form['file']
        seid = request.form['seid']
        sid = "All_Files"  # Replace this with dynamic session ID if available

        # Decode the base64 file data
        if file_data.startswith('data:image/jpeg;base64,'):
            file_data = base64.b64decode(file_data.split(',')[1])
            filename = str(seid)+".jpg"
        else:
            return jsonify({'error': 'Invalid file format.', "success": False}), 400

        # Save the file and get the URL
        file_url = save_file2(file_data, sid, filename)

        return jsonify({'message': 'File stored successfully.', 'file_url': file_url, "success": True}), 200

    except Exception as e:
        return jsonify({'error': str(e), "success": False}), 500
    

@app.route('/registerStudentExam', methods=['POST'])
def register_student_exam():
    try:
        data = request.form
        data = dict(data)
        # Generate a unique ID for the student using UUID
        seid = str(uuid.uuid4().hex)
        data["seid"] = seid
        data["status"] = "not-start"
        exam_students_db = db["exam_students_db"]
        exam_students_db.insert_one(data)
        exam_link = f"https://test-view-exam.bnbdevelopers.in/screening?id={data['exam_id']}&seid={seid}"
        msg = f"""You have successfully registered for {data['exam_name']}.\n\nYou can use below link to directly access the Exam :\n{exam_link}\n\nIf you face any Problem with Above link then You can use https://test-view-exam.bnbdevelopers.in/login \nUserID : {seid}\nPassword : {data['phn']}\n\nPlease Don't Share any of the Above Links or Credentials with Anyone.\n\n Thank You."""
        sub = "Registration Successful and Exam Link."
        # send_email(msg, sub, data['email'])
        thread = threading.Thread(target=send_email, args=(msg, sub, data['email'],))
        # generate_certificate(doc,student_data)
        thread.start()
        return jsonify({"message": "Student registered successfully", "seid": seid})

    except ValueError as ve:
        return jsonify({"error": str(ve)}), 400  # Bad Request

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/getAllExamStudents', methods=['GET'])
def get_all_exam_students():
    try:
        students_db = db["exam_students_db"]
        students = students_db.find({}, {"_id": 0})  # Exclude the _id field from the response

        # Convert the cursor to a list of dictionaries for easier serialization
        exam_list = list(students)

        return jsonify({"students": exam_list})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/getExamStudent', methods=['GET'])
def get_exam_student():
    try:
        seid = request.args.get("seid")
        students_db = db["exam_students_db"]
        student = students_db.find_one({"seid":seid}, {"_id": 0})  # Exclude the _id field from the response

        return jsonify({"student": student})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    

@app.route('/deleteExamStudent', methods=['DELETE'])
def delete_exam_student():
    try:
        # Get the seid from request parameters
        seid = request.args.get('seid')

        if not seid:
            return jsonify({"error": "Missing 'seid' parameter in the request."}), 400  # Bad Request

        # Find the seis based on seid
        exam_students_db = db["exam_students_db"]
        student = exam_students_db.find_one({"seid":  seid})

        if not student:
            return jsonify({"error": f"No student found with  seid: { seid}"}), 404  # Not Found

        # Delete the exam from the database
        exam_students_db.delete_one({"seid": seid})

        return jsonify({"message": f"Student with seid {seid} deleted successfully"})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    

@app.route('/submitAnswers', methods=['POST'])
def submit_answers():
    try:
        data = request.json
        # data = dict(data)
        students_exam_answers_db = db2[data['seid']]

        if (students_exam_answers_db.find_one({"question_id":data['question_id']})):
            students_exam_answers_db.update_one({"question_id":data['question_id']}, {"$set": {"answers":data['answers']}})
            return jsonify({"message": "Answer submitted successfully", "question_id": data['question_id']})
        else:
            students_exam_answers_db.insert_one(data)
            return jsonify({"message": "Answer submitted successfully", "question_id": data['question_id']})

    except ValueError as ve:
        return jsonify({"error": str(ve)}), 400  # Bad Request

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/submitExam', methods=['POST'])
def submit_exam():
    try:
        data = request.json
        # data = dict(data)
        exam_students_db = db["exam_students_db"]

        student = exam_students_db.find_one({"seid":data['seid']})

        if (student):
            exam_students_db.update_one({"seid":data['seid']}, {"$set": {"status":"submitted"}})
            calculate_result(data['exam_id'],data['seid'])
            return jsonify({"message": "Exam submitted successfully", "exam_id": data['exam_id']})
        else:
            return jsonify({"message": "Exam Not submitted successfully"}),401

    except ValueError as ve:
        return jsonify({"error": str(ve)}), 400  # Bad Request

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    

@app.route('/checkStudentExamStatus', methods=['GET'])
def check_student_exam_status():
    try:
        seid = request.args.get("seid")
        exam_students_db = db["exam_students_db"]

        student = exam_students_db.find_one({"seid":seid})

        if (student):
            return jsonify({"status": student['status']})
        else:
            return jsonify({"message": "Student Not Found"}),401

    except ValueError as ve:
        return jsonify({"error": str(ve)}), 400  # Bad Request

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/startExam', methods=['POST'])
def start_exam():
    try:
        data = request.form
        seid = data['seid']
        captured_image = data['captured_image']

        exam_students_db = db["exam_students_db"]
        exams_db = db["exams_db"]

        student = exam_students_db.find_one({"seid":seid})

        if (student):
            exam = exams_db.find_one({"exam_id":student['exam_id']})
            minutes_int = int(exam['exam_duration'])
            seconds = minutes_int * 60  # 1 minute = 60 seconds
            exam_students_db.update_one({"seid":seid}, {"$set": {"captured_image":captured_image, "status" : "started", "remaining_duration":seconds}})
            return jsonify({"status": "started", "success":True})
        else:
            return jsonify({"message": "Student Not Found"}),401

    except ValueError as ve:
        return jsonify({"error": str(ve)}), 400  # Bad Request

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    

@app.route('/updateTimer', methods=['GET'])
def update_timer():
    try:

        seid = request.args.get("seid")
        remaining_duration = request.args.get("remaining_duration")

        exam_students_db = db["exam_students_db"]

        student = exam_students_db.find_one({"seid":seid})

        if (student):
            exam_students_db.update_one({"seid":seid}, {"$set": {"remaining_duration":remaining_duration}})
            return jsonify({"status": "started", "success":True})
        else:
            return jsonify({"message": "Student Not Found"}),401

    except ValueError as ve:
        return jsonify({"error": str(ve)}), 400  # Bad Request

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    

@app.route('/loginStudent', methods=['POST'])
def login_student():
    try:
        data = request.get_json()
        seid = data["username"]
        password = data["password"]

        exam_students_db = db["exam_students_db"]

        student = exam_students_db.find_one({"seid":seid})

        if (student):
            if(str(student['phn']) == str(password)):
                return jsonify({"message": "Authenticated", "success":True, "token":seid, "seid":seid, "exam_id":student['exam_id']})
            else:
                return jsonify({"message": "Not Authenticated", "success":False}), 401
        else:
            return jsonify({"message": "Student Not Found"}),404

    except ValueError as ve:
        return jsonify({"error": str(ve)}), 400  # Bad Request

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    

    



if __name__ == '__main__':
    app.run(host="0.0.0.0",port=8088)





