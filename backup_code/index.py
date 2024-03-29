from flask import Flask
from flask import request, session , make_response
from pymongo import MongoClient
from flask import Flask, request, jsonify, send_file
from flask_login import LoginManager, UserMixin, login_user, logout_user, current_user
from flask_bcrypt import Bcrypt
from flask_cors import CORS
import datetime
from datetime import datetime
import random
import json
from email.mime.text import MIMEText
import smtplib
import uuid
import re
import os
from docx import Document
from docx.shared import Inches, Pt
import requests
from io import BytesIO
from docx.shared import Pt
import subprocess
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from werkzeug.security import generate_password_hash, check_password_hash
import jwt
import threading
import time
import zipfile
import requests
import base64

#--------------------------------------------------------------------------------

# file_dir = "/home/bnbdevelopers-files/htdocs/files.bnbdevelopers.in/mcf_files/"
# files_url = "https://files.bnbdevelopers.in"
# files_base_dir = "/home/bnbdevelopers-files/htdocs/files.bnbdevelopers.in/"
# files_base_url = "https://files.bnbdevelopers.in/mcf_files/"

file_dir = "/home/mcfcamp-files/htdocs/files.mcfcamp.in/mcf_files/"
files_url = "https://files.mcfcamp.in"
files_base_dir = "/home/mcfcamp-files/htdocs/files.mcfcamp.in/"
files_base_url = "https://files.mcfcamp.in/mcf_files/"

#----------------------------------------------------------------------------------



app = Flask(__name__)
CORS(app)

# client = MongoClient(
#     'mongodb+srv://bnbdevs:feLC7m4jiT9zrmHh@cluster0.fjnp4qu.mongodb.net/?retryWrites=true&w=majority')
# app.config['MONGO_URI'] = 'mongodb+srv://bnbdevs:feLC7m4jiT9zrmHh@cluster0.fjnp4qu.mongodb.net/?retryWrites=true&w=majority'

client = MongoClient(
    'mongodb+srv://mcfcamp:mcf123@mcf.nyh46tl.mongodb.net/')
app.config['MONGO_URI'] = 'mongodb+srv://mcfcamp:mcf123@mcf.nyh46tl.mongodb.net/'

app.config['SECRET_KEY'] = 'a6d217d048fdcd227661b755'
db = client['mcf_db']
bcrypt = Bcrypt(app)
login_manager = LoginManager(app)
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 465
app.config['MAIL_USE_SSL'] = True
app.config['MAIL_USERNAME'] = "ic2023wallet@gmail.com"
app.config['MAIL_PASSWORD'] = "irbnexpguzgxwdgx"

host = ""


@app.route('/')
def hello_world():
    return 'Hello World!'


@app.route('/home')
def home():
    return 'home page'


#-----------------------------------------------------------------------------------------

#---------------------- System Synchronization Module ------------------------------------



#-------------- Supporting Functions Start ----------------

def generate_3_digit_number(num):
    num_str = str(num)
    padded_num = num_str.zfill(3)
    return padded_num

def sac_table_generator(batch_id, camp_id, intake):
    sac_table = {
        "batch_id":batch_id,
        "camp_id":camp_id,
    }
    intake = int(intake)
    for i in range(1, intake):
        num = generate_3_digit_number(i)
        sac_table[num] = "-"
    sac_table_db = db["sac_table_db"]
    sac_table_data = sac_table_db.find_one({"batch_id":batch_id})
    if sac_table_data:
        return 1
    else:
        sac_table_db.insert_one(sac_table)
        return 0
    

#-------------- Supporting Functions End ----------------

def sa_module(batch_id, sid):
    pass

def sac(batch_id, sr):
    pass



#------------------------------------------------------------------------------------------

#------------------------------------------------------------------------------------------

def sync_data(original_sid):
    students_db = db["students_db"]
    data = students_db.find_one({"sid":original_sid})
    
    batches_db = db["batches_db"]
    batch = batches_db.find_one({"batch_id":data.get("batch_id")}, {"_id":0})
    camps_db = db["camps_db"]
    camp = camps_db.find_one({"camp_id":data.get("camp_id")}, {"_id":0})
    camp_name = camp["camp_name"]
    camp_short = camp_name.split(" ")[-1].replace("(", "").replace(")", "")
    sid=original_sid
    medical_cert_url = ""
    company = ""

    # Calculate age based on the provided date of birth
    age = calculate_age(data["dob"])
    if age>=7 and age<=11 and data.get("gender").lower() == "male":
        company = "ALPHA"
    elif age>=12 and age<=16 and data.get("gender").lower() == "male":
        company = "BRAVO"
    elif age>=17 and age<=21 and data.get("gender").lower() == "male":
        company = "DELTA"
    elif age>=7 and age<=11 and data.get("gender").lower() == "female":
        company = "CHARLEY"
    elif age>=12 and age<=16 and data.get("gender").lower() == "female":
        company = "ECO"
    elif age>=17 and age<=21 and data.get("gender").lower() == "female":
        company = "FOXFORD"

    isCountInc = False
    if batch:
        if int(batch["students_registered"]) <= int(batch["batch_intake"]):
            sr_no = int(int(batch["students_registered"]))
            start_date = batch["start_date"]
            year = start_date[-2:]
            day = start_date[0:2]
            batch_name = batch["batch_name"].replace(" ", "")

            company_sf = str(company[0])+"C"
            days = str(batch['duration'])+"D"

            sid = str(camp_short)+str(year)+str(days)+str(batch_name)+str(company_sf)+str(sr_no)

            stud = students_db.find_one({"sid":sid})
            if sid != original_sid:
                while stud:
                    sr_no = int(int(batch["students_registered"])+1)
                    isCountInc = True
                    start_date = batch["start_date"]
                    year = start_date[-2:]
                    day = start_date[0:2]
                    batch_name = batch["batch_name"].replace(" ", "")

                    company_sf = str(company[0])+"C"
                    days = str(batch['duration'])+"D"

                    sid = str(camp_short)+str(year)+str(days)+str(batch_name)+str(company_sf)+str(sr_no)
                    stud = students_db.find_one({"sid":sid})

            document_med_path = 'medical_certificate.docx'

            field_values = {
                'CADET_NAME': str(data["first_name"].upper()+" "+data["last_name"].upper()),
                'LOC':  str(data["district"]+", "+data["state"]),
                'DOB':  str(data["dob"]),
                '121212':  str("__________"),
                'C_NAME':  str(camp_name),
                'DATE':  str(start_date),
                'BATCH':  str(batch_name),
                'sid': sid
            }
            replace_fields_in_document_med(document_med_path, field_values)

            # Load the document template
            doc1 = Document('visit_card.docx')

            # Sample student_data
            student_data1 = {
                'CADET_NAME': str(data["first_name"].upper()+" "+data["last_name"].upper()),
                'CAMP_NAME': str(camp_name),
                'BATCH_NO': str(batch_name),
                'ADDRESS': data["address"],
                'CONTACT': data["phn"],
                'WHATSAPP_NO': data["wp_no"],
                'CAMP_DATE':start_date,
                'REG_NO':sid,
                'PICKUP_POINT':data['pick_up_point'],
            }

            for key, value in student_data1.items():
                find_and_replace_paragraphs_visiting_card(doc1.paragraphs, f'{{MERGEFIELD {key}}}', str(value))

            try:
                cadet_photo_url = data["cadetPhoto"]
                cadet_photo_path = cadet_photo_url.replace(files_url,files_base_dir)

                image_url_guardian = data["parentGurdianPhoto"]
                image_path_guardian = image_url_guardian.replace(files_url,files_base_dir)
                
                replace_image_in_cell(doc1, table_index=0, row_index=0, column_index=3, image_path=cadet_photo_path)
                replace_image_in_cell(doc1, table_index=0, row_index=0, column_index=4, image_path=image_path_guardian)

                doc1.save(str(str(file_dir)+f"{sid}_visit_card.docx"))

                convert_to_pdf(str(str(file_dir)+f"{sid}_visit_card.docx"), str(str(file_dir)+f"{sid}_visit_card.pdf"))



                doc = Document('admission_form_new_format.docx')
                student_data1 = {
                "REG_NO": sid,
                "FIRST_NAME": str(data["first_name"].upper()),
                "MIDDLE_NAME": str(data["middle_name"].upper()),
                "LAST_NAME": str(data["last_name"].upper()),
                "EMAIL_ID": str(data["email"]),
                "CONTACT_NO": str(data["phn"]),
                "DATE_OF_BIRTH": str(data["dob"]),
                "ADDRESS": str(data["address"]),
                "HOW_YOU_GOT_TO_KNOW": str(data["how_you_got_to_know"]),
                "EMPLOYEE_WHO_REACHED_OUT": str(data["employee_who_reached_out_to_you"]),
                "DISTRICT": str(data["district"]),
                "STATE": data["state"],
                "PINCODE": data["pincode"],
                "PICKUP_POINT": data["pick_up_point"],
                "BLOOD_GROUP": data["blood_group"],
                "SCHOOL_NAME": data["school_name"],
                "GENDER": data["gender"],
                "STANDARD": data["standard"],
                "WHATSAPP_NO": data["wp_no"],
                "PARENT_NAME":data["middle_name"],
                "CAMP_NAME":camp_name,
                "CAMP_DATE":batch["start_date"],
                "CAMP_DAYS":batch["duration"]
                }
                for key, value in student_data1.items():
                        find_and_replace_tables_admission_form(doc.tables, f'{{MERGEFIELD {key}}}', str(value))

                image_path_sign_url = data["cadetSign"]
                image_path_sign = image_path_sign_url.replace(files_url,files_base_dir)

                image_url_guardian_sign = data["parentGurdianSign"]
                image_path_guardian_sign = image_url_guardian_sign.replace(files_url,files_base_dir)

                replace_image_in_cell_admission_form(doc, table_index=0, row_index=25, column_index=1, image_path=cadet_photo_path,w=1.4,h=1.6)
                replace_image_in_cell_admission_form(doc, table_index=0, row_index=25, column_index=12, image_path=image_path_sign,w=1.8,h=1.0)
                replace_image_in_cell_admission_form(doc, table_index=1, row_index=16, column_index=1, image_path=image_path_guardian,w=1.4,h=1.6)
                replace_image_in_cell_admission_form(doc, table_index=1, row_index=17, column_index=6, image_path=image_path_guardian_sign,w=1.8,h=1.0)

                doc.save(str(str(file_dir)+f"{sid}_admission_form.docx"))

                convert_to_pdf(str(str(file_dir)+f"{sid}_admission_form.docx"), str(str(file_dir)+f"{sid}_admission_form.pdf"))

                final_status = data['status']
                if float(data['total_amount_paid']) >= float(data['total_amount_payable']):
                    camp_id = data['camp_id']
                    camp_db = db["camps_db"]
                    camp_data = camp_db.find_one({"camp_id":camp_id}, {"_id":0})
                    batch_id = data['batch_id']
                    batch_db = db["batches_db"]
                    batch_data = batch_db.find_one({"batch_id":batch_id}, {"_id":0})
                    payment_db = db["all_payments"]
                    payment_data = payment_db.find({"sid":data['sid']}, {"_id":0})
                    receipt_nos = ""
                    payment_data = list(payment_data)
                    print(len(payment_data))
                    for receipt in payment_data:
                        receipt_nos = str(receipt_nos + str(str(receipt['receipt_no'])+ " , "))
                        print(receipt_nos)
                    receipt_nos = str(receipt_nos)

                    student_data_1 = {
                            'CADET_NAME': str(data["first_name"].upper()+" "+data["last_name"].upper()),
                            'REGNO': sid,
                            'RANK': 'CDT',
                            'C_NAME': camp_data['camp_name'],
                            'C_BATCH': batch_data['batch_name'],
                            'C_DAYS': batch_data["duration"],
                            'COMP_N': data['company'],
                            'C_DATE': batch_data['start_date'],
                            'PICKPT': data["pick_up_point"],
                            'PICK_TIME': '',
                            'EMP_NAME':  data["employee_who_reached_out_to_you"],
                            'GAR_NAME': data["middle_name"],
                            'ADDRESS': data["address"],
                            'CITY': data["pick_up_city"],
                            'DISTRICT': data["district"],
                            'STATE': data['state'],
                            'PINCODE': data["pincode"],
                            'EMAIL': data["email"],
                            'C_NUM': str(data["phn"]),
                            'WP_NUM': data.get("wp_no", ""),
                            'FATHER_NUM': '',
                            'MOTHER_NUM': '',
                            'DOB': data["dob"],
                            'BLOOD_GROUP':  data.get("blood_group", ""),
                            'STD': data.get("standard", ""),
                            'SCHOOL': data.get("school_name", ""),
                            'FEE_PAID': data['total_amount_paid'],
                            'BALANCE': float(data['total_amount_payable']) - float(data['total_amount_paid']),
                            'RECEIPT_NUM': receipt_nos,
                            'DATE': '',
                            'TIME': ''
                        }

                    doc = Document('mcf_entrance_card.docx')

                    # Replace text fields in paragraphs
                    find_and_replace_paragraphs_entrance_card(doc.paragraphs, '{MERGEFIELD CADET_NAME}', student_data_1['CADET_NAME'])
                    find_and_replace_paragraphs_entrance_card(doc.paragraphs, '{MERGEFIELD DATE}', student_data_1['DATE'])
                    find_and_replace_paragraphs_entrance_card(doc.paragraphs, '{MERGEFIELD TIME}', student_data_1['TIME'])

                    for key, value in student_data_1.items():
                            find_and_replace_tables_entrance_card(doc.tables, f'{{MERGEFIELD {key}}}', str(value))

                    try:
                        print("replacing image")
                        table = doc.tables[0]  # Assuming the first table
                        cell = table.cell(0, 3)  # Assuming the first cell in the third column

                        # Clear the content of the cell by removing its paragraphs
                        for paragraph in cell.paragraphs:
                            paragraph.clear()

                        # Add a new paragraph and insert the image
                        paragraph = cell.add_paragraph()
                        run = paragraph.add_run()
                        cadet_photo_url = data["cadetPhoto"]
                        cadet_photo_path = cadet_photo_url.replace(files_url,files_base_dir)
                        run.add_picture(cadet_photo_path, width=Inches(0.9))
                    except Exception as e:
                        print("Error : ",str(e))

                    doc.save(str(str(file_dir)+f"{sid}_entrance_card.docx"))

                    convert_to_pdf(str(str(file_dir)+f"{sid}_entrance_card.docx"), str(str(file_dir)+f"{sid}_entrance_card.pdf"))

                    ec = f"{files_base_url}{sid}_entrance_card.pdf"

                    entrance_card = {
                        "entrence_card" : ec,
                        "status":"Active"
                    }
                    final_status = "Active"
                    students_db.update_one({"sid": data['sid']}, {"$set": entrance_card})
                else:
                    if data['status']=="Active" or data['status']=="In Progress":
                        final_status = "In Progress"
                    else:
                        final_status = data['status']
                                
            except Exception as e:
                print("Error : ", str(e))
            medical_cert_url = f"{files_base_url}{sid}_MEDICAL_CER.pdf"
            visiting_card_url = f"{files_base_url}{sid}_visit_card.pdf"
            admission_form = f"{files_base_url}{sid}_admission_form.pdf"

        student = {
            "sid": sid,
            "first_name": data["first_name"].upper(),
            "middle_name": data.get("middle_name", "").upper(),
            "last_name": data["last_name"].upper(),
            "email": data["email"],
            "phn": str(data["phn"]),
            "dob": data["dob"],
            "age": str(age),
            "company":company,
            "address": data["address"].upper(),
            "fathers_occupation": data["fathers_occupation"].upper(),
            "mothers_occupation": data["mothers_occupation"].upper(),
            "how_you_got_to_know": data["how_you_got_to_know"].upper(),
            "employee_who_reached_out_to_you": data["employee_who_reached_out_to_you"].upper(),
            "district": data["district"].upper(),
            "state": data["state"].upper(),
            "pincode": str(data["pincode"]),
            "status": str(final_status),
            "camp_id": data.get("camp_id", ""),
            "camp_category": data.get("camp_category", "").upper(),
            "batch_id": data.get("batch_id", ""),
            "food_option": data.get("food_option", "").upper(),
            # "dress_code": data.get("dress_code", ""),
            "pick_up_city": data.get("pick_up_city", "").upper(),
            "pick_up_point": data.get("pick_up_point", "").upper(),
            "height": data.get("height", ""),
            "weight": data.get("weight", ""),
            "blood_group": data.get("blood_group", ""),
            "payment_option": data.get("payment_option", ""),
            "school_name": data.get("school_name", "").upper(),
            "gender": data.get("gender", "").upper(),
            "standard": data.get("standard", ""),
            "wp_no": data.get("wp_no", ""),
            "medication_physical":data.get("medication_physical"),
            "other_problem":data.get("other_problem"),
            "physical_problem":data.get("physical_problem",""),
            "medication_allergy":data.get("medication_allergy",""),
            "medication_other":data.get("medication_other",""),
            "allergy":data.get("allergy",""),
            "medicalCertificate":medical_cert_url,
            "cadetPhoto":data.get("cadetPhoto",""),
            "cadetSign":data.get("cadetSign",""),
            "parentGurdianPhoto":data.get("parentGurdianPhoto",""),
            "parentGurdianSign":data.get("parentGurdianSign",""),
            "payment_status": data.get("payment_status", "Pending"),
            "visiting_card":visiting_card_url,
            "admission_form":admission_form,
            'total_amount_payable':int(data.get("total_amount_payable", 0)),
            "total_amount_paid":int(data.get("total_amount_paid", 0)),
            "discount_code":data.get("discount_code", ""),
            "discount_amount":int(data['discount_amount']),
            "camp_year":str("20"+str(year))
        }

        # Store the student information in the MongoDB collection
        batches_db = db["batches_db"]
        batch = batches_db.find_one({"batch_id":data.get("batch_id")}, {"_id":0})
        if batch:
            if int(batch["students_registered"]) <= int(batch["batch_intake"]):
                students_db.update_one({'sid': original_sid},{"$set": student})
                if isCountInc:
                    batches_db.update_one({"batch_id": data.get("batch_id")}, {"$set": {"students_registered":int(int(batch["students_registered"])+1)}})
                else:
                    pass

                payment_db = db['all_payments']
                filter_criteria = {'sid': original_sid}
                update_operation = {'$set': {'sid': sid}}
                payment_db.update_many(filter_criteria, update_operation)
                return 0
            else:
                return 1

def sendSMS(msg,phn):
    if msg and phn:
        url = "http://msg.msgclub.net/rest/services/sendSMS/sendGroupSms"
        msg_text = msg
        phn_no = phn
        querystring = {"AUTH_KEY":"2b4186d8fc21f47949e7f5e92b56390","message":msg_text,"senderId":"MCFCMP","routeId":"1","mobileNos":phn_no,"smsContentType":"english"}
        headers = {'Cache-Control': "no-cache"}
        response = requests.request("GET", url, headers=headers, params=querystring)
        print(response.text)
        return 0
    else:
        return 1
    

def send_wp(sms_content, mobile_numbers, file_paths=[]):
    if len(file_paths)>1:
            file_paths.append("THINGS_TO_BRING.pdf")
    api_url = "http://msg.msgclub.net/rest/services/sendSMS/sendGroupSms"
    auth_key = "2b4186d8fc21f47949e7f5e92b56390"
    route_id = "21"
    sender_id = "9604992000"
    sms_content_type = "english"
    payload = {
        "smsContent": sms_content,
        "routeId": route_id,
        "mobileNumbers": mobile_numbers,
        "senderId": sender_id,
        "smsContentType": sms_content_type
    }
    headers = {
        "AUTH_KEY": auth_key,
        "Content-Type": "application/json"
    }

    payload2 = {
        "smsContent": "",
        "routeId": route_id,
        "mobileNumbers": mobile_numbers,
        "senderId": sender_id,
        "smsContentType": sms_content_type
    }

    # Add file data if file_path is provided
    if file_paths:
        if len(file_paths) == 1:
            filedata_encoded = encode_file_to_base64(file_paths[0])
            if filedata_encoded:
                payload["filename"] = file_paths[0].split('/')[-1]  # Extract filename from path
                payload["filedata"] = filedata_encoded
            else:
                print(f"Error: Unable to encode {file_path} to Base64")
                return 1
            response = requests.post(api_url, json=payload, headers=headers)
            if response.status_code == 200:
                response_json = response.json()
                if 'response' in response_json:
                    print("Send WP")
                    return 0
                else:
                    return 1
            else:
                return 1
        elif len(file_paths) > 1:
            for file_path in file_paths:
                filedata_encoded = encode_file_to_base64(file_path)
                if filedata_encoded:
                    payload2["filename"] = file_path.split('/')[-1]  # Extract filename from path
                    payload2["filedata"] = filedata_encoded
                else:
                    print(f"Error: Unable to encode {file_path} to Base64")
                response = requests.post(api_url, json=payload2, headers=headers)
                if response.status_code == 200:
                    response_json = response.json()
                    if 'response' in response_json:
                        print("Send WP")
                    else:
                        print("Error")
    response = requests.post(api_url, json=payload, headers=headers)
    if response.status_code == 200:
        response_json = response.json()
        if 'response' in response_json:
            print("Send WP")
            return 0
        else:
            return 1
    else:
        return 1

def encode_file_to_base64(file_path):
    try:
        with open(file_path, "rb") as file:
            filedata = file.read()
            filedata_encoded = base64.b64encode(filedata).decode('utf-8')
            return filedata_encoded
    except Exception as e:
        print(f"Error encoding file to Base64: {str(e)}")
        return None
    

def send_email(msg, sub, mailToSend):
    try:
        # Send the password reset link via email
        sender_email = "mcfcamp@gmail.com"
        smtp_server = smtplib.SMTP("smtp.gmail.com", 587)
        smtp_server.ehlo()
        smtp_server.starttls()
        smtp_server.login("mcfcamp@gmail.com", "meyv ghup onbl fqhu")

        message_text = msg
        message = MIMEText(message_text)
        message["Subject"] = sub
        message["From"] = sender_email
        message["To"] = mailToSend

        smtp_server.sendmail(sender_email, mailToSend, message.as_string())
        print(mailToSend)
        print("Send Mail")
        smtp_server.quit()
        return 0
    except Exception as e:
        print(str(e))
        return 1
    


import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

def send_email_attachments(msg, sub, mailToSend, files=[]):
    try:
        if len(files)>1:
            files.append("THINGS_TO_BRING.pdf")
        sender_email = "mcfcamp@gmail.com"
        smtp_server = smtplib.SMTP("smtp.gmail.com", 587)
        smtp_server.ehlo()
        smtp_server.starttls()
        smtp_server.login("mcfcamp@gmail.com", "meyv ghup onbl fqhu")

        # Create a multipart message
        message = MIMEMultipart()
        message["Subject"] = sub
        message["From"] = sender_email
        message["To"] = mailToSend

        # Attach message body
        message.attach(MIMEText(msg, "plain"))

        # Attach files
        for file_path in files:
            with open(file_path, "rb") as attachment:
                part = MIMEBase("application", "octet-stream")
                part.set_payload(attachment.read())

            # Encode file in ASCII characters to send by email
            encoders.encode_base64(part)

            # Add header as key/value pair to attachment part
            part.add_header(
                "Content-Disposition",
                f"attachment; filename= {file_path}",
            )

            # Attach the attachment to the message
            message.attach(part)

        smtp_server.sendmail(sender_email, mailToSend, message.as_string())
        print(mailToSend)
        print("Send Mail")
        smtp_server.quit()
        return 0
    except Exception as e:
        print(str(e))
        return 1

# ------------------------------------------------------------------------------------------------------------


@app.route("/getAllStudents", methods=["GET"])
def getAllStudents():
    users = db["students_db"]
    ans = []
    ans = list(users.find({},{'_id':0}))
    return jsonify({"students":ans[::-1]})

@app.route("/getInactiveStudents", methods=["GET"])
def getInactiveStudents():
    users = db["students_db"]
    ans = []
    ans = list(users.find({"status": {"$ne": "Active"}}, {"_id": 0}))
    return jsonify({"students":ans[::-1]})

def calculate_age(dob):
    try:
        birth_date = datetime.strptime(dob, "%d-%m-%Y")
        today = datetime.today()
        age = today.year - birth_date.year - ((today.month, today.day) < (birth_date.month, birth_date.day))
        return age
    except ValueError:
        raise ValueError("Invalid date of birth format. Please use 'dd-mm-yyyy'.")

def convert_to_pdf(docx_file, pdf_file):
    try:
        subprocess.run(['unoconv', '--output', pdf_file, '--format', 'pdf', docx_file], check=True)
        print(f"Conversion successful: {docx_file} -> {pdf_file}")
    except subprocess.CalledProcessError as e:
        print(f"Error during conversion: {e}")

def set_font(run, font_name, font_size):
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = True

def find_and_replace_paragraph_med(paragraph, field_values):
    for run in paragraph.runs:
        print(run.text)
        for field, value in field_values.items():
            run.text = run.text.replace(f"{field}", value)
            set_font(run, 'Montserrat', 14)

def replace_fields_in_document_med(doc_path, field_values):
    doc = Document(doc_path)
    for paragraph in doc.paragraphs:
        find_and_replace_paragraph_med(paragraph, field_values)
    doc.save(str(str(file_dir)+str(f"{field_values['sid']}_MEDICAL_CER.docx")))

    convert_to_pdf(str(str(file_dir)+str(f"{field_values['sid']}_MEDICAL_CER.docx")),str(str(file_dir)+str(f"{field_values['sid']}_MEDICAL_CER.pdf")))

def set_paragraph_font_entrace_card(paragraph, font_name, font_size, bold=False):
    for run in paragraph.runs:
        font = run.font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = bold

def find_and_replace_paragraphs_entrance_card(paragraphs, field, replacement):
    for paragraph in paragraphs:
        if field in paragraph.text:
            paragraph.text = paragraph.text.replace(field, replacement)
            set_paragraph_font_entrace_card(paragraph, 'Arial', 9, False)

def find_and_replace_tables_entrance_card(tables, field, replacement):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    find_and_replace_paragraphs_entrance_card([paragraph], field, replacement)

def set_paragraph_font_visiting_card(paragraph, font_name, font_size, bold=False):
    for run in paragraph.runs:
        font = run.font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = bold                    

def find_and_replace_paragraphs_visiting_card(paragraphs, field, replacement):
    for paragraph in paragraphs:
        if field in paragraph.text:
            paragraph.text = paragraph.text.replace(field, replacement)
            set_paragraph_font_visiting_card(paragraph, 'Times New Roman', 14, False)

def find_and_replace_tables_visiting_card(tables, field, replacement):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    find_and_replace_paragraphs_visiting_card([paragraph], field, replacement)

def replace_image_in_cell(doc, table_index, row_index, column_index, image_path):
    table = doc.tables[table_index]
    cell = table.cell(row_index, column_index)
    for paragraph in cell.paragraphs:
        paragraph.clear()
    paragraph = cell.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(1.4), height=Inches(1.8))
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER



def set_paragraph_font_fee_receipt(paragraph, font_name, font_size, bold=False):
    for run in paragraph.runs:
        font = run.font
        font.name = font_name
        font.size = Pt(font_size)                   

def find_and_replace_paragraphs_fee_receipt(paragraphs, field, replacement):
    for paragraph in paragraphs:
        if field in paragraph.text:
            paragraph.text = paragraph.text.replace(field, replacement)
            set_paragraph_font_fee_receipt(paragraph, 'Times New Roman', 11, False)

def find_and_replace_tables_fee_receipt(tables, field, replacement):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    find_and_replace_paragraphs_fee_receipt([paragraph], field, replacement)


def set_paragraph_font_admission_form(paragraph, font_name, font_size, bold=False):
    for run in paragraph.runs:
        font = run.font
        font.name = font_name
        font.size = Pt(font_size)                   

def find_and_replace_paragraphs_admission_form(paragraphs, field, replacement):
    for paragraph in paragraphs:
        if field in paragraph.text:
            paragraph.text = paragraph.text.replace(field, replacement)
            set_paragraph_font_admission_form(paragraph, 'Times New Roman', 11, False)

def find_and_replace_tables_admission_form(tables, field, replacement):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    find_and_replace_paragraphs_admission_form([paragraph], field, replacement)



def replace_image_in_cell_admission_form(doc, table_index, row_index, column_index, image_path,w,h):
    table = doc.tables[table_index]
    cell = table.cell(row_index, column_index)
    for paragraph in cell.paragraphs:
        paragraph.clear()
    paragraph = cell.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(image_path,width=Inches(w), height=Inches(h))
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER





def set_paragraph_font_cert(paragraph, font_name, font_size, bold):
    for run in paragraph.runs:
        font = run.font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = bold

def find_and_replace_paragraphs_cert(paragraphs, field, replacement, specific_font=None):
    for paragraph in paragraphs:
        if field in paragraph.text:
            paragraph.text = paragraph.text.replace(field, replacement)
            if specific_font is not None:
                set_paragraph_font_cert(paragraph, *specific_font)

def generate_certificate_cert(sid):
    students_db = db["students_db"]
    student_data = students_db.find_one({"sid":sid})
    camp_db = db["camps_db"]
    camp_data = camp_db.find_one({"camp_id":student_data['camp_id']})
    batches_db = db["batches_db"]
    batch_data = batches_db.find_one({"batch_id":student_data['batch_id']})

    if not student_data:
        return 1
    else:
        if 'BTC' in camp_data['camp_name'] or 'btc' in camp_data['camp_name']:
            doc = Document('templates_cert/CER_3D_BTC.docx')
        elif 'RTC' in camp_data['camp_name'] or 'rtc' in camp_data['camp_name']:
            doc = Document('templates_cert/CER_5D_RTC.docx')
        elif 'ATC' in camp_data['camp_name'] or 'atc' in camp_data['camp_name']:
            doc = Document('templates_cert/CER_7D_ATC.docx')
        elif 'CTC' in camp_data['camp_name'] or 'ctc' in camp_data['camp_name']:
            doc = Document('templates_cert/CER_15D_CTC.docx')
        elif 'PDC' in camp_data['camp_name'] or 'pdc' in camp_data['camp_name']:
            doc = Document('templates_cert/CER_15D_PDC.docx')
        elif 'SMTC' in camp_data['camp_name'] or 'smtc' in camp_data['camp_name']:
            doc = Document('templates_cert/CER_30D_SMTC.docx')
        else:
            print("invalid camp name")

        student_data1 = {
            'REG_NO': sid,
            'CADET_NAME': str(student_data['first_name']+" "+student_data['last_name']),
            'START_DATE': batch_data['start_date'],
            'END_DATE':batch_data['end_date'],
            'CQY': ''
        }

        for key, value in student_data1.items():
            if key == 'CADET_NAME':
                find_and_replace_paragraphs_cert(doc.paragraphs, f'{{MERGEFIELD {key}}}', str(value), specific_font=('Times New Roman', 18, True))
            else:
                find_and_replace_paragraphs_cert(doc.paragraphs, f'{{MERGEFIELD {key}}}', str(value), specific_font=('Times New Roman', 14, True))
                
        # doc.save(f"{student_data['REG_NO']}_CER_{student_data['CAMP_NAME']}.docx")
        doc.save(str(str(file_dir)+f"{sid}_Certificate.docx"))
        convert_to_pdf(str(str(file_dir)+f"{sid}_Certificate.docx"), str(str(file_dir)+f"{sid}_Certificate.pdf"))
        cert_url = f"{files_base_url}{sid}_Certificate.pdf"

        students_db.update_one({"sid": sid}, {"$set": {"completion_cert":cert_url}})
        return 0
    


def set_paragraph_font_fdb(paragraph, font_name, font_size, bold=False):
    for run in paragraph.runs:
        font = run.font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = bold                    

def find_and_replace_paragraphs_fdb(paragraphs, field, replacement):
    for paragraph in paragraphs:
        if field in paragraph.text:
            paragraph.text = paragraph.text.replace(field, replacement)
            set_paragraph_font_fdb(paragraph, 'Times New Roman', 11, False)

def find_and_replace_tables_fdb(tables, field, replacement):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    find_and_replace_paragraphs_fdb([paragraph], field, replacement)

def replace_image_in_cell_fdb(doc, table_index, row_index, column_index, image_path,w):
    table = doc.tables[table_index]
    cell = table.cell(row_index, column_index)
    for paragraph in cell.paragraphs:
        paragraph.clear()
    paragraph = cell.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(w))
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER




def number_to_words(num):
    # Define lists of words for numbers
    units = ['Zero', 'One', 'Two', 'Three', 'Four', 'Five', 'Six', 'Seven', 'Eight', 'Nine']
    teens = ['Ten', 'Eleven', 'Twelve', 'Thirteen', 'Fourteen', 'Fifteen', 'Sixteen', 'Seventeen', 'Eighteen', 'Nineteen']
    tens = ['Twenty', 'Thirty', 'Forty', 'Fifty', 'Sixty', 'Seventy', 'Eighty', 'Ninety']
    thousands = ['', 'Thousand', 'Million', 'Billion', 'Trillion']

    # Function to convert numbers less than 1000 to words
    def helper(num):
        if num == 0:
            return ''
        elif num < 10:
            return units[num]
        elif num < 20:
            return teens[num - 10]
        elif num < 100:
            return tens[num // 10 - 2] + ' ' + helper(num % 10)
        else:
            return units[num // 100] + ' Hundred ' + helper(num % 100)

    if num == 0:
        return 'Zero'

    words = ''
    for i in range(len(thousands)):
        if num % 1000 != 0:
            words = helper(num % 1000) + thousands[i] + ' ' + words
        num //= 1000

    return words.strip()

        


@app.route('/registerStudent', methods=['POST'])
def register_student():
    try:
        data = request.form

        student_raw_data = {
            "first_name": data["first_name"],
            "middle_name": data.get("middle_name", ""),
            "last_name": data["last_name"],
            "email": data["email"],
            "phn": str(data["phn"]),
            "dob": data["dob"],
            "address": data["address"],
            "fathers_occupation": data["fathers_occupation"],
            "mothers_occupation": data["mothers_occupation"],
            "how_you_got_to_know": data["how_you_got_to_know"],
            "employee_who_reached_out_to_you": data["employee_who_reached_out_to_you"],
            "district": data["district"],
            "state": data["state"],
            "pincode": str(data["pincode"]),
            "camp_id": data.get("camp_id", ""),
            "camp_category": data.get("camp_category", ""),
            "batch_id": data.get("batch_id", ""),
            "food_option": data.get("food_option", ""),
            "pick_up_city": data.get("pick_up_city", ""),
            "pick_up_point": data.get("pick_up_point", ""),
            "height": data.get("height", ""),
            "weight": data.get("weight", ""),
            "blood_group": data.get("blood_group", ""),
            "school_name": data.get("school_name", ""),
            "gender": data.get("gender", ""),
            "standard": data.get("standard", ""),
            "wp_no": data.get("wp_no", ""),
            "medication_physical":data.get("medication_physical"),
            "other_problem":data.get("other_problem"),
            "physical_problem":data.get("physical_problem",""),
            "medication_allergy":data.get("medication_allergy",""),
            "medication_other":data.get("medication_other",""),
            "allergy":data.get("allergy",""),
            "cadetPhoto":data.get("cadetPhoto",""),
            "cadetSign":data.get("cadetSign",""),
            "parentGurdianPhoto":data.get("parentGurdianPhoto",""),
            "parentGurdianSign":data.get("parentGurdianSign",""),
            "payment_status": data.get("payment_status", "Pending"),
            'total_amount_payable':int(data.get("total_amount_payable", 0)),
        }

        all_keys = student_raw_data.keys()
        required_keys = list(all_keys)
        for key in required_keys:
            if not student_raw_data.get(key) or student_raw_data[key] == "":
                return jsonify({"error": f"Please fill in the '{key}' field."}), 400

        # Check if email and phn are not repeating
        students_db = db["students_db"]
        existing_student_email = students_db.find_one({"email": data["email"]})
        existing_student_phn = students_db.find_one({"phn": data["phn"]})

        if existing_student_email:
            raise ValueError(f"Email '{data['email']}' is already registered.")

        if existing_student_phn:
            raise ValueError(f"Phone number '{data['phn']}' is already registered.")

        # Generate a unique ID for the student using UUID
        sid = str(uuid.uuid4().hex)

        company = ""
        # Calculate age based on the provided date of birth
        age = calculate_age(data["dob"])
        if age>=7 and age<=11 and data.get("gender") == "male":
            company = "ALPHA"
        elif age>=12 and age<=16 and data.get("gender") == "male":
            company = "BRAVO"
        elif age>=17 and age<=21 and data.get("gender") == "male":
            company = "DELTA"
        elif age>=7 and age<=11 and data.get("gender") == "female":
            company = "CHARLEY"
        elif age>=12 and age<=16 and data.get("gender") == "female":
            company = "ECO"
        elif age>=17 and age<=21 and data.get("gender") == "female":
            company = "FOXFORD"

        batches_db = db["batches_db"]
        batch = batches_db.find_one({"batch_id":data.get("batch_id")}, {"_id":0})
        camps_db = db["camps_db"]
        camp = camps_db.find_one({"camp_id":data.get("camp_id")}, {"_id":0})
        camp_name = camp["camp_name"]
        camp_short = camp_name.split(" ")[-1].replace("(", "").replace(")", "")
        sid=""
        medical_cert_url = ""
        if batch:
            if int(batch["students_registered"]) <= int(batch["batch_intake"]):
                sr_no = int(int(batch["students_registered"])+1)
                start_date = batch["start_date"]
                year = start_date[-2:]
                day = start_date[0:2]
                batch_name = batch["batch_name"].replace(" ", "")

                company_sf = str(company[0])+"C"
                days = str(batch['duration'])+"D"

                sid = str(camp_short)+str(year)+str(days)+str(batch_name)+str(company_sf)+str(sr_no)

                document_med_path = 'medical_certificate.docx'

                field_values = {
                    'CADET_NAME': str(data["first_name"].upper()+" "+data["last_name"].upper()),
                    'LOC':  str(data["district"]+", "+data["state"]),
                    'DOB':  str(data["dob"]),
                    '121212':  str("__________"),
                    'C_NAME':  str(camp_name),
                    'DATE':  str(start_date),
                    'BATCH':  str(batch_name),
                    'sid': sid
                }
                replace_fields_in_document_med(document_med_path, field_values)

                # Load the document template
                doc1 = Document('visit_card.docx')

                # Sample student_data
                student_data1 = {
                    'CADET_NAME': str(data["first_name"].upper()+" "+data["last_name"].upper()),
                    'CAMP_NAME': str(camp_name),
                    'BATCH_NO': str(batch_name),
                    'ADDRESS': data["address"],
                    'CONTACT': data["phn"],
                    'WHATSAPP_NO': data["wp_no"],
                    'CAMP_DATE':start_date,
                    'REG_NO':sid,
                    'PICKUP_POINT':data['pick_up_point'],
                }

                for key, value in student_data1.items():
                    find_and_replace_paragraphs_visiting_card(doc1.paragraphs, f'{{MERGEFIELD {key}}}', str(value))

                try:
                    cadet_photo_url = data["cadetPhoto"]
                    cadet_photo_path = cadet_photo_url.replace(files_url,files_base_dir)

                    image_url_guardian = data["parentGurdianPhoto"]
                    image_path_guardian = image_url_guardian.replace(files_url,files_base_dir)
                    
                    replace_image_in_cell(doc1, table_index=0, row_index=0, column_index=3, image_path=cadet_photo_path)
                    replace_image_in_cell(doc1, table_index=0, row_index=0, column_index=4, image_path=image_path_guardian)

                    doc1.save(str(str(file_dir)+f"{sid}_visit_card.docx"))

                    convert_to_pdf(str(str(file_dir)+f"{sid}_visit_card.docx"), str(str(file_dir)+f"{sid}_visit_card.pdf"))



                    doc = Document('admission_form_new_format.docx')
                    student_data1 = {
                    "REG_NO": sid,
                    "FIRST_NAME": str(data["first_name"].upper()),
                    "MIDDLE_NAME": str(data["middle_name"].upper()),
                    "LAST_NAME": str(data["last_name"].upper()),
                    "EMAIL_ID": str(data["email"]),
                    "CONTACT_NO": str(data["phn"]),
                    "DATE_OF_BIRTH": str(data["dob"]),
                    "ADDRESS": str(data["address"]),
                    "HOW_YOU_GOT_TO_KNOW": str(data["how_you_got_to_know"]),
                    "EMPLOYEE_WHO_REACHED_OUT": str(data["employee_who_reached_out_to_you"]),
                    "DISTRICT": str(data["district"]),
                    "STATE": data["state"],
                    "PINCODE": data["pincode"],
                    "PICKUP_POINT": data["pick_up_point"],
                    "BLOOD_GROUP": data["blood_group"],
                    "SCHOOL_NAME": data["school_name"],
                    "GENDER": data["gender"],
                    "STANDARD": data["standard"],
                    "WHATSAPP_NO": data["wp_no"],
                    "PARENT_NAME":data["middle_name"],
                    "CAMP_NAME":camp_name,
                    "CAMP_DATE":batch["start_date"],
                    "CAMP_DAYS":batch["duration"]
                    }
                    for key, value in student_data1.items():
                            find_and_replace_tables_admission_form(doc.tables, f'{{MERGEFIELD {key}}}', str(value))

                    image_path_sign_url = data["cadetSign"]
                    image_path_sign = image_path_sign_url.replace(files_url,files_base_dir)

                    image_url_guardian_sign = data["parentGurdianSign"]
                    image_path_guardian_sign = image_url_guardian_sign.replace(files_url,files_base_dir)

                    replace_image_in_cell_admission_form(doc, table_index=0, row_index=25, column_index=1, image_path=cadet_photo_path,w=1.4,h=1.6)
                    replace_image_in_cell_admission_form(doc, table_index=0, row_index=25, column_index=12, image_path=image_path_sign,w=1.8,h=1.0)
                    replace_image_in_cell_admission_form(doc, table_index=1, row_index=16, column_index=1, image_path=image_path_guardian,w=1.4,h=1.6)
                    replace_image_in_cell_admission_form(doc, table_index=1, row_index=17, column_index=6, image_path=image_path_guardian_sign,w=1.8,h=1.0)

                    doc.save(str(str(file_dir)+f"{sid}_admission_form.docx"))

                    convert_to_pdf(str(str(file_dir)+f"{sid}_admission_form.docx"), str(str(file_dir)+f"{sid}_admission_form.pdf"))
                                   
                except Exception as e:
                    print("Error : ", str(e))
                medical_cert_url = f"{files_base_url}{sid}_MEDICAL_CER.pdf"
                visiting_card_url = f"{files_base_url}{sid}_visit_card.pdf"
                admission_form = f"{files_base_url}{sid}_admission_form.pdf"

            else:
                return jsonify({"message": "Batch is Already Full !"})
        

        student = {
            "sid": sid,
            "first_name": data["first_name"].upper(),
            "middle_name": data.get("middle_name", "").upper(),
            "last_name": data["last_name"].upper(),
            "email": data["email"],
            "phn": str(data["phn"]),
            "dob": data["dob"],
            "age": str(age),
            "company":company,
            "address": data["address"].upper(),
            "fathers_occupation": data["fathers_occupation"].upper(),
            "mothers_occupation": data["mothers_occupation"].upper(),
            "how_you_got_to_know": data["how_you_got_to_know"].upper(),
            "employee_who_reached_out_to_you": data["employee_who_reached_out_to_you"].upper(),
            "district": data["district"].upper(),
            "state": data["state"].upper(),
            "pincode": str(data["pincode"]),
            "status": "In Progress",
            "camp_id": data.get("camp_id", ""),
            "camp_category": data.get("camp_category", "").upper(),
            "batch_id": data.get("batch_id", ""),
            "food_option": data.get("food_option", "").upper(),
            # "dress_code": data.get("dress_code", ""),
            "pick_up_city": data.get("pick_up_city", "").upper(),
            "pick_up_point": data.get("pick_up_point", "").upper(),
            "height": data.get("height", ""),
            "weight": data.get("weight", ""),
            "blood_group": data.get("blood_group", ""),
            "payment_option": data.get("payment_option", ""),
            "school_name": data.get("school_name", "").upper(),
            "gender": data.get("gender", "").upper(),
            "standard": data.get("standard", ""),
            "wp_no": data.get("wp_no", ""),
            "medication_physical":data.get("medication_physical"),
            "other_problem":data.get("other_problem"),
            "physical_problem":data.get("physical_problem",""),
            "medication_allergy":data.get("medication_allergy",""),
            "medication_other":data.get("medication_other",""),
            "allergy":data.get("allergy",""),
            "medicalCertificate":medical_cert_url,
            "cadetPhoto":data.get("cadetPhoto",""),
            "cadetSign":data.get("cadetSign",""),
            "parentGurdianPhoto":data.get("parentGurdianPhoto",""),
            "parentGurdianSign":data.get("parentGurdianSign",""),
            "payment_status": data.get("payment_status", "Pending"),
            "visiting_card":visiting_card_url,
            "admission_form":admission_form,
            'total_amount_payable':int(data.get("total_amount_payable", 0)),
            "total_amount_paid":0,
            "discount_code":data.get("discount_code", ""),
            "discount_amount":int(0),
            "camp_year":str("20"+str(year))
        }

        # Store the student information in the MongoDB collection
        batches_db = db["batches_db"]
        batch = batches_db.find_one({"batch_id":data.get("batch_id")}, {"_id":0})
        if batch:
            if int(batch["students_registered"]) <= int(batch["batch_intake"]):
                students_db.insert_one(student)
                batches_db.update_one({"batch_id": data.get("batch_id")}, {"$set": {"students_registered":int(int(batch["students_registered"])+1)}})
                msg = "Dear Parent, Thank you for registering with MCF Camp, for any registration and payment-related query please visit us at www.mcfcamp.in. Or contact us at 9604087000/9604082000, or email us at mcfcamp@gmail.com MCF Summer Camp"
                sub = "Registration Successful !"
                mailToSend = data['email']
                sendSMS(msg,data["phn"])
                send_wp(msg,data["wp_no"])
                send_email(msg, sub, mailToSend)

                msg2 = f"New Student Registered \n\n Name - {data['first_name']} {data['last_name']} \n Camp Name - {camp_name} \n Batch Name - {batch_name}"
                send_wp(msg2,"9604084000")
                send_email(msg2, sub, "infomcfcamp@gmail.com")
                return jsonify({"message": "Student registered successfully", "sid": sid})
            else:
                return jsonify({"message": "Batch is Already Full !"})

        return jsonify({"message": "Student registered successfully", "sid": sid})

    except ValueError as ve:
        return jsonify({"error": str(ve)}), 400  # Bad Request

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error


@app.route('/updateStudent', methods=['PUT'])
def update_student():
    try:
        data = request.form

        # Check if sid is provided
        if 'sid' not in data:
            raise ValueError("Missing 'sid' in the request.")

        # Find the student based on sid
        students_db = db["students_db"]
        student = students_db.find_one({"sid": data['sid']})

        if not student:
            return jsonify({"error": f"No student found with sid: {data['sid']}"}), 404  # Not Found

        # Update the student information with the received data
        for key, value in data.items():
            if key != 'sid':
                student[key] = value

        # Update the student in the database
        batches_db = db["batches_db"]
        batch = batches_db.find_one({"batch_id":data.get("batch_id")}, {"_id":0})
        if batch:
            if int(batch["students_registered"]) <= int(batch["batch_intake"]):
                students_db.update_one({"sid": data['sid']}, {"$set": student})
                result = sync_data(data['sid'])
                return jsonify({"message": f"Student with sid {data['sid']} updated successfully"})
            else:
                return jsonify({"message": "Batch is Already Full !"}),400
        else:
            return jsonify({"message": "Batch not Found !"}),400
            
        # students_db.update_one({"sid": data['sid']}, {"$set": student})
        # return jsonify({"message": f"Student with sid {data['sid']} updated successfully"})

    except ValueError as ve:
        return jsonify({"error": str(ve)}), 400  # Bad Request

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/deleteStudent', methods=['DELETE'])
def delete_student():
    try:
        students_db = db["students_db"]
        # Get the sid from request parameters
        sid = request.args.get('sid')

        if not sid:
            return jsonify({"error": "Missing 'sid' parameter in the request."}), 400  # Bad Request

        # Find the student based on sid
        student = students_db.find_one({"sid": sid})

        if not student:
            return jsonify({"error": f"No student found with sid: {sid}"}), 404  # Not Found

        # Delete the student from the database
        students_db.delete_one({"sid": sid})

        return jsonify({"message": f"Student with sid {sid} deleted successfully"})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/getStudent', methods=['GET'])
def get_student():
    try:
        # Get the sid from request parameters
        sid = request.args.get('sid')

        if not sid:
            return jsonify({"error": "Missing 'sid' parameter in the request."}), 400  # Bad Request

        # Find the student based on sid
        students_db = db["students_db"]
        student = students_db.find_one({"sid": sid}, {"_id": 0})  # Exclude the _id field from the response

        if not student:
            return jsonify({"error": f"No student found with sid: {sid}"}), 404  # Not Found

        # If camp_id and batch_id are available in the student data, fetch additional information
        camp_id = student.get("camp_id", "")
        batch_id = student.get("batch_id", "")

        # Fetch camp details if camp_id is present
        camp_details = {}
        if camp_id:
            camps_db = db["camps_db"]
            camp_details = camps_db.find_one({"camp_id": camp_id}, {"_id": 0})

        # Fetch batch details if batch_id is present
        batch_details = {}
        if batch_id:
            batches_db = db["batches_db"]
            batch_details = batches_db.find_one({"batch_id": batch_id}, {"_id": 0})

        # Combine student, camp, and batch details
        response_data = {
            "student": student,
            "camp_details": camp_details,
            "batch_details": batch_details
        }

        return jsonify(response_data)

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    

def build_filter_query(params):
    filter_query = {}

    for key, value in params.items():
        if value:
            # For 'age-group', parse l_age and g_age and add to the filter query
            if key == 'age-group':
                l_age, g_age = value.split(',')
                filter_query['age'] = {"$gte": int(l_age), "$lte": int(g_age)}

            # For other parameters, use regex for partial matching
            else:
                filter_query[key] = re.compile(f".*{re.escape(value)}.*", re.IGNORECASE)

    return filter_query

@app.route('/filterStudents', methods=['POST'])
def filter_students():
    try:
        # Get filter parameters from request parameters
        filter_params = request.json

        # Build the filter query
        filter_query = build_filter_query(filter_params)

        # Find students based on the filter query
        students_db = db["students_db"]
        students = students_db.find(filter_query, {"_id": 0})  # Exclude the _id field from the response

        # Convert the cursor to a list of dictionaries for easier serialization
        student_list = list(students)

        return jsonify({"students": student_list})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/addCamp', methods=['POST'])
def add_camp():
    try:
        data = request.form
        print("Data Recieved : ",data)

        # Validate required fields
        # required_fields = ["camp_name", "chess_prefix", "camp_place", "camp_fee", "camp_description", "fee_discount", "discount_date", "final_fee"]
        # for field in required_fields:
        #     if field not in data or not data[field]:
        #         raise ValueError(f"Missing or empty value for the required field: {field}")

        # Generate a unique ID for the camp using UUID
        camp_id = str(uuid.uuid4().hex)

        camp = {
            "camp_id": camp_id,
            "camp_name": data["camp_name"].strip(),
            "camp_place": data["camp_place"],
            "camp_fee": float(data["camp_fee"]),  # assuming camp_fee is a float
            "camp_description": data["camp_description"],
            "fee_discount": float(data["fee_discount"]),  # assuming fee_discount is a float
            # "discount_date": data["discount_date"],
            "camp_status" : data["camp_status"],
            "final_fee": float(data["final_fee"])  # assuming final_fee is a float
        }

        # Store the camp information in the MongoDB collection
        camps_db = db["camps_db"]
        camps_db.insert_one(camp)

        return jsonify({"message": "Camp added successfully", "camp_id": camp_id})

    except ValueError as ve:
        return jsonify({"error": str(ve)}), 400  # Bad Request

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/updateCamp', methods=['PUT'])
def update_camp():
    try:
        data = request.form

        # Check if camp_id is provided
        if 'camp_id' not in data:
            raise ValueError("Missing 'camp_id' in the request.")

        # Find the camp based on camp_id
        camps_db = db["camps_db"]
        camp = camps_db.find_one({"camp_id": data['camp_id']})

        if not camp:
            return jsonify({"error": f"No camp found with camp_id: {data['camp_id']}"}), 404  # Not Found

        # Update the camp information with the received data
        for key, value in data.items():
            if key != 'camp_id':
                # If the value is provided, update the field; otherwise, keep the existing value
                if value:
                    camp[key] = float(value) if key.endswith('_fee') else value
                    if camp['camp_status'] == "on":
                        camp["camp_status"] = "Active"
                    else:
                        camp['camp_status'] = "Inactive"

        # Update the camp in the database
        camps_db.update_one({"camp_id": data['camp_id']}, {"$set": camp})

        return jsonify({"message": f"Camp with camp_id {data['camp_id']} updated successfully"})

    except ValueError as ve:
        return jsonify({"error": str(ve)}), 400  # Bad Request

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/getAllCamps', methods=['GET'])
def get_all_camps():
    try:
        camps_db = db["camps_db"]
        camps = camps_db.find({}, {"_id": 0})  # Exclude the _id field from the response

        # Convert the cursor to a list of dictionaries for easier serialization
        camp_list = list(camps)

        return jsonify({"camps": camp_list})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/getAllCampsActive', methods=['GET'])
def get_all_camps_active():
    try:
        camps_db = db["camps_db"]
        camps = camps_db.find({"camp_status":"Active"}, {"_id": 0})  # Exclude the _id field from the response

        # Convert the cursor to a list of dictionaries for easier serialization
        camp_list = list(camps)

        return jsonify({"camps": camp_list})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/getAllBatches', methods=['GET'])
def get_all_batches():
    try:
        batches_db = db["batches_db"]
        batches = batches_db.find({}, {"_id": 0})  # Exclude the _id field from the response

        # Convert the cursor to a list of dictionaries for easier serialization
        batches_list = list(batches)

        return jsonify({"camps": batches_list})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/getCamp', methods=['GET'])
def get_camp():
    try:
        # Get the camp_id from request parameters
        camp_id = request.args.get('camp_id')

        if not camp_id:
            return jsonify({"error": "Missing 'camp_id' parameter in the request."}), 400  # Bad Request

        # Find the camp based on camp_id
        camps_db = db["camps_db"]
        camp = camps_db.find_one({"camp_id": camp_id}, {"_id": 0})  # Exclude the _id field from the response

        if not camp:
            return jsonify({"error": f"No camp found with camp_id: {camp_id}"}), 404  # Not Found

        return jsonify({"camp": camp})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/deleteCamp', methods=['DELETE'])
def delete_camp():
    try:
        # Get the camp_id from request parameters
        camp_id = request.args.get('camp_id')

        if not camp_id:
            return jsonify({"error": "Missing 'camp_id' parameter in the request."}), 400  # Bad Request

        # Find the camp based on camp_id
        camps_db = db["camps_db"]
        camp = camps_db.find_one({"camp_id": camp_id})

        if not camp:
            return jsonify({"error": f"No camp found with camp_id: {camp_id}"}), 404  # Not Found

        # Delete the camp from the database
        camps_db.delete_one({"camp_id": camp_id})

        return jsonify({"message": f"Camp with camp_id {camp_id} deleted successfully"})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/addBatch', methods=['POST'])
def add_batch():
    try:
        data = request.form

        # Validate required fields
        required_fields = ["batch_name", "start_date", "end_date", "duration", "batch_intake", "camp_id"]
        for field in required_fields:
            if field not in data or not data[field]:
                raise ValueError(f"Missing or empty value for the required field: {field}")

        # Parse start_date and end_date to datetime objects
        start_date = data["start_date"]
        end_date = data["end_date"]

        # Generate a unique ID for the batch using UUID
        batch_id = str(uuid.uuid4().hex)

        batch = {
            "batch_id": batch_id,
            "batch_name": data["batch_name"],
            "start_date": start_date,
            "end_date": end_date,
            "company": data["company"],
            "duration": data["duration"],
            "batch_intake": int(data["batch_intake"]),
            "students_registered": 0,
            "camp_id": data["camp_id"]
        }

        # Store the batch information in the MongoDB collection
        batches_db = db["batches_db"]
        batches_db.insert_one(batch)

        return jsonify({"message": "Batch added successfully", "batch_id": batch_id})

    except ValueError as ve:
        return jsonify({"error": str(ve)}), 400  # Bad Request

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error

@app.route('/updateBatch', methods=['POST'])
def update_batch():
    try:
        data = request.form

        # Check if batch_id is provided
        if 'batch_id' not in data:
            raise ValueError("Missing 'batch_id' in the request.")

        # Find the batch based on batch_id
        batches_db = db["batches_db"]
        batch = batches_db.find_one({"batch_id": data['batch_id']})

        if not batch:
            return jsonify({"error": f"No batch found with batch_id: {data['batch_id']}"}), 404  # Not Found

        # Update the batch information with the received data
        for key, value in data.items():
            if key != 'batch_id':
                # If the value is provided, update the field; otherwise, keep the existing value
                if value:
                    batch[key] = int(value) if key == 'batch_intake' else value

        # Update the batch in the database
        batches_db.update_one({"batch_id": data['batch_id']}, {"$set": batch})

        return jsonify({"message": f"Batch with batch_id {data['batch_id']} updated successfully", "camp_id":batch['camp_id']})

    except ValueError as ve:
        return jsonify({"error": str(ve)}), 400  # Bad Request

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/getBatches', methods=['GET'])
def get_batches():
    try:
        # Get the camp_id from request parameters
        camp_id = request.args.get('camp_id')

        if not camp_id:
            return jsonify({"error": "Missing 'camp_id' parameter in the request."}), 400  # Bad Request

        # Find batches based on camp_id
        batches_db = db["batches_db"]
        batches = batches_db.find({"camp_id": camp_id}, {"_id": 0})  # Exclude the _id field from the response

        # Convert the cursor to a list of dictionaries for easier serialization
        batch_list = list(batches)

        return jsonify({"batches": batch_list})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/getBatch', methods=['GET'])
def get_batch():
    try:
        # Get the batch_id from request parameters
        batch_id = request.args.get('batch_id')

        if not batch_id:
            return jsonify({"error": "Missing 'batch_id' parameter in the request."}), 400  # Bad Request

        # Find the batch based on batch_id
        batches_db = db["batches_db"]
        batch = batches_db.find_one({"batch_id": batch_id}, {"_id": 0})  # Exclude the _id field from the response

        if not batch:
            return jsonify({"error": f"No batch found with batch_id: {batch_id}"}), 404  # Not Found

        return jsonify({"batch": batch})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/deleteBatch', methods=['DELETE'])
def delete_batch():
    try:
        # Get the batch_id from request parameters
        batch_id = request.args.get('batch_id')

        if not batch_id:
            return jsonify({"error": "Missing 'batch_id' parameter in the request."}), 400  # Bad Request

        # Find the batch based on batch_id
        batches_db = db["batches_db"]
        batch = batches_db.find_one({"batch_id": batch_id})

        if not batch:
            return jsonify({"error": f"No batch found with batch_id: {batch_id}"}), 404  # Not Found

        # Delete the batch from the database
        batches_db.delete_one({"batch_id": batch_id})

        return jsonify({"message": f"Batch with batch_id {batch_id} deleted successfully"})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/activateStudent', methods=['GET'])
def activate_student():
    try:
        # Get the sid from request parameters
        sid = request.args.get('sid')

        if not sid:
            return jsonify({"error": "Missing 'sid' parameter in the request."}), 400  # Bad Request

        # Find the student based on sid
        students_db = db["students_db"]
        student = students_db.find_one({"sid": sid}, {"_id": 0})

        amt_paid = student['total_amount_paid']
        amt_payable = student['total_amount_payable']

        if not student:
            return jsonify({"error": f"No student found with sid: {sid}"}), 404  # Not Found

        # Update the status to "Active"
        if amt_paid >= amt_payable:
            students_db.update_one({"sid": sid}, {"$set": {"status": "Active"}})
        else:
            students_db.update_one({"sid": sid}, {"$set": {"status": "In Progress"}})

        return jsonify({"message": f"Student with sid {sid} is now Active."})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/generateReport', methods=['POST'])
def generate_report():
    try:
        data = request.get_json()

        # Get parameters from the JSON data
        sid = data.get('sid')
        rank = data.get('rank')
        report_date = data.get('report_date')
        report_camp_name = data.get('report_camp_name')
        in_charge = data.get('in_charge')
        cqy = data.get('cqy')
        discipline = data.get('discipline')
        physical_fitness = data.get('physical_fitness')
        courage = data.get('courage')
        leadership = data.get('leadership')
        initiative = data.get('initiative')
        interpersonal_relations = data.get('interpersonal_relations')
        team_building = data.get('team_building')
        training = data.get('training')
        remark = data.get('remark')

        # Check if sid is provided
        if not sid:
            return jsonify({"error": "Missing 'sid' parameter in the request."}), 400  # Bad Request

        # Find the student based on sid
        students_db = db["students_db"]
        student = students_db.find_one({"sid": sid}, {"_id": 0})

        if not student:
            return jsonify({"error": f"No student found with sid: {sid}"}), 404  # Not Found

        # Update the student record with the report details
        students_db.update_one(
            {"sid": sid},
            {"$set": {
                "rank": rank,
                "report_date": report_date,
                "report_camp_name": report_camp_name,
                "in_charge": in_charge,
                "cqy": cqy,
                "discipline": discipline,
                "physical_fitness": physical_fitness,
                "courage": courage,
                "leadership": leadership,
                "initiative": initiative,
                "interpersonal_relations": interpersonal_relations,
                "team_building": team_building,
                "training": training,
                "remark": remark
            }}
        )

        return jsonify({"message": f"Report generated for student with sid {sid}."})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error

@app.route('/addAdmin', methods=['POST'])
def add_admin():
    try:
        data = request.get_json()

        # Get parameters from the JSON data
        admins_db = db['admins_db']
        username = data.get('username')
        password = data.get('password')
        email = data.get('email')
        admin_id = str(uuid.uuid4().hex)
        admin = admins_db.find_one({"username": username}, {"_id": 0})
        if admin:
            return jsonify({"success":False,"error":"Username Already Exist"})

        # Check if username and password are provided
        if not username or not password:
            return jsonify({"error": "Username and password are required."}), 400  # Bad Request

        # Hash the password before storing it
        hashed_password = generate_password_hash(password, method='pbkdf2:sha256')

        # Store admin information in the MongoDB collection
        admins_db = db['admins_db']
        admins_db.insert_one({"username": username, "password": hashed_password, "email": email, "admin_id":admin_id})

        return jsonify({"message": "Admin added successfully.","success":True, "admin_id":admin_id})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error

def create_jwt_token(admin_id):
    import datetime
    payload = {
        'admin_id': admin_id,
        'exp': datetime.datetime.utcnow() + datetime.timedelta(days=1)  # Token expiration time
    }
    token = jwt.encode(payload, app.config['SECRET_KEY'], algorithm='HS256')
    return token

@app.route('/loginAdmin', methods=['POST'])
def login_admin():
    try:
        data = request.get_json()

        # Get parameters from the JSON data
        username = data.get('username')
        password = data.get('password')

        # Check if username and password are provided
        if not username or not password:
            return jsonify({"error": "Username and password are required.", "success": False}), 400  # Bad Request

        # Find the admin based on username
        admins_db = db["admins_db"]
        admin = admins_db.find_one({"username": username}, {"_id": 0})

        if not admin or not check_password_hash(admin.get("password", ""), password):
            return jsonify({"error": "Invalid username or password.", "success": False}), 401  # Unauthorized

        # Generate JWT token
        token = create_jwt_token(admin['admin_id'])

        return jsonify({"message": "Login successful.", "success": True, "admin_id": admin['admin_id'], "token": token})

    except Exception as e:
        return jsonify({"error": str(e), "success": False}), 500  # Internal Server Error

    
@app.route('/getAllAdmin', methods=['GET'])
def get_all_admin():
    try:
        # Retrieve all admin records from the MongoDB collection
        admins_db = db["admins_db"]
        admins = list(admins_db.find({}, {"_id": 0}))

        return jsonify({"admins": admins})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/getStudentCounts', methods=['GET'])
def get_student_counts():
    try:
        # Retrieve count of canceled students
        total_students_count = db["students_db"].count_documents({})

        active_students_count = db["students_db"].count_documents({"status": "Active"})

        canceled_students_count = db["students_db"].count_documents({"status": "Cancel"})

        # Retrieve count of refunded students
        refunded_students_count = db["students_db"].count_documents({"status": "Refund"})

        # Retrieve count of extended students
        extended_students_count = db["students_db"].count_documents({"status": "Extend"})

        return jsonify({
            "active_students_count": active_students_count,
            "canceled_students_count": canceled_students_count,
            "refunded_students_count": refunded_students_count,
            "extended_students_count": extended_students_count,
            "total_students_count": total_students_count
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/getCanceledStudents', methods=['GET'])
def get_canceled_students():
    try:
        canceled_students_count = db["students_db"].find({"status": "Cancel"},{"_id":0})

        return jsonify({
            "canceled_students_count": canceled_students_count
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/getRefundStudents', methods=['GET'])
def get_refund_students():
    try:
        # Retrieve count of refunded students
        refunded_students_count = db["students_db"].find({"status": "Refund"},{"_id":0})
        return jsonify({
            "refunded_students_count": refunded_students_count
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/getExtendedStudents', methods=['GET'])
def get_extended_students():
    try:

        # Retrieve count of extended students
        extended_students_count = db["students_db"].find({"status": "Extend"},{"_id":0})

        return jsonify({
            "extended_students_count": extended_students_count
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/changeStatus', methods=['PUT'])
def change_student_status():
    try:
        data = request.get_json()

        # Check if sid and new_status are provided
        if 'sid' not in data or 'new_status' not in data or 'reason' not in data:
            return jsonify({"error": "Both 'sid' and 'new_status' are required."}), 400  # Bad Request

        # Find the student based on sid
        students_db = db["students_db"]
        student = students_db.find_one({"sid": data['sid']})

        if not student:
            return jsonify({"error": f"No student found with sid: {data['sid']}"}), 404  # Not Found

        # Update the status of the student
        students_db.update_one({"sid": data['sid']}, {"$set": {"status": data['new_status'],"reason":data['reason']}})

        return jsonify({"message": f"Status for student with sid {data['sid']} updated to {data['new_status']} successfully"})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
#-----------------------------------------------------------------------------------

@app.route('/addDiscountCodes', methods=['POST'])
def add_discount_codes():
    try:
        data = request.get_json()

        # Check if discount_code and discount_amount are provided
        if 'discount_code' not in data or 'discount_amount' not in data:
            return jsonify({"error": "Both 'discount_code' and 'discount_amount' are required."}), 400  # Bad Request

        # Generate a unique ID for the discount using UUID
        discount_id = str(uuid.uuid4().hex)

        # Store discount code information in the MongoDB collection
        discount_codes_db = db["discount_codes_db"]
        discount_codes_db.insert_one({
            "discount_id": discount_id,
            "discount_code": data['discount_code'],
            "discount_amount": data['discount_amount']
        })

        return jsonify({"message": f"Discount code '{data['discount_code']}' added successfully.", "discount_id": discount_id})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error

@app.route('/getAllDiscounts', methods=['GET'])
def get_all_discounts():
    try:
        discount_codes_db = db["discount_codes_db"]
        discounts = list(discount_codes_db.find({}, {"_id": 0}))

        return jsonify({"discounts": discounts})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error

@app.route('/getDiscount', methods=['GET'])
def get_discount():
    try:
        discount_id = request.args.get('discount_id')
        if not discount_id:
            return jsonify({"error": "Missing 'discount_id' parameter in the request."}), 400  # Bad Request

        discount_codes_db = db["discount_codes_db"]
        discount = discount_codes_db.find_one({"discount_id": discount_id}, {"_id": 0})

        if not discount:
            return jsonify({"error": f"No discount found with discount_id: {discount_id}"}), 404  # Not Found

        return jsonify({"discount": discount})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error


@app.route('/updateDiscount', methods=['PUT'])
def update_discount():
    try:
        data = request.get_json()

        # Check if discount_id and new_discount_code are provided
        if 'discount_id' not in data or 'new_discount_code' not in data or "discount_amount" not in data:
            return jsonify({"error": "Both 'discount_id' and 'new_discount_code' and 'discount_amount' are required."}), 400  # Bad Request

        discount_codes_db = db["discount_codes_db"]
        discount = discount_codes_db.find_one({"discount_id": data['discount_id']})

        if not discount:
            return jsonify({"error": f"No discount found with discount_id: {data['discount_id']}"}), 404  # Not Found

        # Update the discount_code of the discount
        discount_codes_db.update_one({"discount_id": data['discount_id']}, {"$set": {"discount_code": data['new_discount_code']},"discount_amount":data["discount_amount"]})

        return jsonify({"message": f"Discount with discount_id {data['discount_id']} updated successfully"})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error

@app.route('/deleteDiscount', methods=['DELETE'])
def delete_discount():
    try:
        discount_id = request.args.get('discount_id')
        if not discount_id:
            return jsonify({"error": "Missing 'discount_id' parameter in the request."}), 400  # Bad Request

        discount_codes_db = db["discount_codes_db"]
        discount = discount_codes_db.find_one({"discount_id": discount_id})

        if not discount:
            return jsonify({"error": f"No discount found with discount_id: {discount_id}"}), 404  # Not Found

        # Delete the discount from the database
        discount_codes_db.delete_one({"discount_id": discount_id})

        return jsonify({"message": f"Discount with discount_id {discount_id} deleted successfully"})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
@app.route('/checkDiscountCode', methods=['POST'])
def check_discount_code():
    try:
        data = request.get_json()

        # Check if discount_code is provided
        if 'discount_code' not in data:
            return jsonify({"error": "Missing 'discount_code' parameter in the request."}), 400  # Bad Request

        discount_codes_db = db["discount_codes_db"]
        discount = discount_codes_db.find_one({"discount_code": data['discount_code']}, {"_id": 0, "discount_amount": 1})

        if discount:
            return jsonify({"message": "Valid code", "discount_amount": discount["discount_amount"],"success":True})
        else:
            return jsonify({"message": "Invalid code","success":False})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    
file_directory = file_dir
    
def save_file(file, uid):
    try:
        # Get the file extension from the original filename
        original_filename = file.filename
        _, file_extension = os.path.splitext(original_filename)

        # Generate a unique filename using UUID and append the original file extension
        filename = str(uuid.uuid4()) + file_extension

        file_path = os.path.join(file_directory, uid, filename)
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        file.save(file_path)

        return f'{files_base_url}{uid}/{filename}'
    except Exception as e:
        raise e
    
@app.route('/uploadFile', methods=['POST'])
def upload_file():
    try:
        # Check if 'file' and 'sid' parameters are present in the form data
        if 'file' not in request.files:
            return jsonify({'error': 'Missing parameters: file',"success":False}), 400

        uploaded_file = request.files['file']
        sid = "All_Files"

        # Check if the file is an allowed type (e.g., image or pdf)
        allowed_extensions = {'png', 'jpg', 'jpeg', 'gif', 'pdf'}
        if (
            '.' in uploaded_file.filename
            and uploaded_file.filename.rsplit('.', 1)[1].lower() not in allowed_extensions
        ):
            return jsonify({'error': 'Invalid file type. Only allowed: png, jpg, jpeg, gif, pdf.',"success":False}), 400

        # Save the file and get the URL
        file_url = save_file(uploaded_file, sid)

        return jsonify({'message': 'File stored successfully.', 'file_url': file_url,"success":True}), 200

    except Exception as e:
        return jsonify({'error': str(e),"success":False}), 500
    
# Endpoint for requesting password reset (for admin)
@app.route("/sendEntranceCard", methods=["GET"])
def send_entrance_card():
    try:
        # collection = db["students_db"]
        sid = request.args.get('sid')
        students_db = db["students_db"]
        student_data = students_db.find_one({"sid":sid}, {"_id":0})
        mailToSend = student_data['email']
        # Send the password reset link via email
        sender_email = "mcfcamp@gmail.com"
        smtp_server = smtplib.SMTP("smtp.gmail.com", 587)
        smtp_server.ehlo()
        smtp_server.starttls()
        smtp_server.login("mcfcamp@gmail.com", "meyvghuponblfqhu")

        message_text = f"Hello {student_data['first_name']}, \n\n You can download your Entrance Card from below Link. \n {student_data['entrence_card']} \n\n You need to print the Entrance Card and Bring to camp in Hardcopy."
        message = MIMEText(message_text)
        message["Subject"] = "MCF Camp Entrance Card"
        message["From"] = sender_email
        message["To"] = mailToSend

        smtp_server.sendmail(sender_email, mailToSend, message.as_string())
        smtp_server.quit()

        return jsonify({'success': True, 'msg': 'Mail Send'}), 200

    except Exception as e:
        return jsonify({'success': False, 'msg': 'Something Went Wrong.', 'reason': str(e)}), 500
    

# Endpoint for requesting password reset (for admin)
@app.route("/sendMedicalCertificate", methods=["GET"])
def send_medical_certificate():
    try:
        # collection = db["students_db"]
        sid = request.args.get('sid')
        students_db = db["students_db"]
        student_data = students_db.find_one({"sid":sid}, {"_id":0})
        mailToSend = student_data['email']
        # Send the password reset link via email
        sender_email = "mcfcamp@gmail.com"
        smtp_server = smtplib.SMTP("smtp.gmail.com", 587)
        smtp_server.ehlo()
        smtp_server.starttls()
        smtp_server.login("mcfcamp@gmail.com", "meyv ghup onbl fqhu")

        message_text = f"Hello {student_data['first_name']}, \n\n You can download your Medical Certificate from below Link. \n {student_data['medicalCertificate']} \n\n You need to print the Medical Certificate and Take a signature from your Doctor and Bring to camp in Hardcopy."
        message = MIMEText(message_text)
        message["Subject"] = "MCF Camp Medical Certificate"
        message["From"] = sender_email
        message["To"] = mailToSend

        smtp_server.sendmail(sender_email, mailToSend, message.as_string())
        smtp_server.quit()

        return jsonify({'success': True, 'msg': 'Mail Send'}), 200

    except Exception as e:
        return jsonify({'success': False, 'msg': 'Something Went Wrong.', 'reason': str(e)}), 500
    
@app.route("/sendVisitingCard", methods=["GET"])
def send_visiting_card():
    try:
        # collection = db["students_db"]
        sid = request.args.get('sid')
        students_db = db["students_db"]
        student_data = students_db.find_one({"sid":sid}, {"_id":0})
        mailToSend = student_data['email']
        # Send the password reset link via email
        sender_email = "mcfcamp@gmail.com"
        smtp_server = smtplib.SMTP("smtp.gmail.com", 587)
        smtp_server.ehlo()
        smtp_server.starttls()
        smtp_server.login("mcfcamp@gmail.com", "meyv ghup onbl fqhu")

        message_text = f"Hello {student_data['first_name']}, \n\n You can download your Visiting Card from below Link. \n {student_data['visiting_card']} \n\n You need to print the Visiting Card and bring to camp for parents to visit Student"
        message = MIMEText(message_text)
        message["Subject"] = "MCF Camp Visiting Card"
        message["From"] = sender_email
        message["To"] = mailToSend

        smtp_server.sendmail(sender_email, mailToSend, message.as_string())
        smtp_server.quit()

        return jsonify({'success': True, 'msg': 'Mail Send'}), 200

    except Exception as e:
        return jsonify({'success': False, 'msg': 'Something Went Wrong.', 'reason': str(e)}), 500
    






# Endpoint for requesting password reset (for admin)
@app.route("/sendEntranceCard_sms", methods=["GET"])
def send_entrance_card_sms():
    try:
        # collection = db["students_db"]
        sid = request.args.get('sid')
        students_db = db["students_db"]
        student_data = students_db.find_one({"sid":sid}, {"_id":0})
        entrence_card = student_data["entrence_card"]
        entrence_card_srt = ''
        url = "https://s.mcfcamp.in/shorten"
        data = {
            "url": entrence_card
        }
        response = requests.post(url, json=data)
        if response.status_code == 200:
            json_data = response.json()
            entrence_card_srt = json_data["short_url"]
            print(entrence_card_srt)
        else:
            print("Error:", response.status_code)
        msg = f"Hello, Download Link for your Entrance Card is {entrence_card_srt} \n Team MCF CAMP"
        phn = student_data['phn']
        sendSMS(msg,phn)

        return jsonify({'success': True, 'msg': 'SMS Send'}), 200

    except Exception as e:
        return jsonify({'success': False, 'msg': 'Something Went Wrong.', 'reason': str(e)}), 500


# import requests
# Endpoint for requesting password reset (for admin)
@app.route("/sendMedicalCertificate_sms", methods=["GET"])
def send_medical_certificate_sms():
    try:
        # collection = db["students_db"]
        sid = request.args.get('sid')
        students_db = db["students_db"]
        student_data = students_db.find_one({"sid":sid}, {"_id":0})
        medical_cert = student_data["medicalCertificate"]
        medical_cert_srt = ''
        url = "https://s.mcfcamp.in/shorten"
        data = {
            "url": medical_cert
        }
        response = requests.post(url, json=data)
        if response.status_code == 200:
            json_data = response.json()
            medical_cert_srt = json_data["short_url"]
            print(medical_cert_srt)
        else:
            print("Error:", response.status_code)
        msg = f"Hello, Download Link for your Medical Certificate is {medical_cert_srt} \n Team MCF CAMP"
        phn = student_data['phn']
        sendSMS(msg,phn)

        return jsonify({'success': True, 'msg': 'SMS Send'}), 200

    except Exception as e:
        return jsonify({'success': False, 'msg': 'Something Went Wrong.', 'reason': str(e)}), 500
        
    
@app.route("/sendVisitingCard_sms", methods=["GET"])
def send_visiting_card_sms():
    try:
        # collection = db["students_db"]
        sid = request.args.get('sid')
        students_db = db["students_db"]
        student_data = students_db.find_one({"sid":sid}, {"_id":0})
        visiting_card = student_data["visiting_card"]
        visiting_card_srt = ''
        url = "https://s.mcfcamp.in/shorten"
        data = {
            "url": visiting_card
        }
        response = requests.post(url, json=data)
        if response.status_code == 200:
            json_data = response.json()
            visiting_card_srt = json_data["short_url"]
            print(visiting_card_srt)
        else:
            print("Error:", response.status_code)
        msg = f"Hello, Download Link for your Visiting Card is {visiting_card_srt} \n Team MCF CAMP"
        phn = student_data['phn']
        sendSMS(msg,phn)

        return jsonify({'success': True, 'msg': 'SMS Send'}), 200

    except Exception as e:
        return jsonify({'success': False, 'msg': 'Something Went Wrong.', 'reason': str(e)}), 500









@app.route("/sendEntranceCard_wp", methods=["GET"])
def send_entrance_card_wp():
    try:
        # collection = db["students_db"]
        sid = request.args.get('sid')
        students_db = db["students_db"]
        student_data = students_db.find_one({"sid":sid}, {"_id":0})
        entrence_card = student_data["entrence_card"]
        ec = entrence_card.replace(files_url,files_base_dir)
        entrence_card_srt = ''
        url = "https://s.mcfcamp.in/shorten"
        data = {
            "url": entrence_card
        }
        response = requests.post(url, json=data)
        if response.status_code == 200:
            json_data = response.json()
            entrence_card_srt = json_data["short_url"]
            print(entrence_card_srt)
        else:
            print("Error:", response.status_code)
        msg = f"Hello, Download Link for your Entrance Card is {entrence_card_srt} \n Team MCF CAMP"
        phn = student_data['wp_no']
        send_wp(msg,phn,file_paths=[ec])

        return jsonify({'success': True, 'msg': 'SMS Send'}), 200

    except Exception as e:
        return jsonify({'success': False, 'msg': 'Something Went Wrong.', 'reason': str(e)}), 500
    

@app.route("/sendMedicalCertificate_wp", methods=["GET"])
def send_medical_certificate_wp():
    try:
        # collection = db["students_db"]
        sid = request.args.get('sid')
        students_db = db["students_db"]
        student_data = students_db.find_one({"sid":sid}, {"_id":0})
        medical_cert = student_data["medicalCertificate"]
        mcert = medical_cert.replace(files_url,files_base_dir)
        medical_cert_srt = ''
        url = "https://s.mcfcamp.in/shorten"
        data = {
            "url": medical_cert
        }
        response = requests.post(url, json=data)
        if response.status_code == 200:
            json_data = response.json()
            medical_cert_srt = json_data["short_url"]
            print(medical_cert_srt)
        else:
            print("Error:", response.status_code)
        msg = f"Hello, Download Link for your Medical Certificate is {medical_cert_srt} \n Team MCF CAMP"
        phn = student_data['wp_no']
        send_wp(msg,phn,file_paths=[mcert])

        return jsonify({'success': True, 'msg': 'SMS Send'}), 200

    except Exception as e:
        return jsonify({'success': False, 'msg': 'Something Went Wrong.', 'reason': str(e)}), 500
    

@app.route("/sendVisitingCard_wp", methods=["GET"])
def send_visiting_card_wp():
    try:
        # collection = db["students_db"]
        sid = request.args.get('sid')
        students_db = db["students_db"]
        student_data = students_db.find_one({"sid":sid}, {"_id":0})
        visiting_card = student_data["visiting_card"]
        vc = visiting_card.replace(files_url,files_base_dir)
        visiting_card_srt = ''
        url = "https://s.mcfcamp.in/shorten"
        data = {
            "url": visiting_card
        }
        response = requests.post(url, json=data)
        if response.status_code == 200:
            json_data = response.json()
            visiting_card_srt = json_data["short_url"]
            print(visiting_card_srt)
        else:
            print("Error:", response.status_code)
        msg = f"Hello, Download Link for your Visiting Card is {visiting_card_srt} \n Team MCF CAMP"
        phn = student_data['wp_no']
        send_wp(msg,phn,file_paths=[vc])

        return jsonify({'success': True, 'msg': 'SMS Send'}), 200

    except Exception as e:
        return jsonify({'success': False, 'msg': 'Something Went Wrong.', 'reason': str(e)}), 500





@app.route("/sendReceipt_wp", methods=["GET"])
def sendReceipt_wp():
    try:
        # collection = db["students_db"]
        payment_id = request.args.get('payment_id')
        sid = request.args.get('sid')
        payments_db = db["all_payments"]
        payment_data = payments_db.find_one({"payment_id":payment_id}, {"_id":0})
        students_db = db["students_db"]
        student_data = students_db.find_one({"sid":sid}, {"_id":0})

        receipt = payment_data["receipt_url"]
        payment_receipt = receipt.replace(files_url,files_base_dir)
        receipt_srt = ''
        url = "https://s.mcfcamp.in/shorten"
        data = {
            "url": receipt
        }
        response = requests.post(url, json=data)
        if response.status_code == 200:
            json_data = response.json()
            receipt_srt = json_data["short_url"]
            print(receipt_srt)
        else:
            print("Error:", response.status_code)
        msg = f"Hello, Download Link for your Payment Receipt is {receipt_srt} \n Team MCF CAMP"
        phn = student_data['wp_no']
        send_wp(msg,phn,file_paths=[payment_receipt])

        return jsonify({'success': True, 'msg': 'SMS Send'}), 200

    except Exception as e:
        return jsonify({'success': False, 'msg': 'Something Went Wrong.', 'reason': str(e)}), 500
    


@app.route("/sendReceipt_email", methods=["GET"])
def sendReceipt_email():
    try:
        # collection = db["students_db"]
        payment_id = request.args.get('payment_id')
        sid = request.args.get('sid')
        payments_db = db["all_payments"]
        payment_data = payments_db.find_one({"payment_id":payment_id}, {"_id":0})
        students_db = db["students_db"]
        student_data = students_db.find_one({"sid":sid}, {"_id":0})

        receipt = payment_data["receipt_url"]
        receipt_srt = ''
        url = "https://s.mcfcamp.in/shorten"
        data = {
            "url": receipt
        }
        response = requests.post(url, json=data)
        if response.status_code == 200:
            json_data = response.json()
            receipt_srt = json_data["short_url"]
            print(receipt_srt)
        else:
            print("Error:", response.status_code)
        msg = f"Hello, Download Link for your Payment Receipt is {receipt_srt} \n Team MCF CAMP"
        email = student_data['email']
        fn = receipt.replace(files_url,files_base_dir)
        send_email_attachments(msg=msg, sub="Payment Receipt Download Link", mailToSend=email, files=[fn])

        return jsonify({'success': True, 'msg': 'SMS Send'}), 200

    except Exception as e:
        return jsonify({'success': False, 'msg': 'Something Went Wrong.', 'reason': str(e)}), 500
        
    




#----------------------------------------------------------------------------------
    


@app.route('/createPayment', methods=['POST'])
def createPayment():
    try:
        data = request.get_json()
        if 'payment_option' not in data or 'payment_amount' not in data or 'sid' not in data:
            return jsonify({"error": "Both 'payment_option' and 'payment_amount' and 'sid' are required."}), 400
        payment_db = db["all_payments"]
        payment = payment_db.find_one({"transaction_id":data['transaction_id']})
        if payment:
            return jsonify({"message": f"Payment with transaction_id - {data['transaction_id']} already exist."}), 400
        payment_id = str(uuid.uuid4().hex)
        receipt_no = str(uuid.uuid4().hex)[:10]
        all_payments = db["all_payments"]
        students_db = db['students_db']
        student_data = students_db.find_one({"sid":data['sid']}, {"_id":0})
        total_paid = student_data['total_amount_paid']
        total_paid = float(total_paid) + float(data['payment_amount'])
        student_data['total_amount_paid'] = total_paid
        if float(total_paid) >= float(student_data['total_amount_payable']):
            student_data['status'] = "Active"
            
        students_db.update_one({"sid": data['sid']}, {"$set": student_data})
        # Load the document template

        batch_id = student_data['batch_id']
        batch_db = db["batches_db"]
        batch_data = batch_db.find_one({"batch_id":batch_id}, {"_id":0})
        batch_dur = batch_data['duration']

        camp_id = student_data['camp_id']
        camp_db = db["camps_db"]
        camp_data = camp_db.find_one({"camp_id":camp_id}, {"_id":0})

        if "7" in batch_dur or "5" in batch_dur or "3" in batch_dur:
            doc = Document('fee_receipt_7.docx')
            final_data = {
                "REG_NO": student_data['sid'],
                "RECEIPT_NO": receipt_no,
                "DATE":data["payment_date"],
                "C_BATCH":batch_data['batch_name'],
                "CADET_NAME": str(student_data['first_name']+" "+student_data['last_name']),
                "ADDRESS": student_data['address'],
                "CONTACT_NO": student_data['phn'],
                "CAMP_NAME":camp_data['camp_name'],
                "CAMP_DATE":batch_data['start_date'],
                "TRANS_ID":data['transaction_id'],
                "TRANS_AMT":data['payment_amount'],
                "PAY_MODE":data["payment_mode"],
                "AMOUNT":camp_data['camp_fee'],
                "DISCOUNT":student_data['discount_amount'],
                "PAYABLE_FEES":student_data['total_amount_payable'],
                "NET_PAID":total_paid,
                "BALANCE":float(student_data['total_amount_payable'])-float(total_paid),
                "AMOUNT_INWORDS":number_to_words(int(data['payment_amount'])),
                "EMPLOYEE_NAME":student_data['employee_who_reached_out_to_you'],
                "1_INST_AMT":"-",
                "1_INST_DT":"-",
                "2_INST_AMT":"-",
                "2_INST_DT":"-"
                }
            
            if "totalPayment" in data['payment_option']:
                for key, value in final_data.items():
                        find_and_replace_tables_fee_receipt(doc.tables, f'{{MERGEFIELD {key}}}', str(value))
                doc.save(str(file_dir)+str(f"{student_data['sid']}_fee_receipt_{data['payment_option']}.docx"))

                convert_to_pdf(str(str(file_dir)+str(f"{student_data['sid']}_fee_receipt_{data['payment_option']}.docx")), str(str(file_dir)+str(f"{student_data['sid']}_fee_receipt_{data['payment_option']}.pdf")))
                

                if float(total_paid) >= float(student_data['total_amount_payable']):
                    camp_id = student_data['camp_id']
                    camp_db = db["camps_db"]
                    camp_data = camp_db.find_one({"camp_id":camp_id}, {"_id":0})
                    batch_id = student_data['batch_id']
                    batch_db = db["batches_db"]
                    batch_data = batch_db.find_one({"batch_id":batch_id}, {"_id":0})
                    payment_db = db["all_payments"]
                    payment_data = payment_db.find({"sid":data['sid']}, {"_id":0})
                    receipt_nos = ""
                    payment_data = list(payment_data)
                    print(len(payment_data))
                    for receipt in payment_data:
                        receipt_nos = str(receipt_nos + str(str(receipt['receipt_no'])+ " , "))
                        print(receipt_nos)
                    receipt_nos = str(receipt_nos + receipt_no)

                    student_data_1 = {
                            'CADET_NAME': str(student_data["first_name"].upper()+" "+student_data["last_name"].upper()),
                            'REGNO': data['sid'],
                            'RANK': 'CDT',
                            'C_NAME': camp_data['camp_name'],
                            'C_BATCH': batch_data['batch_name'],
                            'C_DAYS': batch_data["duration"],
                            'COMP_N': student_data['company'],
                            'C_DATE': batch_data['start_date'],
                            'PICKPT': student_data["pick_up_point"],
                            'PICK_TIME': '',
                            'EMP_NAME':  student_data["employee_who_reached_out_to_you"],
                            'GAR_NAME': student_data["middle_name"],
                            'ADDRESS': student_data["address"],
                            'CITY': student_data["pick_up_city"],
                            'DISTRICT': student_data["district"],
                            'STATE': student_data['state'],
                            'PINCODE': student_data["pincode"],
                            'EMAIL': student_data["email"],
                            'C_NUM': str(student_data["phn"]),
                            'WP_NUM': student_data.get("wp_no", ""),
                            'FATHER_NUM': '',
                            'MOTHER_NUM': '',
                            'DOB': student_data["dob"],
                            'BLOOD_GROUP':  student_data.get("blood_group", ""),
                            'STD': student_data.get("standard", ""),
                            'SCHOOL': student_data.get("school_name", ""),
                            'FEE_PAID': student_data['total_amount_paid'],
                            'BALANCE': float(student_data['total_amount_payable']) - float(student_data['total_amount_paid']),
                            'RECEIPT_NUM': receipt_nos,
                            'DATE': '',
                            'TIME': ''
                        }

                    doc = Document('mcf_entrance_card.docx')

                    # Replace text fields in paragraphs
                    find_and_replace_paragraphs_entrance_card(doc.paragraphs, '{MERGEFIELD CADET_NAME}', student_data_1['CADET_NAME'])
                    find_and_replace_paragraphs_entrance_card(doc.paragraphs, '{MERGEFIELD DATE}', student_data_1['DATE'])
                    find_and_replace_paragraphs_entrance_card(doc.paragraphs, '{MERGEFIELD TIME}', student_data_1['TIME'])

                    for key, value in student_data_1.items():
                            find_and_replace_tables_entrance_card(doc.tables, f'{{MERGEFIELD {key}}}', str(value))

                    try:
                        print("replacing image")
                        table = doc.tables[0]  # Assuming the first table
                        cell = table.cell(0, 3)  # Assuming the first cell in the third column

                        # Clear the content of the cell by removing its paragraphs
                        for paragraph in cell.paragraphs:
                            paragraph.clear()

                        # Add a new paragraph and insert the image
                        paragraph = cell.add_paragraph()
                        run = paragraph.add_run()
                        cadet_photo_url = student_data["cadetPhoto"]
                        cadet_photo_path = cadet_photo_url.replace(files_url,files_base_dir)
                        run.add_picture(cadet_photo_path, width=Inches(0.9))
                    except Exception as e:
                        print("Error : ",str(e))

                    doc.save(str(str(file_dir)+f"{data['sid']}_entrance_card.docx"))

                    convert_to_pdf(str(str(file_dir)+f"{data['sid']}_entrance_card.docx"), str(str(file_dir)+f"{data['sid']}_entrance_card.pdf"))

                    ec = f"{files_base_url}{data['sid']}_entrance_card.pdf"

                    entrance_card = {
                        "entrence_card" : ec
                    }
                    

                    students_db.update_one({"sid": data['sid']}, {"$set": entrance_card})

                    payment_receipt_url = f"{files_base_url}{student_data['sid']}_fee_receipt_{data['payment_option']}.pdf"

                    msg = f"Hello,\n {student_data['first_name']} {student_data['last_name']} your admission for camp - {camp_data['camp_name']} and Batch - {batch_data['batch_name']} for Duration of {batch_data['duration']} is Succesfull. \n Following are the Documents and List of Things to Bring for Camp. \n Download Links for Your Documents are Shared Below : \nPayment Receipt - {payment_receipt_url}\n Medical Certificate - {student_data['medicalCertificate']} \nEntrance Card - {ec} \nVisiting Card - {student_data['visiting_card']} \nAdmission Form - {student_data['admission_form']}\n\nTeam MCF Camp"

                    fns = []
                    payment_receipt = payment_receipt_url.replace(files_url,files_base_dir)
                    mcert = student_data['medicalCertificate'].replace(files_url,files_base_dir)
                    ecfn = ec.replace(files_url,files_base_dir)
                    vc = student_data['visiting_card'].replace(files_url,files_base_dir)
                    af = student_data['admission_form'].replace(files_url,files_base_dir)
                    fns = [payment_receipt, mcert, ecfn, vc, af]
            
                    send_email_attachments(msg=msg, sub="Payment Receipt and Other Documents", mailToSend=student_data['email'], files=fns)
                    send_wp(msg,student_data['wp_no'],file_paths=fns)

                    send_wp(msg,"9604084000")
                    send_email(msg, f"Payment Complete of Student - {student_data['first_name']} {student_data['last_name']}", "infomcfcamp@gmail.com")


                all_payments.insert_one({
                "payment_id": payment_id,
                "payment_option": data['payment_option'],
                "payment_amount": data['payment_amount'],
                "payment_date":data["payment_date"],
                "transaction_id":data["transaction_id"],
                "payment_mode":data["payment_mode"],
                "sid":data['sid'],
                "receipt_no":receipt_no,
                "receipt_url":f"{files_base_url}{student_data['sid']}_fee_receipt_{data['payment_option']}.pdf"
            })
                
            elif "1installment" in data['payment_option']:
                final_data['1_INST_AMT'] = data['payment_amount']
                final_data['1_INST_DT'] = data["payment_date"]
                for key, value in final_data.items():
                        find_and_replace_tables_fee_receipt(doc.tables, f'{{MERGEFIELD {key}}}', str(value))
                doc.save(str(file_dir)+str(f"{student_data['sid']}_fee_receipt_{data['payment_option']}.docx"))

                convert_to_pdf(str(str(file_dir)+str(f"{student_data['sid']}_fee_receipt_{data['payment_option']}.docx")), str(str(file_dir)+str(f"{student_data['sid']}_fee_receipt_{data['payment_option']}.pdf")))


                if float(total_paid) >= float(student_data['total_amount_payable']):
                    camp_id = student_data['camp_id']
                    camp_db = db["camps_db"]
                    camp_data = camp_db.find_one({"camp_id":camp_id}, {"_id":0})
                    batch_id = student_data['batch_id']
                    batch_db = db["batches_db"]
                    batch_data = batch_db.find_one({"batch_id":batch_id}, {"_id":0})
                    payment_db = db["all_payments"]
                    payment_data = payment_db.find({"sid":data['sid']}, {"_id":0})
                    receipt_nos = ""
                    payment_data = list(payment_data)
                    print(len(payment_data))
                    for receipt in payment_data:
                        receipt_nos = str(receipt_nos + str(str(receipt['receipt_no'])+ " , "))
                        print(receipt_nos)
                    receipt_nos = str(receipt_nos + receipt_no)

                    student_data_1 = {
                            'CADET_NAME': str(student_data["first_name"].upper()+" "+student_data["last_name"].upper()),
                            'REGNO': data['sid'],
                            'RANK': 'CDT',
                            'C_NAME': camp_data['camp_name'],
                            'C_BATCH': batch_data['batch_name'],
                            'C_DAYS': batch_data["duration"],
                            'COMP_N': student_data['company'],
                            'C_DATE': batch_data['start_date'],
                            'PICKPT': student_data["pick_up_point"],
                            'PICK_TIME': '',
                            'EMP_NAME':  student_data["employee_who_reached_out_to_you"],
                            'GAR_NAME': student_data["middle_name"],
                            'ADDRESS': student_data["address"],
                            'CITY': student_data["pick_up_city"],
                            'DISTRICT': student_data["district"],
                            'STATE': student_data['state'],
                            'PINCODE': student_data["pincode"],
                            'EMAIL': student_data["email"],
                            'C_NUM': str(student_data["phn"]),
                            'WP_NUM': student_data.get("wp_no", ""),
                            'FATHER_NUM': '',
                            'MOTHER_NUM': '',
                            'DOB': student_data["dob"],
                            'BLOOD_GROUP':  student_data.get("blood_group", ""),
                            'STD': student_data.get("standard", ""),
                            'SCHOOL': student_data.get("school_name", ""),
                            'FEE_PAID': student_data['total_amount_paid'],
                            'BALANCE': float(student_data['total_amount_payable']) - float(student_data['total_amount_paid']),
                            'RECEIPT_NUM': receipt_nos,
                            'DATE': '',
                            'TIME': ''
                        }

                    doc = Document('mcf_entrance_card.docx')

                    # Replace text fields in paragraphs
                    find_and_replace_paragraphs_entrance_card(doc.paragraphs, '{MERGEFIELD CADET_NAME}', student_data_1['CADET_NAME'])
                    find_and_replace_paragraphs_entrance_card(doc.paragraphs, '{MERGEFIELD DATE}', student_data_1['DATE'])
                    find_and_replace_paragraphs_entrance_card(doc.paragraphs, '{MERGEFIELD TIME}', student_data_1['TIME'])

                    for key, value in student_data_1.items():
                            find_and_replace_tables_entrance_card(doc.tables, f'{{MERGEFIELD {key}}}', str(value))

                    try:
                        print("replacing image")
                        table = doc.tables[0]  # Assuming the first table
                        cell = table.cell(0, 3)  # Assuming the first cell in the third column

                        # Clear the content of the cell by removing its paragraphs
                        for paragraph in cell.paragraphs:
                            paragraph.clear()

                        # Add a new paragraph and insert the image
                        paragraph = cell.add_paragraph()
                        run = paragraph.add_run()
                        cadet_photo_url = student_data["cadetPhoto"]
                        cadet_photo_path = cadet_photo_url.replace(files_url,files_base_dir)
                        run.add_picture(cadet_photo_path, width=Inches(0.9))
                    except Exception as e:
                        print("Error : ",str(e))

                    doc.save(str(str(file_dir)+f"{data['sid']}_entrance_card.docx"))

                    convert_to_pdf(str(str(file_dir)+f"{data['sid']}_entrance_card.docx"), str(str(file_dir)+f"{data['sid']}_entrance_card.pdf"))

                    ec = f"{files_base_url}{data['sid']}_entrance_card.pdf"

                    entrance_card = {
                        "entrence_card" : ec
                    }
                    

                    students_db.update_one({"sid": data['sid']}, {"$set": entrance_card})

                    payment_receipt_url = f"{files_base_url}{student_data['sid']}_fee_receipt_{data['payment_option']}.pdf"

                    msg = f"Hello,\n {student_data['first_name']} {student_data['last_name']} your admission for camp - {camp_data['camp_name']} and Batch - {batch_data['batch_name']} for Duration of {batch_data['duration']} is Succesfull. \n Following are the Documents and List of Things to Bring for Camp. \n Download Links for Your Documents are Shared Below : \nPayment Receipt - {payment_receipt_url}\n Medical Certificate - {student_data['medicalCertificate']} \nEntrance Card - {ec} \nVisiting Card - {student_data['visiting_card']} \nAdmission Form - {student_data['admission_form']}\n\nTeam MCF Camp"

                    fns = []
                    payment_receipt = payment_receipt_url.replace(files_url,files_base_dir)
                    mcert = student_data['medicalCertificate'].replace(files_url,files_base_dir)
                    ecfn = ec.replace(files_url,files_base_dir)
                    vc = student_data['visiting_card'].replace(files_url,files_base_dir)
                    af = student_data['admission_form'].replace(files_url,files_base_dir)
                    fns = [payment_receipt, mcert, ecfn, vc, af]
            
                    send_email_attachments(msg=msg, sub="Payment Receipt and Other Documents", mailToSend=student_data['email'], files=fns)
                    send_wp(msg,student_data['wp_no'],file_paths=fns)

                    send_wp(msg,"9604084000")
                    send_email(msg, f"Payment Complete of Student - {student_data['first_name']} {student_data['last_name']}", "infomcfcamp@gmail.com")


                all_payments.insert_one({
                "payment_id": payment_id,
                "payment_option": data['payment_option'],
                "payment_amount": data['payment_amount'],
                "payment_date":data["payment_date"],
                "transaction_id":data["transaction_id"],
                "payment_mode":data["payment_mode"],
                "sid":data['sid'],
                "receipt_no":receipt_no,
                "receipt_url":f"{files_base_url}{student_data['sid']}_fee_receipt_{data['payment_option']}.pdf"
            })
                

            elif '2installment' in data['payment_option']:
                final_data['2_INST_AMT'] = data['payment_amount']
                final_data['2_INST_DT'] = data["payment_date"]
                for key, value in final_data.items():
                        find_and_replace_tables_fee_receipt(doc.tables, f'{{MERGEFIELD {key}}}', str(value))
                doc.save(str(file_dir)+str(f"{student_data['sid']}_fee_receipt_{data['payment_option']}.docx"))

                convert_to_pdf(str(str(file_dir)+str(f"{student_data['sid']}_fee_receipt_{data['payment_option']}.docx")), str(str(file_dir)+str(f"{student_data['sid']}_fee_receipt_{data['payment_option']}.pdf")))


                if float(total_paid) >= float(student_data['total_amount_payable']):
                    camp_id = student_data['camp_id']
                    camp_db = db["camps_db"]
                    camp_data = camp_db.find_one({"camp_id":camp_id}, {"_id":0})
                    batch_id = student_data['batch_id']
                    batch_db = db["batches_db"]
                    batch_data = batch_db.find_one({"batch_id":batch_id}, {"_id":0})
                    payment_db = db["all_payments"]
                    payment_data = payment_db.find({"sid":data['sid']}, {"_id":0})
                    receipt_nos = ""
                    payment_data = list(payment_data)
                    print(len(payment_data))
                    for receipt in payment_data:
                        receipt_nos = str(receipt_nos + str(str(receipt['receipt_no'])+ " , "))
                        print(receipt_nos)
                    receipt_nos = str(receipt_nos + receipt_no)

                    student_data_1 = {
                            'CADET_NAME': str(student_data["first_name"].upper()+" "+student_data["last_name"].upper()),
                            'REGNO': data['sid'],
                            'RANK': 'CDT',
                            'C_NAME': camp_data['camp_name'],
                            'C_BATCH': batch_data['batch_name'],
                            'C_DAYS': batch_data["duration"],
                            'COMP_N': student_data['company'],
                            'C_DATE': batch_data['start_date'],
                            'PICKPT': student_data["pick_up_point"],
                            'PICK_TIME': '',
                            'EMP_NAME':  student_data["employee_who_reached_out_to_you"],
                            'GAR_NAME': student_data["middle_name"],
                            'ADDRESS': student_data["address"],
                            'CITY': student_data["pick_up_city"],
                            'DISTRICT': student_data["district"],
                            'STATE': student_data['state'],
                            'PINCODE': student_data["pincode"],
                            'EMAIL': student_data["email"],
                            'C_NUM': str(student_data["phn"]),
                            'WP_NUM': student_data.get("wp_no", ""),
                            'FATHER_NUM': '',
                            'MOTHER_NUM': '',
                            'DOB': student_data["dob"],
                            'BLOOD_GROUP':  student_data.get("blood_group", ""),
                            'STD': student_data.get("standard", ""),
                            'SCHOOL': student_data.get("school_name", ""),
                            'FEE_PAID': student_data['total_amount_paid'],
                            'BALANCE': float(student_data['total_amount_payable']) - float(student_data['total_amount_paid']),
                            'RECEIPT_NUM': receipt_nos,
                            'DATE': '',
                            'TIME': ''
                        }

                    doc = Document('mcf_entrance_card.docx')

                    # Replace text fields in paragraphs
                    find_and_replace_paragraphs_entrance_card(doc.paragraphs, '{MERGEFIELD CADET_NAME}', student_data_1['CADET_NAME'])
                    find_and_replace_paragraphs_entrance_card(doc.paragraphs, '{MERGEFIELD DATE}', student_data_1['DATE'])
                    find_and_replace_paragraphs_entrance_card(doc.paragraphs, '{MERGEFIELD TIME}', student_data_1['TIME'])

                    for key, value in student_data_1.items():
                            find_and_replace_tables_entrance_card(doc.tables, f'{{MERGEFIELD {key}}}', str(value))

                    try:
                        print("replacing image")
                        table = doc.tables[0]  # Assuming the first table
                        cell = table.cell(0, 3)  # Assuming the first cell in the third column

                        # Clear the content of the cell by removing its paragraphs
                        for paragraph in cell.paragraphs:
                            paragraph.clear()

                        # Add a new paragraph and insert the image
                        paragraph = cell.add_paragraph()
                        run = paragraph.add_run()
                        cadet_photo_url = student_data["cadetPhoto"]
                        cadet_photo_path = cadet_photo_url.replace(files_url,files_base_dir)
                        run.add_picture(cadet_photo_path, width=Inches(0.9))
                    except Exception as e:
                        print("Error : ",str(e))

                    doc.save(str(str(file_dir)+f"{data['sid']}_entrance_card.docx"))

                    convert_to_pdf(str(str(file_dir)+f"{data['sid']}_entrance_card.docx"), str(str(file_dir)+f"{data['sid']}_entrance_card.pdf"))

                    ec = f"{files_base_url}{data['sid']}_entrance_card.pdf"

                    entrance_card = {
                        "entrence_card" : ec
                    }
                    

                    students_db.update_one({"sid": data['sid']}, {"$set": entrance_card})

                    payment_receipt_url = f"{files_base_url}{student_data['sid']}_fee_receipt_{data['payment_option']}.pdf"

                    msg = f"Hello,\n {student_data['first_name']} {student_data['last_name']} your admission for camp - {camp_data['camp_name']} and Batch - {batch_data['batch_name']} for Duration of {batch_data['duration']} is Succesfull. \n Following are the Documents and List of Things to Bring for Camp. \n Download Links for Your Documents are Shared Below : \nPayment Receipt - {payment_receipt_url}\n Medical Certificate - {student_data['medicalCertificate']} \nEntrance Card - {ec} \nVisiting Card - {student_data['visiting_card']} \nAdmission Form - {student_data['admission_form']}\n\nTeam MCF Camp"

                    fns = []
                    payment_receipt = payment_receipt_url.replace(files_url,files_base_dir)
                    mcert = student_data['medicalCertificate'].replace(files_url,files_base_dir)
                    ecfn = ec.replace(files_url,files_base_dir)
                    vc = student_data['visiting_card'].replace(files_url,files_base_dir)
                    af = student_data['admission_form'].replace(files_url,files_base_dir)
                    fns = [payment_receipt, mcert, ecfn, vc, af]
            
                    send_email_attachments(msg=msg, sub="Payment Receipt and Other Documents", mailToSend=student_data['email'], files=fns)
                    send_wp(msg,student_data['wp_no'],file_paths=fns)

                    send_wp(msg,"9604084000")
                    send_email(msg, f"Payment Complete of Student - {student_data['first_name']} {student_data['last_name']}", "infomcfcamp@gmail.com")


                all_payments.insert_one({
                "payment_id": payment_id,
                "payment_option": data['payment_option'],
                "payment_amount": data['payment_amount'],
                "payment_date":data["payment_date"],
                "transaction_id":data["transaction_id"],
                "payment_mode":data["payment_mode"],
                "sid":data['sid'],
                "receipt_no":receipt_no,
                "receipt_url":f"{files_base_url}{student_data['sid']}_fee_receipt_{data['payment_option']}.pdf"
            })
                

            else:
                return {"error":"Please Specify Payment Option"}

            return jsonify({"message": f"Payment added successfully.", "payment_id": payment_id})
        
        elif "15" in batch_dur or "30" in batch_dur:
            doc = Document('fee_receipt_15_30.docx')

            final_data = {
                "REG_NO": student_data['sid'],
                "RECEIPT_NO": receipt_no,
                "DATE":data["payment_date"],
                "C_BATCH":batch_data['batch_name'],
                "CADET_NAME": str(student_data['first_name']+" "+student_data['last_name']),
                "ADDRESS": student_data['address'],
                "CONTACT_NO": student_data['phn'],
                "CAMP_NAME":camp_data['camp_name'],
                "CAMP_DATE":batch_data['start_date'],
                "TRANS_ID":data['transaction_id'],
                "TRANS_AMT":data['payment_amount'],
                "PAY_MODE":data["payment_mode"],
                "AMOUNT":camp_data['camp_fee'],
                "DISCOUNT":student_data['discount_amount'],
                "PAYABLE_FEES":student_data['total_amount_payable'],
                "NET_PAID":total_paid,
                "BALANCE":float(student_data['total_amount_payable'])-float(total_paid),
                "AMOUNT_INWORDS":number_to_words(int(data['payment_amount'])),
                "EMPLOYEE_NAME":student_data['employee_who_reached_out_to_you'],
                "1_INST_AMT":"-",
                "1_INST_DT":"-",
                "2_INST_AMT":"-",
                "2_INST_DT":"-"
                }
            
            if "totalPayment" in data['payment_option']:
                for key, value in final_data.items():
                        find_and_replace_tables_fee_receipt(doc.tables, f'{{MERGEFIELD {key}}}', str(value))
                doc.save(str(file_dir)+str(f"{student_data['sid']}_fee_receipt_{data['payment_option']}.docx"))

                convert_to_pdf(str(str(file_dir)+str(f"{student_data['sid']}_fee_receipt_{data['payment_option']}.docx")), str(str(file_dir)+str(f"{student_data['sid']}_fee_receipt_{data['payment_option']}.pdf")))


                if float(total_paid) >= float(student_data['total_amount_payable']):
                    camp_id = student_data['camp_id']
                    camp_db = db["camps_db"]
                    camp_data = camp_db.find_one({"camp_id":camp_id}, {"_id":0})
                    batch_id = student_data['batch_id']
                    batch_db = db["batches_db"]
                    batch_data = batch_db.find_one({"batch_id":batch_id}, {"_id":0})
                    payment_db = db["all_payments"]
                    payment_data = payment_db.find({"sid":data['sid']}, {"_id":0})
                    receipt_nos = ""
                    payment_data = list(payment_data)
                    print(len(payment_data))
                    for receipt in payment_data:
                        receipt_nos = str(receipt_nos + str(str(receipt['receipt_no'])+ " , "))
                        print(receipt_nos)
                    receipt_nos = str(receipt_nos + receipt_no)

                    student_data_1 = {
                            'CADET_NAME': str(student_data["first_name"].upper()+" "+student_data["last_name"].upper()),
                            'REGNO': data['sid'],
                            'RANK': 'CDT',
                            'C_NAME': camp_data['camp_name'],
                            'C_BATCH': batch_data['batch_name'],
                            'C_DAYS': batch_data["duration"],
                            'COMP_N': student_data['company'],
                            'C_DATE': batch_data['start_date'],
                            'PICKPT': student_data["pick_up_point"],
                            'PICK_TIME': '',
                            'EMP_NAME':  student_data["employee_who_reached_out_to_you"],
                            'GAR_NAME': student_data["middle_name"],
                            'ADDRESS': student_data["address"],
                            'CITY': student_data["pick_up_city"],
                            'DISTRICT': student_data["district"],
                            'STATE': student_data['state'],
                            'PINCODE': student_data["pincode"],
                            'EMAIL': student_data["email"],
                            'C_NUM': str(student_data["phn"]),
                            'WP_NUM': student_data.get("wp_no", ""),
                            'FATHER_NUM': '',
                            'MOTHER_NUM': '',
                            'DOB': student_data["dob"],
                            'BLOOD_GROUP':  student_data.get("blood_group", ""),
                            'STD': student_data.get("standard", ""),
                            'SCHOOL': student_data.get("school_name", ""),
                            'FEE_PAID': student_data['total_amount_paid'],
                            'BALANCE': float(student_data['total_amount_payable']) - float(student_data['total_amount_paid']),
                            'RECEIPT_NUM': receipt_nos,
                            'DATE': '',
                            'TIME': ''
                        }

                    doc = Document('mcf_entrance_card.docx')

                    # Replace text fields in paragraphs
                    find_and_replace_paragraphs_entrance_card(doc.paragraphs, '{MERGEFIELD CADET_NAME}', student_data_1['CADET_NAME'])
                    find_and_replace_paragraphs_entrance_card(doc.paragraphs, '{MERGEFIELD DATE}', student_data_1['DATE'])
                    find_and_replace_paragraphs_entrance_card(doc.paragraphs, '{MERGEFIELD TIME}', student_data_1['TIME'])

                    for key, value in student_data_1.items():
                            find_and_replace_tables_entrance_card(doc.tables, f'{{MERGEFIELD {key}}}', str(value))

                    try:
                        print("replacing image")
                        table = doc.tables[0]  # Assuming the first table
                        cell = table.cell(0, 3)  # Assuming the first cell in the third column

                        # Clear the content of the cell by removing its paragraphs
                        for paragraph in cell.paragraphs:
                            paragraph.clear()

                        # Add a new paragraph and insert the image
                        paragraph = cell.add_paragraph()
                        run = paragraph.add_run()
                        cadet_photo_url = student_data["cadetPhoto"]
                        cadet_photo_path = cadet_photo_url.replace(files_url,files_base_dir)
                        run.add_picture(cadet_photo_path, width=Inches(0.9))
                    except Exception as e:
                        print("Error : ",str(e))

                    doc.save(str(str(file_dir)+f"{data['sid']}_entrance_card.docx"))

                    convert_to_pdf(str(str(file_dir)+f"{data['sid']}_entrance_card.docx"), str(str(file_dir)+f"{data['sid']}_entrance_card.pdf"))

                    ec = f"{files_base_url}{data['sid']}_entrance_card.pdf"

                    entrance_card = {
                        "entrence_card" : ec
                    }
                    

                    students_db.update_one({"sid": data['sid']}, {"$set": entrance_card})

                    payment_receipt_url = f"{files_base_url}{student_data['sid']}_fee_receipt_{data['payment_option']}.pdf"
                    
                    msg = f"Hello,\n {student_data['first_name']} {student_data['last_name']} your admission for camp - {camp_data['camp_name']} and Batch - {batch_data['batch_name']} for Duration of {batch_data['duration']} is Succesfull. \n Following are the Documents and List of Things to Bring for Camp. \n Download Links for Your Documents are Shared Below : \nPayment Receipt - {payment_receipt_url}\n Medical Certificate - {student_data['medicalCertificate']} \nEntrance Card - {ec} \nVisiting Card - {student_data['visiting_card']} \nAdmission Form - {student_data['admission_form']}\n\nTeam MCF Camp"

                    fns = []
                    payment_receipt = payment_receipt_url.replace(files_url,files_base_dir)
                    mcert = student_data['medicalCertificate'].replace(files_url,files_base_dir)
                    ecfn = ec.replace(files_url,files_base_dir)
                    vc = student_data['visiting_card'].replace(files_url,files_base_dir)
                    af = student_data['admission_form'].replace(files_url,files_base_dir)
                    fns = [payment_receipt, mcert, ecfn, vc, af]
            
                    send_email_attachments(msg=msg, sub="Payment Receipt and Other Documents", mailToSend=student_data['email'], files=fns)
                    send_wp(msg,student_data['wp_no'],file_paths=fns)

                    send_wp(msg,"9604084000")
                    send_email(msg, f"Payment Complete of Student - {student_data['first_name']} {student_data['last_name']}", "infomcfcamp@gmail.com")


                all_payments.insert_one({
                "payment_id": payment_id,
                "payment_option": data['payment_option'],
                "payment_amount": data['payment_amount'],
                "payment_date":data["payment_date"],
                "transaction_id":data["transaction_id"],
                "payment_mode":data["payment_mode"],
                "sid":data['sid'],
                "receipt_no":receipt_no,
                "receipt_url":f"{files_base_url}{student_data['sid']}_fee_receipt_{data['payment_option']}.pdf"
            })

            elif "1installment" in data['payment_option']:
                final_data['1_INST_AMT'] = data['payment_amount']
                final_data['1_INST_DT'] = data["payment_date"]
                for key, value in final_data.items():
                        find_and_replace_tables_fee_receipt(doc.tables, f'{{MERGEFIELD {key}}}', str(value))
                doc.save(str(file_dir)+str(f"{student_data['sid']}_fee_receipt_{data['payment_option']}.docx"))

                convert_to_pdf(str(str(file_dir)+str(f"{student_data['sid']}_fee_receipt_{data['payment_option']}.docx")), str(str(file_dir)+str(f"{student_data['sid']}_fee_receipt_{data['payment_option']}.pdf")))


                if float(total_paid) >= float(student_data['total_amount_payable']):
                    camp_id = student_data['camp_id']
                    camp_db = db["camps_db"]
                    camp_data = camp_db.find_one({"camp_id":camp_id}, {"_id":0})
                    batch_id = student_data['batch_id']
                    batch_db = db["batches_db"]
                    batch_data = batch_db.find_one({"batch_id":batch_id}, {"_id":0})
                    payment_db = db["all_payments"]
                    payment_data = payment_db.find({"sid":data['sid']}, {"_id":0})
                    receipt_nos = ""
                    payment_data = list(payment_data)
                    print(len(payment_data))
                    for receipt in payment_data:
                        receipt_nos = str(receipt_nos + str(str(receipt['receipt_no'])+ " , "))
                        print(receipt_nos)
                    receipt_nos = str(receipt_nos + receipt_no)

                    student_data_1 = {
                            'CADET_NAME': str(student_data["first_name"].upper()+" "+student_data["last_name"].upper()),
                            'REGNO': data['sid'],
                            'RANK': 'CDT',
                            'C_NAME': camp_data['camp_name'],
                            'C_BATCH': batch_data['batch_name'],
                            'C_DAYS': batch_data["duration"],
                            'COMP_N': student_data['company'],
                            'C_DATE': batch_data['start_date'],
                            'PICKPT': student_data["pick_up_point"],
                            'PICK_TIME': '',
                            'EMP_NAME':  student_data["employee_who_reached_out_to_you"],
                            'GAR_NAME': student_data["middle_name"],
                            'ADDRESS': student_data["address"],
                            'CITY': student_data["pick_up_city"],
                            'DISTRICT': student_data["district"],
                            'STATE': student_data['state'],
                            'PINCODE': student_data["pincode"],
                            'EMAIL': student_data["email"],
                            'C_NUM': str(student_data["phn"]),
                            'WP_NUM': student_data.get("wp_no", ""),
                            'FATHER_NUM': '',
                            'MOTHER_NUM': '',
                            'DOB': student_data["dob"],
                            'BLOOD_GROUP':  student_data.get("blood_group", ""),
                            'STD': student_data.get("standard", ""),
                            'SCHOOL': student_data.get("school_name", ""),
                            'FEE_PAID': student_data['total_amount_paid'],
                            'BALANCE': float(student_data['total_amount_payable']) - float(student_data['total_amount_paid']),
                            'RECEIPT_NUM': receipt_nos,
                            'DATE': '',
                            'TIME': ''
                        }

                    doc = Document('mcf_entrance_card.docx')

                    # Replace text fields in paragraphs
                    find_and_replace_paragraphs_entrance_card(doc.paragraphs, '{MERGEFIELD CADET_NAME}', student_data_1['CADET_NAME'])
                    find_and_replace_paragraphs_entrance_card(doc.paragraphs, '{MERGEFIELD DATE}', student_data_1['DATE'])
                    find_and_replace_paragraphs_entrance_card(doc.paragraphs, '{MERGEFIELD TIME}', student_data_1['TIME'])

                    for key, value in student_data_1.items():
                            find_and_replace_tables_entrance_card(doc.tables, f'{{MERGEFIELD {key}}}', str(value))

                    try:
                        print("replacing image")
                        table = doc.tables[0]  # Assuming the first table
                        cell = table.cell(0, 3)  # Assuming the first cell in the third column

                        # Clear the content of the cell by removing its paragraphs
                        for paragraph in cell.paragraphs:
                            paragraph.clear()

                        # Add a new paragraph and insert the image
                        paragraph = cell.add_paragraph()
                        run = paragraph.add_run()
                        cadet_photo_url = student_data["cadetPhoto"]
                        cadet_photo_path = cadet_photo_url.replace(files_url,files_base_dir)
                        run.add_picture(cadet_photo_path, width=Inches(0.9))
                    except Exception as e:
                        print("Error : ",str(e))

                    doc.save(str(str(file_dir)+f"{data['sid']}_entrance_card.docx"))

                    convert_to_pdf(str(str(file_dir)+f"{data['sid']}_entrance_card.docx"), str(str(file_dir)+f"{data['sid']}_entrance_card.pdf"))

                    ec = f"{files_base_url}{data['sid']}_entrance_card.pdf"

                    entrance_card = {
                        "entrence_card" : ec
                    }
                    

                    students_db.update_one({"sid": data['sid']}, {"$set": entrance_card})

                    payment_receipt_url = f"{files_base_url}{student_data['sid']}_fee_receipt_{data['payment_option']}.pdf"

                    msg = f"Hello,\n {student_data['first_name']} {student_data['last_name']} your admission for camp - {camp_data['camp_name']} and Batch - {batch_data['batch_name']} for Duration of {batch_data['duration']} is Succesfull. \n Following are the Documents and List of Things to Bring for Camp. \n Download Links for Your Documents are Shared Below : \nPayment Receipt - {payment_receipt_url}\n Medical Certificate - {student_data['medicalCertificate']} \nEntrance Card - {ec} \nVisiting Card - {student_data['visiting_card']} \nAdmission Form - {student_data['admission_form']}\n\nTeam MCF Camp"

                    fns = []
                    payment_receipt = payment_receipt_url.replace(files_url,files_base_dir)
                    mcert = student_data['medicalCertificate'].replace(files_url,files_base_dir)
                    ecfn = ec.replace(files_url,files_base_dir)
                    vc = student_data['visiting_card'].replace(files_url,files_base_dir)
                    af = student_data['admission_form'].replace(files_url,files_base_dir)
                    fns = [payment_receipt, mcert, ecfn, vc, af]
            
                    send_email_attachments(msg=msg, sub="Payment Receipt and Other Documents", mailToSend=student_data['email'], files=fns)
                    send_wp(msg,student_data['wp_no'],file_paths=fns)

                    send_wp(msg,"9604084000")
                    send_email(msg, f"Payment Complete of Student - {student_data['first_name']} {student_data['last_name']}", "infomcfcamp@gmail.com")


                all_payments.insert_one({
                "payment_id": payment_id,
                "payment_option": data['payment_option'],
                "payment_amount": data['payment_amount'],
                "payment_date":data["payment_date"],
                "transaction_id":data["transaction_id"],
                "payment_mode":data["payment_mode"],
                "sid":data['sid'],
                "receipt_no":receipt_no,
                "receipt_url":f"{files_base_url}{student_data['sid']}_fee_receipt_{data['payment_option']}.pdf"
            })

            elif '2installment' in data['payment_option']:
                final_data['2_INST_AMT'] = data['payment_amount']
                final_data['2_INST_DT'] = data["payment_date"]
                for key, value in final_data.items():
                        find_and_replace_tables_fee_receipt(doc.tables, f'{{MERGEFIELD {key}}}', str(value))
                doc.save(str(file_dir)+str(f"{student_data['sid']}_fee_receipt_{data['payment_option']}.docx"))

                convert_to_pdf(str(str(file_dir)+str(f"{student_data['sid']}_fee_receipt_{data['payment_option']}.docx")), str(str(file_dir)+str(f"{student_data['sid']}_fee_receipt_{data['payment_option']}.pdf")))


                if float(total_paid) >= float(student_data['total_amount_payable']):
                    camp_id = student_data['camp_id']
                    camp_db = db["camps_db"]
                    camp_data = camp_db.find_one({"camp_id":camp_id}, {"_id":0})
                    batch_id = student_data['batch_id']
                    batch_db = db["batches_db"]
                    batch_data = batch_db.find_one({"batch_id":batch_id}, {"_id":0})
                    payment_db = db["all_payments"]
                    payment_data = payment_db.find({"sid":data['sid']}, {"_id":0})
                    receipt_nos = ""
                    payment_data = list(payment_data)
                    print(len(payment_data))
                    for receipt in payment_data:
                        receipt_nos = str(receipt_nos + str(str(receipt['receipt_no'])+ " , "))
                        print(receipt_nos)
                    receipt_nos = str(receipt_nos + receipt_no)

                    student_data_1 = {
                            'CADET_NAME': str(student_data["first_name"].upper()+" "+student_data["last_name"].upper()),
                            'REGNO': data['sid'],
                            'RANK': 'CDT',
                            'C_NAME': camp_data['camp_name'],
                            'C_BATCH': batch_data['batch_name'],
                            'C_DAYS': batch_data["duration"],
                            'COMP_N': student_data['company'],
                            'C_DATE': batch_data['start_date'],
                            'PICKPT': student_data["pick_up_point"],
                            'PICK_TIME': '',
                            'EMP_NAME':  student_data["employee_who_reached_out_to_you"],
                            'GAR_NAME': student_data["middle_name"],
                            'ADDRESS': student_data["address"],
                            'CITY': student_data["pick_up_city"],
                            'DISTRICT': student_data["district"],
                            'STATE': student_data['state'],
                            'PINCODE': student_data["pincode"],
                            'EMAIL': student_data["email"],
                            'C_NUM': str(student_data["phn"]),
                            'WP_NUM': student_data.get("wp_no", ""),
                            'FATHER_NUM': '',
                            'MOTHER_NUM': '',
                            'DOB': student_data["dob"],
                            'BLOOD_GROUP':  student_data.get("blood_group", ""),
                            'STD': student_data.get("standard", ""),
                            'SCHOOL': student_data.get("school_name", ""),
                            'FEE_PAID': student_data['total_amount_paid'],
                            'BALANCE': float(student_data['total_amount_payable']) - float(student_data['total_amount_paid']),
                            'RECEIPT_NUM': receipt_nos,
                            'DATE': '',
                            'TIME': ''
                        }

                    doc = Document('mcf_entrance_card.docx')

                    # Replace text fields in paragraphs
                    find_and_replace_paragraphs_entrance_card(doc.paragraphs, '{MERGEFIELD CADET_NAME}', student_data_1['CADET_NAME'])
                    find_and_replace_paragraphs_entrance_card(doc.paragraphs, '{MERGEFIELD DATE}', student_data_1['DATE'])
                    find_and_replace_paragraphs_entrance_card(doc.paragraphs, '{MERGEFIELD TIME}', student_data_1['TIME'])

                    for key, value in student_data_1.items():
                            find_and_replace_tables_entrance_card(doc.tables, f'{{MERGEFIELD {key}}}', str(value))

                    try:
                        print("replacing image")
                        table = doc.tables[0]  # Assuming the first table
                        cell = table.cell(0, 3)  # Assuming the first cell in the third column

                        # Clear the content of the cell by removing its paragraphs
                        for paragraph in cell.paragraphs:
                            paragraph.clear()

                        # Add a new paragraph and insert the image
                        paragraph = cell.add_paragraph()
                        run = paragraph.add_run()
                        cadet_photo_url = student_data["cadetPhoto"]
                        cadet_photo_path = cadet_photo_url.replace(files_url,files_base_dir)
                        run.add_picture(cadet_photo_path, width=Inches(0.9))
                    except Exception as e:
                        print("Error : ",str(e))

                    doc.save(str(str(file_dir)+f"{data['sid']}_entrance_card.docx"))

                    convert_to_pdf(str(str(file_dir)+f"{data['sid']}_entrance_card.docx"), str(str(file_dir)+f"{data['sid']}_entrance_card.pdf"))

                    ec = f"{files_base_url}{data['sid']}_entrance_card.pdf"

                    entrance_card = {
                        "entrence_card" : ec
                    }
                    
                    payment_receipt_url = f"{files_base_url}{student_data['sid']}_fee_receipt_{data['payment_option']}.pdf"

                    students_db.update_one({"sid": data['sid']}, {"$set": entrance_card})
                    msg = f"Hello,\n {student_data['first_name']} {student_data['last_name']} your admission for camp - {camp_data['camp_name']} and Batch - {batch_data['batch_name']} for Duration of {batch_data['duration']} is Succesfull. \n Following are the Documents and List of Things to Bring for Camp. \n Download Links for Your Documents are Shared Below : \nPayment Receipt - {payment_receipt_url}\n Medical Certificate - {student_data['medicalCertificate']} \nEntrance Card - {ec} \nVisiting Card - {student_data['visiting_card']} \nAdmission Form - {student_data['admission_form']}\n\nTeam MCF Camp"

                    fns = []
                    payment_receipt = payment_receipt_url.replace(files_url,files_base_dir)
                    mcert = student_data['medicalCertificate'].replace(files_url,files_base_dir)
                    ecfn = ec.replace(files_url,files_base_dir)
                    vc = student_data['visiting_card'].replace(files_url,files_base_dir)
                    af = student_data['admission_form'].replace(files_url,files_base_dir)
                    fns = [payment_receipt, mcert, ecfn, vc, af]
            
                    send_email_attachments(msg=msg, sub="Payment Receipt and Other Documents", mailToSend=student_data['email'], files=fns)
                    send_wp(msg,student_data['wp_no'],file_paths=fns)

                    send_wp(msg,"9604084000")
                    send_email(msg, f"Payment Complete of Student - {student_data['first_name']} {student_data['last_name']}", "infomcfcamp@gmail.com")


                all_payments.insert_one({
                "payment_id": payment_id,
                "payment_option": data['payment_option'],
                "payment_amount": data['payment_amount'],
                "payment_date":data["payment_date"],
                "transaction_id":data["transaction_id"],
                "payment_mode":data["payment_mode"],
                "sid":data['sid'],
                "receipt_no":receipt_no,
                "receipt_url":f"{files_base_url}{student_data['sid']}_fee_receipt_{data['payment_option']}.pdf"
            })
                
            else:
                return {"error":"Please Specify Payment Option"}

            return jsonify({"message": f"Payment added successfully.", "payment_id": payment_id})
        else:
            return {"error":"Invalid Batch Duration"}

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    

@app.route("/getStudentPayment", methods=["GET"])
def getStudentPayment():
    try:
        # collection = db["students_db"]
        sid = request.args.get('sid')
        students_db = db["all_payments"]
        payment_data = students_db.find({"sid":sid}, {"_id":0})
        payment_data = list(payment_data)

        return jsonify({'success': True, "payments":payment_data}), 200

    except Exception as e:
        return jsonify({'success': False, 'msg': 'Something Went Wrong.', 'reason': str(e)}), 500
    


@app.route('/deletePayment', methods=['DELETE'])
def delete_payment():
    try:
        payment_db = db["all_payments"]
        payment_id = request.args.get('payment_id')
        sid = request.args.get('sid')
        payment_amount = request.args.get('payment_amount')

        if not payment_id:
            return jsonify({"error": "Missing 'payment_id' parameter in the request."}), 400  # Bad Request
        payment = payment_db.find_one({"payment_id": payment_id})

        if not payment:
            return jsonify({"error": f"No student found with payment_id: {payment_id}"}), 404  # Not Found
        payment_db.delete_one({"payment_id": payment_id})

        students_db = db['students_db']
        student_data = students_db.find_one({"sid":sid}, {"_id":0})
        total_paid = student_data['total_amount_paid']
        total_paid = float(total_paid) - float(payment_amount)
        student_data['total_amount_paid'] = total_paid
        students_db.update_one({"sid": sid}, {"$set": student_data})

        return jsonify({"message": f"Student with payment_id {payment_id} deleted successfully"})

    except Exception as e:
        return jsonify({"error": str(e)}), 500  # Internal Server Error
    


@app.route("/sendAllDocuments", methods=["GET"])
def sendAllDocuments():
    try:
        # collection = db["students_db"]
        sid = request.args.get('sid')
        students_db = db["students_db"]
        student_data = students_db.find_one({"sid":sid}, {"_id":0})

        mailToSend = student_data['email']

        payment_db = db["all_payments"]
        payment_data = payment_db.find({"sid":sid},{"_id":0})
        payment_data = list(payment_data)

        payment_links = ''

        mcert = student_data['medicalCertificate'].replace(files_url,files_base_dir)
        ecfn = student_data['entrence_card'].replace(files_url,files_base_dir)
        vc = student_data['visiting_card'].replace(files_url,files_base_dir)
        af = student_data['admission_form'].replace(files_url,files_base_dir)
        fns = [mcert, ecfn, vc, af]

        c=1
        for data in payment_data:
            payment_links += f"{c}) - {data['receipt_url']}\n"
            fns.append(data['receipt_url'].replace(files_url,files_base_dir))
            c+=1


        msg = f"Hello,\n Download Links for Your Documents are Shared Below : \nPayment Receipt - {payment_links}\n\n Medical Certificate - {student_data['medicalCertificate']} \n\nEntrance Card - {student_data['entrence_card']} \n\nVisiting Card - {student_data['visiting_card']} \n\nAdmission Form - {student_data['admission_form']}\n\nTeam MCF Summer Camp"
            
        send_email_attachments(msg=msg, sub="All Documents Download Links", mailToSend=mailToSend, files=fns)
        send_wp(msg,student_data['wp_no'],file_paths=fns)

        return jsonify({'success': True, 'msg': 'Mail Send'}), 200

    except Exception as e:
        return jsonify({'success': False, 'msg': 'Something Went Wrong.', 'reason': str(e)}), 500
    
def process_data(body):
    students_db = db["students_db"]
    payment_db = db["all_payments"]
    for dt in body:
        student_data = students_db.find_one({"sid":dt['sid']}, {"_id":0})
        mailToSend = student_data['email']
        payment_data = payment_db.find({"sid":dt['sid']},{"_id":0})
        payment_data = list(payment_data)
        payment_links = ''
        mcert = student_data['medicalCertificate'].replace(files_url,files_base_dir)
        ecfn = student_data['entrence_card'].replace(files_url,files_base_dir)
        vc = student_data['visiting_card'].replace(files_url,files_base_dir)
        af = student_data['admission_form'].replace(files_url,files_base_dir)
        fns = [mcert, ecfn, vc, af]
        c=1
        for data in payment_data:
            payment_links += f"{c}) - {data['receipt_url']}\n"
            fns.append(data['receipt_url'].replace(files_url,files_base_dir))
            c+=1
        msg = f"Hello,\n Download Links for Your Documents are Shared Below : \nPayment Receipt - {payment_links}\n\n Medical Certificate - {student_data['medicalCertificate']} \n\nEntrance Card - {student_data['entrence_card']} \n\nVisiting Card - {student_data['visiting_card']} \n\nAdmission Form - {student_data['admission_form']}\n\nTeam MCF Camp"

        send_email_attachments(msg=msg, sub="All Documents Download Links", mailToSend=mailToSend, files=fns)
        send_wp(msg,student_data['wp_no'],file_paths=fns)



def send_certs_bulk(body):
    students_db = db["students_db"]
    for dt in body:
        student_data = students_db.find_one({"sid":dt['sid']}, {"_id":0})
        mailToSend = student_data['email']
        if student_data['completion_cert']:
            cc = student_data['completion_cert'].replace(files_url,files_base_dir)
            ncc = cc.replace(files_url,files_base_dir)
            msg = f"Hello,\n Download Links for Your Documents are Shared Below : \nCertificate - {cc}\n\nTeam MCF Camp"

            send_email_attachments(msg=msg, sub="Camp Completion Certificate from MCF Camp", mailToSend=mailToSend, files=[ncc])
            send_wp(msg,student_data['wp_no'],file_paths=[ncc])
        else:
            print("Generate Certificate First")


@app.route("/sendAllStudentsDocs", methods=["POST"])
def sendAllStudentsDocs():
    try:
        data = request.json
        body = data['body']
        thread = threading.Thread(target=process_data, args=(body,))
        thread.start()
        return jsonify({'success': True, 'msg': 'Process Started'}), 200

    except Exception as e:
        return jsonify({'success': False, 'msg': 'Something Went Wrong.', 'reason': str(e)}), 500
    

@app.route("/sendAllStudentsCert", methods=["POST"])
def sendAllStudentsCert():
    try:
        data = request.json
        body = data['body']
        thread = threading.Thread(target=send_certs_bulk, args=(body,))
        thread.start()
        return jsonify({'success': True, 'msg': 'Process Started'}), 200

    except Exception as e:
        return jsonify({'success': False, 'msg': 'Something Went Wrong.', 'reason': str(e)}), 500
    


@app.route("/resetDiscount", methods=["GET"])
def resetDiscount():
    try:
        # collection = db["students_db"]
        sid = request.args.get('sid')
        camp_id = request.args.get('camp_id')
        students_db = db["students_db"]
        camps_db = db["camps_db"]
        student_data = students_db.find({"sid":sid}, {"_id":0})
        camp_data = camps_db.find_one({"camp_id":camp_id}, {"_id":0})

        original_amt = camp_data['camp_fee']
        reset_data = {
            "total_amount_payable":original_amt,
            "discount_amount":0
        }

        students_db.update_one({"sid": sid}, {"$set": reset_data})


        return jsonify({'success': True, "msg":'Reset Successful'}), 200

    except Exception as e:
        return jsonify({'success': False, 'msg': 'Something Went Wrong.', 'reason': str(e)}), 500

def create_zip(directory, selected_files, zip_filename):
    if os.path.exists(zip_filename):
        os.remove(zip_filename)
    with zipfile.ZipFile(zip_filename, 'w', allowZip64=True) as zipf:
        for filename in selected_files:
            filepath = os.path.join(directory, filename)
            if os.path.exists(filepath):
                zipf.write(filepath, arcname=filename)
            else:
                print(f"File '{filename}' not found in directory '{directory}'.")

@app.route("/bulkDownloadAdmissionCard", methods=["POST"])
def bulkDownloadAdmissionCard():
    try:
        data = request.json
        body = data['body']
        filter = data['filter']
        fns = []
        for dt in body:
            admission_link = dt["admission_form"]
            fn = admission_link.replace(files_base_url,"")
            fns.append(fn)

        # Example usage
        directory = file_dir
        zip_filename = f"{file_dir}All_{filter['camp_id']}_{filter['batch_id']}_{filter['status']}_Admission_Cards.zip"

        create_zip(directory, fns, zip_filename)

        zip_url = f"{files_base_url}All_{filter['camp_id']}_{filter['batch_id']}_{filter['status']}_Admission_Cards.zip"

        rint = str(uuid.uuid4().hex)

        # return send_file(zip_filename, as_attachment=True)
        data = {
            'success': True,
            'msg': zip_url + f"?nocache={rint}",
            'filename': f"All_{filter['camp_id']}_{filter['batch_id']}_{filter['status']}_Admission_Cards.zip"
        }

        # Create a JSON response
        response = make_response(jsonify(data), 200)

        # Set Cache-Control header to no-store
        response.headers['Cache-Control'] = 'no-store'

        return response

    except Exception as e:
        return jsonify({'success': False, 'msg': 'Something Went Wrong.', 'reason': str(e)}), 500
    


@app.route("/syncStudent", methods=["GET"])
def sync_Student():
    try:
        sid = request.args.get('sid')
        result = sync_data(sid)
        if result == 0:
            return jsonify({'success': True, "msg":'Sync Successful'}), 200
        else:
            return jsonify({"success":False,"msg":"Not Sync"}), 400
    except Exception as e:
        return jsonify({'success': False, 'msg': 'Something Went Wrong.', 'reason': str(e)}), 500
    

@app.route('/submitFeedback', methods=['POST'])
def submit_feedback():
    try:
        json_data = request.get_json()
        feedback_db = db['feedback_db']
        students_db = db["students_db"]
        feedback=feedback_db.find_one({"sid":json_data['sid']})
        student = students_db.find_one({"sid":json_data['sid']})
        if student and not feedback:
            if len(json_data) >= 17:

                doc = Document('feedback_files/SUMMER_CAMP_FEEDBACK.docx')

                # Sample student_data
                student_data1 = {
                    'CADET_NAME': str(json_data['name']),
                    'CAMP_NAME': json_data['camp_name'],
                    'DAYS' : json_data['duration'],
                    'BATCH': json_data['batch_name'],
                    'REG_NO': json_data['sid'],
                    'TRANING' : json_data['training'],
                    'CAMPUS' : json_data['campus'],
                    'FOOD': json_data['food'],
                    'DORMITORY':json_data['dormitory'],
                    'TRAVELLING':json_data['traveling'],
                    'CAMP_EXPERIENCE':json_data['campExperience'],
                    'SUGGESTIONS' : json_data['suggestions'],
                    'PARENT_NAME': json_data['middle_name'],

                }

                for key, value in student_data1.items():
                        find_and_replace_tables_fdb(doc.tables, f'{{MERGEFIELD {key}}}', str(value))

                yes_star = "feedback_files/fill_star.png"
                no_star = "feedback_files/hollow_star.png"
                no_of_stars = int(json_data['camp_rating'])
                i_index = 1
                for i in range (no_of_stars):
                    replace_image_in_cell_fdb(doc, table_index=1, row_index=0, column_index=i_index, image_path=yes_star,w=0.23)
                    i_index = i_index + 1
                no_star_num = 5 - no_of_stars
                ns_index = i_index
                for i in range (no_star_num) :
                    replace_image_in_cell_fdb(doc, table_index=1, row_index=0, column_index=ns_index, image_path=no_star,w=0.23)
                    ns_index = ns_index + 1


                image_path_guardian = json_data['parent_sign'].replace(files_url,files_base_dir)
                replace_image_in_cell_fdb(doc, table_index=2, row_index=1, column_index=3, image_path=image_path_guardian,w=1.8)

                docx_path = f"{file_dir}{student_data1['REG_NO']}_feedback.docx"
                doc.save(docx_path)
                output_path = f"{file_dir}{student_data1['REG_NO']}_feedback.pdf"
                convert_to_pdf(docx_path,output_path)

                feedback_url = f"{files_base_url}{student_data1['REG_NO']}_feedback.pdf"

                json_data['feedback_form'] = feedback_url

                students_db.update_one({"sid": json_data['sid']}, {"$set": {"feedback_form":feedback_url}})

                feedback_db.insert_one(json_data)

                response = {
                'success': True,
                'message': 'Feedback stored successfully',
                }


                return jsonify(response), 200
            else:
                response = {
                'success': False,
                'message': 'Some Fields are not filled...',
                }
                return jsonify(response), 400
        else:
            response = {
            'success': False,
            'message': 'Feedback from this MRN already Exist',
        }
            return jsonify(response), 400
    except Exception as e:
        error_response = {
            'success': False,
            'message': f'Error storing JSON data: {str(e)}'
        }
        return jsonify(error_response), 500
    


@app.route("/generateCampCertificate", methods=["GET"])
def generate_camp_certificates():
    try:
        sid = request.args.get('sid')
        result = generate_certificate_cert(sid)
        if result == 0:
            return jsonify({'success': True, "msg":'Generate Successful'}), 200
        else:
            return jsonify({"success":False,"msg":"Generate Failed"}), 400
    except Exception as e:
        return jsonify({'success': False, 'msg': 'Something Went Wrong.', 'reason': str(e)}), 500



def bulk_generate_cert(data):
    for dt in data:
        generate_certificate_cert(dt['sid'])


@app.route("/bulkGenerateCampCertificate", methods=["POST"])
def bulkGenerateCampCertificate():
    try:
        data = request.json
        body = data['body']
        thread = threading.Thread(target=bulk_generate_cert, args=(body,))
        thread.start()
        return jsonify({'success': True, 'msg': 'Process Started'}), 200

    except Exception as e:
        return jsonify({'success': False, 'msg': 'Something Went Wrong.', 'reason': str(e)}), 500
    

@app.route("/sendCampCertificate", methods=["GET"])
def sendCampCertificate():
    try:
        sid = request.args.get('sid')
        send_certs_bulk([{"sid":sid}])
        return jsonify({'success': True, 'msg': 'Send Succefully'}), 200

    except Exception as e:
        return jsonify({'success': False, 'msg': 'Something Went Wrong.', 'reason': str(e)}), 500
    

@app.route("/getAllFeedbacks", methods=["GET"])
def getAllFeedbacks():
    try:
        feedback_db = db["feedback_db"]
        feedbacks = feedback_db.find({},{'_id':0})
        return jsonify({'success': True, "feedbacks":list(feedbacks)}), 200
    except Exception as e:
        return jsonify({'success': False, 'msg': 'Something Went Wrong.', 'reason': str(e)}), 500
    


#-----------------------------------------------------------------------------------------------------------
    
#---------------------------------- Payment Section Start --------------------------------------------

import requests
import json
import hashlib
import uuid

def generate_hash(data, salt):
    # key|merchant_txn|name|email|phone|amount|udf1|udf2|udf3|udf4|udf5|message|salt
    concat_string = f"{data[0]}|{data[1]}|{data[2]}|{data[3]}|{data[4]}|{data[5]}|{data[6]}|{data[7]}|{data[8]}|{data[9]}|{data[10]}|{data[11]}|{salt}"
    hashed = hashlib.sha512(concat_string.encode()).hexdigest()
    print(hashed)
    return hashed

def initiate_payment(name, email, phn, camp_name, amt, sid):
        url = "https://dashboard.easebuzz.in/easycollect/v1/create"
        msg = camp_name
        transaction_id = uuid.uuid4().hex
        key = "1YUG4UBN1Q"

        # Define the data in the specified sequence for hashing
        data_sequence = [
            key,
            transaction_id,
            name,
            email,
            phn,
            amt,
            sid,
            "",
            "",
            "",
            "",
            msg,
        ]
        salt = "KPRYL60DC1"

        # Generate hash using the data sequence and salt
        hash_value = generate_hash(data_sequence, salt)
        # key|merchant_txn|name|email|phone|amount|udf1|udf2|udf3|udf4|udf5|message|salt
        payload = {
            "key": key,
            "merchant_txn": transaction_id,
            "name": name,
            "email": email,
            "phone": phn,
            "amount": amt,
            "udf1": sid,
            "udf2": " ",
            "udf3": " ",
            "udf4": " ",
            "udf5": " ",
            "message": msg,
            "accept_partial_payment": False,
            "hash": hash_value  # Include the generated hash in the payload
        }

        headers = {
            "Accept": "application/json",
            "Content-Type": "application/json"
        }

        try:
            response = requests.post(url, headers=headers, json=payload)
            response_data = response.json()
            return jsonify({"success":True,"payment_url":response_data['data']['payment_url']}),200
        except Exception as e:
            print("Error:", e)
            return jsonify({"success":False, "error":str(e)}),400
    
#---------------------------------- Payment Section End ----------------------------------------------

#-----------------------------------------------------------------------------------------------------------


    


if __name__ == '__main__':
    app.run(host="0.0.0.0")





