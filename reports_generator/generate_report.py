from docx import Document
from docx.shared import Pt ,Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import subprocess

def set_paragraph_font(paragraph, font_name, font_size,):
    for run in paragraph.runs:
        font = run.font
        font.name = font_name
        font.size = Pt(font_size)
        
def find_and_replace_paragraphs(paragraphs, field, replacement, specific_font=None):
    for paragraph in paragraphs:
        if field in paragraph.text:
            paragraph.text = paragraph.text.replace(field, replacement)
            if specific_font is not None:
                set_paragraph_font(paragraph, *specific_font)

def find_and_replace_tables(tables, field, replacement,specific_font):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    find_and_replace_paragraphs([paragraph], field, replacement, specific_font=specific_font)

def replace_image_in_cell(doc, table_index, row_index, column_index, image_path,w,h):
    table = doc.tables[table_index]
    cell = table.cell(row_index, column_index)
    for paragraph in cell.paragraphs:
        paragraph.clear()
    paragraph = cell.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(image_path,width=Inches(w), height=Inches(h))
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def convert_to_pdf(docx_file, pdf_file):
    try:
        subprocess.run(['unoconv', '--output', pdf_file, '--format', 'pdf', docx_file], check=True)
        print(f"Conversion successful: {docx_file} -> {pdf_file}")
    except subprocess.CalledProcessError as e:
        print(f"Error during conversion: {e}")


def generate_report(sid, data, image_path_photo, days, camp_date, batch_name, docx_path, output_path): 
    Report_Data = {
        #details
        "REG_NO" : sid,
        "CADET_NAME" : data["details"]["name"],
        "ADDRESS" : data["details"]["address"],
        "CONTACT_NO" : data["details"]["phone"],
        "EMAIL_ID" : data["details"]["email"],
        "CAMP_NAME" : data["details"]["camp_name"],
        "DAYS" : days,
        "CAMP_DATE" : camp_date,
        "BATCH" : batch_name,
        "CQY_NAME" :data["details"]["cqy_name"],
        "INCHARGE_NAME" : data["details"]["incharge_name"],

        #dates
        "CHECK_IN_DATE" : data["dates"]["checkin_Date"],
        "PICK_UP_DATE" : data["dates"]["pickup_Date"],
        "CHECK_OUT_DATE" : data["dates"]["checkout_Date"],
        "DROP_DATE" : data["dates"]["drop_Date"],

        "LAST_CLOSE_DATE" : data["dates"]["last_closing_ceremony"],
        "LAST_CLOSE_TIME" : data["dates"]["last_closing_time"],
        "PARENT_P_A" : data["final_remarks"]["Parents presence"],
        "BEST_ACTIVITY" : data["final_remarks"]["best activity"],

        "REMARKS" : data["final_remarks"]["remarks"],

        }

    
    
    if 'SMTC' in Report_Data["CAMP_NAME"] or 'smtc' in Report_Data["CAMP_NAME"]:
        doc = Document('report_card_SMTC.docx')
        camp_report_data ={
            #skill_activities
            "ARCHERY_TTR" : data["activities"]["skill_activities"][0]["TIMES TO REPEAT"],
            "ARCHERY_TBI" : data["activities"]["skill_activities"][0]["TRAINED BY INS"],
            "LATHI_KATHI_TTR" : data["activities"]["skill_activities"][1]["TIMES TO REPEAT"],
            "LATHI_KATHI_TBI" : data["activities"]["skill_activities"][1]["TRAINED BY INS"],
            "RIFLE_SHOOTING_TTR" : data["activities"]["skill_activities"][2]["TIMES TO REPEAT"],
            "RIFLE_SHOOTING_TBI" : data["activities"]["skill_activities"][2]["TRAINED BY INS"],
            "MARTIAL_ARTS_TTR" : data["activities"]["skill_activities"][3]["TIMES TO REPEAT"],
            "MARTIAL_ARTS_TBI" : data["activities"]["skill_activities"][3]["TRAINED BY INS"],
            "HORSE_RIDING_TTR" : data["activities"]["skill_activities"][4]["TIMES TO REPEAT"],
            "HORSE_RIDING_TBI" : data["activities"]["skill_activities"][4]["TRAINED BY INS"],
            "PISTOL_SHOOTING_TTR" : data["activities"]["skill_activities"][5]["TIMES TO REPEAT"],
            "PISTOL_SHOOTING_TBI" : data["activities"]["skill_activities"][5]["TRAINED BY INS"],
            "MAP_READING_TTR" : data["activities"]["skill_activities"][6]["TIMES TO REPEAT"],
            "MAP_READING_TBI" : data["activities"]["skill_activities"][6]["TRAINED BY INS"],

            #physical_activities
            "TREKKING_TTR" : data["activities"]["physical_activities"][0]["TIMES TO REPEAT"],
            "TREKKING_TBI" : data["activities"]["physical_activities"][0]["TRAINED BY INS"],
            "AEROBICS_TTR" : data["activities"]["physical_activities"][1]["TIMES TO REPEAT"],
            "AEROBICS_TBI" : data["activities"]["physical_activities"][1]["TRAINED BY INS"],
            "PT_EXERCISE_TTR" : data["activities"]["physical_activities"][2]["TIMES TO REPEAT"],
            "PT_EXERCISE_TBI" : data["activities"]["physical_activities"][2]["TRAINED BY INS"],
            "COMMANDO_ACTIVITIES_TTR" : data["activities"]["physical_activities"][3]["TIMES TO REPEAT"],
            "COMMANDO_ACTIVITIES_TBI" : data["activities"]["physical_activities"][3]["TRAINED BY INS"],
            "FCBC_TTR" : data["activities"]["physical_activities"][4]["TIMES TO REPEAT"],
            "FCBC_TBI" : data["activities"]["physical_activities"][4]["TRAINED BY INS"],
            #water_activities
            "RAIN_DANCE_TTR" : data["activities"]["water_activities"][0]["TIMES TO REPEAT"],
            "RAIN_DANCE_TBI" : data["activities"]["water_activities"][0]["TRAINED BY INS"],
            "SWIMMING_TTR" : data["activities"]["water_activities"][1]["TIMES TO REPEAT"],
            "SWIMMING_TBI" : data["activities"]["water_activities"][1]["TRAINED BY INS"],
            "BOATING_TTR" : data["activities"]["water_activities"][2]["TIMES TO REPEAT"],
            "BOATING_TBI" : data["activities"]["water_activities"][2]["TRAINED BY INS"],
            #adventure_activities
            "ROCK_CLIMBING_TTR" : data["activities"]["adventure_activities"][0]["TIMES TO REPEAT"],
            "ROCK_CLIMBING_TBI" : data["activities"]["adventure_activities"][0]["TRAINED BY INS"],
            "ZIP_LINE_TTR" : data["activities"]["adventure_activities"][1]["TIMES TO REPEAT"],
            "ZIP_LINE_TBI" : data["activities"]["adventure_activities"][1]["TRAINED BY INS"],
            "RAPPELLING_TTR" : data["activities"]["adventure_activities"][2]["TIMES TO REPEAT"],
            "RAPPELLING_TBI" : data["activities"]["adventure_activities"][2]["TRAINED BY INS"],
            "ROPE_CLIMBING_TTR" : data["activities"]["adventure_activities"][3]["TIMES TO REPEAT"],
            "ROPE_CLIMBING_TBI" : data["activities"]["adventure_activities"][3]["TRAINED BY INS"],
            "WATER_RAPPELLING_TTR" : data["activities"]["adventure_activities"][4]["TIMES TO REPEAT"],
            "WATER_RAPPELLING_TBI" : data["activities"]["adventure_activities"][4]["TRAINED BY INS"],
            "PARAGLIDING_TTR" : data["activities"]["adventure_activities"][5]["TIMES TO REPEAT"],
            "PARAGLIDING_TBI" : data["activities"]["adventure_activities"][5]["TRAINED BY INS"],

            #mcf_rope_course_activities    
            "VERTICAL_ROPE_CLIMBING_TTR" : data["activities"]["mcf_rope_course_activities"][0]["TIMES TO REPEAT"],
            "VERTICAL_ROPE_CLIMBING_TBI" : data["activities"]["mcf_rope_course_activities"][0]["TRAINED BY INS"],
            "ROPE_BRIDGE_TTR" : data["activities"]["mcf_rope_course_activities"][1]["TIMES TO REPEAT"],
            "ROPE_BRIDGE_TBI" : data["activities"]["mcf_rope_course_activities"][1]["TRAINED BY INS"],
            "LADDER_WALKING_TTR" : data["activities"]["mcf_rope_course_activities"][2]["TIMES TO REPEAT"],
            "LADDER_WALKING_TBI" : data["activities"]["mcf_rope_course_activities"][2]["TRAINED BY INS"],
            "BARREL_CRAWLING_TTR" : data["activities"]["mcf_rope_course_activities"][3]["TIMES TO REPEAT"],
            "BARREL_CRAWLING_TBI" : data["activities"]["mcf_rope_course_activities"][3]["TRAINED BY INS"],
            "TARZAN_SWING_TTR" : data["activities"]["mcf_rope_course_activities"][4]["TIMES TO REPEAT"],
            "TARZAN_SWING_TBI" : data["activities"]["mcf_rope_course_activities"][4]["TRAINED BY INS"],
            "AUSTRALIAN_WALK_TTR" : data["activities"]["mcf_rope_course_activities"][5]["TIMES TO REPEAT"], 
            "AUSTRALIAN_WALK_TBI" : data["activities"]["mcf_rope_course_activities"][5]["TRAINED BY INS"], 
            "WALL_JUMP_TTR" : data["activities"]["mcf_rope_course_activities"][6]["TIMES TO REPEAT"],
            "WALL_JUMP_TBI" : data["activities"]["mcf_rope_course_activities"][6]["TRAINED BY INS"],
            "TAWA_WALK_TTR" : data["activities"]["mcf_rope_course_activities"][7]["TIMES TO REPEAT"],
            "TAWA_WALK_TBI" : data["activities"]["mcf_rope_course_activities"][7]["TRAINED BY INS"],
            "TYRE_WALK_TTR" : data["activities"]["mcf_rope_course_activities"][8]["TIMES TO REPEAT"],
            "TYRE_WALK_TBI" : data["activities"]["mcf_rope_course_activities"][8]["TRAINED BY INS"],
            "RING_SWING_TTR" : data["activities"]["mcf_rope_course_activities"][9]["TIMES TO REPEAT"],
            "RING_SWING_TBI" : data["activities"]["mcf_rope_course_activities"][9]["TRAINED BY INS"],
            "STRAIGHT_BALANCE_TTR" : data["activities"]["mcf_rope_course_activities"][10]["TIMES TO REPEAT"],
            "STRAIGHT_BALANCE_TBI" : data["activities"]["mcf_rope_course_activities"][10]["TRAINED BY INS"],
            "SINGLE_ROPE_WALK_TTR" : data["activities"]["mcf_rope_course_activities"][11]["TIMES TO REPEAT"],
            "SINGLE_ROPE_WALK_TBI" : data["activities"]["mcf_rope_course_activities"][11]["TRAINED BY INS"],
            "ZIGZAG_LADDER_WALK_TTR" : data["activities"]["mcf_rope_course_activities"][12]["TIMES TO REPEAT"],
            "ZIGZAG_LADDER_WALK_TBI" : data["activities"]["mcf_rope_course_activities"][12]["TRAINED BY INS"],
            "ONE_FEET_WALK_TTR" : data["activities"]["mcf_rope_course_activities"][13]["TIMES TO REPEAT"],
            "ONE_FEET_WALK_TBI" : data["activities"]["mcf_rope_course_activities"][13]["TRAINED BY INS"],
            
            #military_obstacle_activities
            "COMMANDO_NT_TTR" : data["activities"]["military_obstacle_activities"][0]["TIMES TO REPEAT"],
            "COMMANDO_NT_TBI" : data["activities"]["military_obstacle_activities"][0]["TRAINED BY INS"],
            "TYRE_CLIMBING_TTR" : data["activities"]["military_obstacle_activities"][1]["TIMES TO REPEAT"],
            "TYRE_CLIMBING_TBI" : data["activities"]["military_obstacle_activities"][1]["TRAINED BY INS"],
            "SPIDER_NET_TTR" : data["activities"]["military_obstacle_activities"][2]["TIMES TO REPEAT"],
            "SPIDER_NET_TBI" : data["activities"]["military_obstacle_activities"][2]["TRAINED BY INS"],
            "VERTICAL_NET_TTR" : data["activities"]["military_obstacle_activities"][3]["TIMES TO REPEAT"],
            "VERTICAL_NET_TBI" : data["activities"]["military_obstacle_activities"][3]["TRAINED BY INS"],
            "LADDER_CLIMBING_TTR" : data["activities"]["military_obstacle_activities"][4]["TIMES TO REPEAT"],
            "LADDER_CLIMBING_TBI" : data["activities"]["military_obstacle_activities"][4]["TRAINED BY INS"],
            
            #commando_training_activities
            "RIFLE_DRILL_TTR" : data["activities"]["commando_training_activities"][0]["TIMES TO REPEAT"],
            "RIFLE_DRILL_TBI" : data["activities"]["commando_training_activities"][0]["TRAINED BY INS"],
            "MARCH_PAST_TTR" : data["activities"]["commando_training_activities"][1]["TIMES TO REPEAT"],
            "MARCH_PAST_TBI" : data["activities"]["commando_training_activities"][1]["TRAINED BY INS"],
            "DHAVA_POSITION_TTR" : data["activities"]["commando_training_activities"][2]["TIMES TO REPEAT"],
            "DHAVA_POSITION_TBI" : data["activities"]["commando_training_activities"][2]["TRAINED BY INS"],
            "FIELD_CRAFT_TTR" : data["activities"]["commando_training_activities"][3]["TIMES TO REPEAT"],
            "FIELD_CRAFT_TBI" : data["activities"]["commando_training_activities"][3]["TRAINED BY INS"],
            "ATTACK_SKILL_TTR" : data["activities"]["commando_training_activities"][4]["TIMES TO REPEAT"],
            "ATTACK_SKILL_TBI" : data["activities"]["commando_training_activities"][4]["TRAINED BY INS"],
            "SELF_DEFENSE_TTR" : data["activities"]["commando_training_activities"][5]["TIMES TO REPEAT"],
            "SELF_DEFENSE_TBI" : data["activities"]["commando_training_activities"][5]["TRAINED BY INS"],

            #survival_training_activities
            "SNAKE_BITE_TTR" : data["activities"]["survival_training_activities"][0]["TIMES TO REPEAT"],
            "SNAKE_BITE_TBI" : data["activities"]["survival_training_activities"][0]["TRAINED BY INS"],
            "NATURAL_PLANT_INFO_TTR" : data["activities"]["survival_training_activities"][1]["TIMES TO REPEAT"],
            "NATURAL_PLANT_INFO_TBI" : data["activities"]["survival_training_activities"][1]["TRAINED BY INS"],
            "ANIMAL_INFO_TTR" : data["activities"]["survival_training_activities"][2]["TIMES TO REPEAT"],
            "ANIMAL_INFO_TBI" : data["activities"]["survival_training_activities"][2]["TRAINED BY INS"],
            "BIO_DI_INFO_TTR" : data["activities"]["survival_training_activities"][3]["TIMES TO REPEAT"],
            "BIO_DI_INFO_TBI" : data["activities"]["survival_training_activities"][3]["TRAINED BY INS"],
            "TRACK_ART_OBV_TTR" : data["activities"]["survival_training_activities"][4]["TIMES TO REPEAT"],
            "TRACK_ART_OBV_TBI" : data["activities"]["survival_training_activities"][4]["TRAINED BY INS"],
            "WILD_AWARE_TTR" : data["activities"]["survival_training_activities"][5]["TIMES TO REPEAT"],
            "WILD_AWARE_TBI" : data["activities"]["survival_training_activities"][5]["TRAINED BY INS"],


            "FUTURE_CAREER" : data["individual_remarks_form"]["future_career"],

            "TIME_MANAGEMENT" : data["individual_remarks_form"]["time_management"],
            "ACCOMMODATION" : data["individual_remarks_form"]["accommodation"],
            "FACILITIES" : data["individual_remarks_form"]["facilities"],
            "INS_TRAINING" : data["individual_remarks_form"]["ins_training"],
            "ATMOSPHERE" : data["individual_remarks_form"]["atmosphere"],
            "ACTIVITY" : data["individual_remarks_form"]["activity"],
            "UNIFORM" : data["individual_remarks_form"]["uniform"],
            "CERTIFICATION" : data["individual_remarks_form"]["certification"],
            "SUGGESTION" : data["individual_remarks_form"]["suggestion"],
            "BEST_ACHIEVEMENT" : data["individual_remarks_form"]["best_achievement"],
            "ACHIEVEMENT" : data["individual_remarks_form"]["achievement"],
            "PERSONALITY_DIMENSIONS" : data["individual_remarks_form"]["personality_dimensions"],

            "LEADER_P_S" : data["individual_remarks_table"]["LEADERSHIP POTENTIAL"]["SCORE"],
            "LEADER_P_I" : data["individual_remarks_table"]["LEADERSHIP POTENTIAL"]["INTERPRETATION"],
            "COMM_SKILLS_S" : data["individual_remarks_table"]["COMMUNICATION SKILLS"]["SCORE"],
            "COMM_SKILLS_I" : data["individual_remarks_table"]["COMMUNICATION SKILLS"]["INTERPRETATION"],
            "TEAM_COOR_S" : data["individual_remarks_table"]["TEAMWORK AND COOPERATION"]["SCORE"],
            "TEAM_COOR_I" : data["individual_remarks_table"]["TEAMWORK AND COOPERATION"]["INTERPRETATION"],
            "ADAPT_FLEX_S" : data["individual_remarks_table"]["ADAPTABILITY AND FLEXIBILITY"]["SCORE"],
            "ADAPT_FLEX_I" : data["individual_remarks_table"]["ADAPTABILITY AND FLEXIBILITY"]["INTERPRETATION"],
            "PRO_SOL_ABILITY_S" : data["individual_remarks_table"]["PROBLEM-SOLVING ABILITY"]["SCORE"],
            "PRO_SOL_ABILITY_I" : data["individual_remarks_table"]["PROBLEM-SOLVING ABILITY"]["INTERPRETATION"],

            "OVERALL_ASSESSMENT" : data["final_remarks"]["overall assessment"],

            
            "LEADER_DEV" : data["recommendations"]["leadership development"],
            "COMM_ENHANCE" : data["recommendations"]["communication enhancement"],
            "TEAM_ENHANCE" : data["recommendations"]["teamwork enhancement"],
            "ADAPT_TRAIN" : data["recommendations"]["adaptibility training"],
            "PRO_SOL_SKILLS_DEV" : data["recommendations"]["problem solving"],


            "CHECKED_BY_NAME" : data["final_remarks"]["checked by name"],
            "RANK" : data["final_remarks"]["rank"],

            }
    elif 'CTC' in Report_Data["CAMP_NAME"] or 'ctc' in Report_Data["CAMP_NAME"] :
        doc = Document('report_card_CTC.docx')
        camp_report_data ={
            "ARCHERY_TTR" : data["activities"]["skill_activities"][0]["TIMES TO REPEAT"] ,
            "ARCHERY_TBI" : data["activities"]["skill_activities"][0]["TRAINED BY INS"] ,
            "LATHI_KATHI_TTR" : data["activities"]["skill_activities"][1]["TIMES TO REPEAT"] ,
            "LATHI_KATHI_TBI" : data["activities"]["skill_activities"][1]["TRAINED BY INS"] ,
            "RIFLE_SHOOTING_TTR" : data["activities"]["skill_activities"][2]["TIMES TO REPEAT"] ,
            "RIFLE_SHOOTING_TBI" : data["activities"]["skill_activities"][2]["TRAINED BY INS"] ,
            "MARTIAL_ARTS_TTR" : data["activities"]["skill_activities"][3]["TIMES TO REPEAT"] ,
            "MARTIAL_ARTS_TBI" : data["activities"]["skill_activities"][3]["TRAINED BY INS"] ,
            "HORSE_RIDING_TTR" : data["activities"]["skill_activities"][4]["TIMES TO REPEAT"] ,
            "HORSE_RIDING_TBI" : data ["activities"]["skill_activities"][4]["TRAINED BY INS"],
            "PISTOL_SHOOTING_TTR" : data["activities"]["skill_activities"][5]["TIMES TO REPEAT"],
            "PISTOL_SHOOTING_TBI" : data["activities"]["skill_activities"][5]["TRAINED BY INS"],
            
            "TREKKING_TTR" : data["activities"]["physical_activities"][0]["TIMES TO REPEAT"],
            "TREKKING_TBI" : data["activities"]["physical_activities"][0]["TRAINED BY INS"],
            "AEROBICS_TTR" : data["activities"]["physical_activities"][1]["TIMES TO REPEAT"],
            "AEROBICS_TBI" : data["activities"]["physical_activities"][1]["TRAINED BY INS"],
            "PT_EXERCISE_TTR" : data["activities"]["physical_activities"][2]["TIMES TO REPEAT"],
            "PT_EXERCISE_TBI" : data["activities"]["physical_activities"][2]["TRAINED BY INS"],
            "MARCH_PAST_TTR" : data["activities"]["physical_activities"][3]["TIMES TO REPEAT"],
            "MARCH_PAST_TBI" : data["activities"]["physical_activities"][3]["TRAINED BY INS"],
            "COMMANDO_ACTIVITIES_TTR" : data["activities"]["physical_activities"][4]["TIMES TO REPEAT"],
            "COMMANDO_ACTIVITIES_TBI" : data["activities"]["physical_activities"][4]["TRAINED BY INS"],
            
            
            
            "RAIN_DANCE_TTR" : data["activities"]["water_activities"][0]["TIMES TO REPEAT"],
            "RAIN_DANCE_TBI" : data["activities"]["water_activities"][0]["TRAINED BY INS"],
            "SWIMMING_TTR" : data["activities"]["water_activities"][1]["TIMES TO REPEAT"],
            "SWIMMING_TBI" : data["activities"]["water_activities"][1]["TRAINED BY INS"],

            "ROCK_CLIMBING_TTR" : data["activities"]["adventure_activities"][0]["TIMES TO REPEAT"],
            "ROCK_CLIMBING_TBI" : data["activities"]["adventure_activities"][0]["TRAINED BY INS"],
            "ZIP_LINE_TTR" : data["activities"]["adventure_activities"][1]["TIMES TO REPEAT"],
            "ZIP_LINE_TBI" : data["activities"]["adventure_activities"][1]["TRAINED BY INS"],
            "RAPPELLING_TTR" : data["activities"]["adventure_activities"][2]["TIMES TO REPEAT"],
            "RAPPELLING_TBI" : data["activities"]["adventure_activities"][2]["TRAINED BY INS"],
                
            "ROPE_BRIDGE_TTR" : data["activities"]["mcf_rope_course_activities"][0]["TIMES TO REPEAT"],
            "ROPE_BRIDGE_TBI" : data["activities"]["mcf_rope_course_activities"][0]["TRAINED BY INS"],
            "LADDER_WALKING_TTR" : data["activities"]["mcf_rope_course_activities"][1]["TIMES TO REPEAT"],
            "LADDER_WALKING_TBI" : data["activities"]["mcf_rope_course_activities"][1]["TRAINED BY INS"],
            "SINGLE_ROPE_WALK_TTR" : data["activities"]["mcf_rope_course_activities"][2]["TIMES TO REPEAT"],
            "SINGLE_ROPE_WALK_TBI" : data["activities"]["mcf_rope_course_activities"][2]["TRAINED BY INS"],
            "ZIGZAG_LADDER_WALK_TTR" : data["activities"]["mcf_rope_course_activities"][3]["TIMES TO REPEAT"],
            "ZIGZAG_LADDER_WALK_TBI" : data["activities"]["mcf_rope_course_activities"][3]["TRAINED BY INS"],
            "ONE_FEET_WALK_TTR" : data["activities"]["mcf_rope_course_activities"][4]["TIMES TO REPEAT"],
            "ONE_FEET_WALK_TBI" : data["activities"]["mcf_rope_course_activities"][4]["TRAINED BY INS"],
            "STRAIGHT_LINE_WALK_TTR" : data["activities"]["mcf_rope_course_activities"][5]["TIMES TO REPEAT"],
            "STRAIGHT_LINE_WALK_TBI" : data["activities"]["mcf_rope_course_activities"][5]["TRAINED BY INS"],
            
            "CAMP_FIRE_TTR" : data["activities"]["cultural_activities"][0]["TIMES TO REPEAT"],
            "CAMP_FIRE_TBI" : data["activities"]["cultural_activities"][0]["TRAINED BY INS"],
            "KARAOKE_TTR" : data["activities"]["cultural_activities"][1]["TIMES TO REPEAT"],
            "KARAOKE_TBI" : data["activities"]["cultural_activities"][1]["TRAINED BY INS"],

            "STRAIGHT_BALANCE_TTR" : data["activities"]["Obstacle_course_activities"][0]["TIMES TO REPEAT"], 
            "STRAIGHT_BALANCE_TBI" : data["activities"]["Obstacle_course_activities"][0]["TRAINED BY INS"], 
            "CLEAR_JUMP_TTR" : data["activities"]["Obstacle_course_activities"][1]["TIMES TO REPEAT"], 
            "CLEAR_JUMP_TBI" : data["activities"]["Obstacle_course_activities"][1]["TRAINED BY INS"], 
            "DOUBLE_WALT_TTR" : data["activities"]["Obstacle_course_activities"][2]["TIMES TO REPEAT"], 
            "DOUBLE_WALT_TBI" : data["activities"]["Obstacle_course_activities"][2]["TRAINED BY INS"], 
            "ZIGZAG_TTR" : data["activities"]["Obstacle_course_activities"][3]["TIMES TO REPEAT"], 
            "ZIGZAG_TBI" : data["activities"]["Obstacle_course_activities"][3]["TRAINED BY INS"], 
            "DOUBLE_JUMP_TTR" : data["activities"]["Obstacle_course_activities"][4]["TIMES TO REPEAT"], 
            "DOUBLE_JUMP_TBI" : data["activities"]["Obstacle_course_activities"][4]["TRAINED BY INS"], 
            "WALL_CLIMBING_TTR" : data["activities"]["Obstacle_course_activities"][5]["TIMES TO REPEAT"], 
            "WALL_CLIMBING_TBI" : data["activities"]["Obstacle_course_activities"][5]["TRAINED BY INS"], 
            "TYRE_JUMP_TTR" : data["activities"]["Obstacle_course_activities"][6]["TIMES TO REPEAT"], 
            "TYRE_JUMP_TBI" : data["activities"]["Obstacle_course_activities"][6]["TRAINED BY INS"], 
            "TARZAN_SWING_TTR" : data["activities"]["Obstacle_course_activities"][7]["TIMES TO REPEAT"], 
            "TARZAN_SWING_TBI" : data["activities"]["Obstacle_course_activities"][7]["TRAINED BY INS"], 

            "FIRST_AID_TTR" : data["activities"]["disaster_management_activities"][0]["TIMES TO REPEAT"], 
            "FIRST_AID_TBI" : data["activities"]["disaster_management_activities"][0]["TRAINED BY INS"], 
            "BANDAGE_TTR" : data["activities"]["disaster_management_activities"][1]["TIMES TO REPEAT"], 
            "BANDAGE_TBI" : data["activities"]["disaster_management_activities"][1]["TRAINED BY INS"], 
            "KNOTS_TTR" : data["activities"]["disaster_management_activities"][2]["TIMES TO REPEAT"], 
            "KNOTS_TBI" : data["activities"]["disaster_management_activities"][2]["TRAINED BY INS"], 

            "COMMANDO_NT_TTR" : data["activities"]["military_obstacle_activities"][0]["TIMES TO REPEAT"],
            "COMMANDO_NT_TBI" : data["activities"]["military_obstacle_activities"][0]["TRAINED BY INS"],
            "SPIDER_NET_TTR" : data["activities"]["military_obstacle_activities"][1]["TIMES TO REPEAT"],
            "SPIDER_NET_TBI" : data["activities"]["military_obstacle_activities"][1]["TRAINED BY INS"],
            "VERTICAL_NET_TTR" : data["activities"]["military_obstacle_activities"][2]["TIMES TO REPEAT"],
            "VERTICAL_NET_TBI" : data["activities"]["military_obstacle_activities"][2]["TRAINED BY INS"],

            "GROUP_ACTIVITIES_TTR" : data["activities"]["Team_building_activities"][0]["TIMES TO REPEAT"] , 
            "GROUP_ACTIVITIES_TBI" : data["activities"]["Team_building_activities"][0]["TRAINED BY INS"], 
            "SPORTS_ACTIVITIES_TTR" : data["activities"]["Team_building_activities"][1]["TIMES TO REPEAT"] , 
            "SPORTS_ACTIVITIES_TBI" : data["activities"]["Team_building_activities"][1]["TRAINED BY INS"] , 

            "FUTURE_CAREER" : data["individual_remarks_form"]["future_career"],

            "TIME_MANAGEMENT" : data["individual_remarks_form"]["time_management"],
            "ACCOMMODATION" : data["individual_remarks_form"]["accommodation"],
            "FACILITIES" : data["individual_remarks_form"]["facilities"],
            "INS_TRAINING" : data["individual_remarks_form"]["ins_training"],
            "ATMOSPHERE" : data["individual_remarks_form"]["atmosphere"],
            "ACTIVITY" : data["individual_remarks_form"]["activity"],
            "UNIFORM" : data["individual_remarks_form"]["uniform"],
            "CERTIFICATION" : data["individual_remarks_form"]["certification"],
            "SUGGESTION" : data["individual_remarks_form"]["suggestion"],
            "BEST_ACHIEVEMENT" : data["individual_remarks_form"]["best_achievement"],
            "ACHIEVEMENT" : data["individual_remarks_form"]["achievement"],
            "PERSONALITY_DIMENSIONS" : data["individual_remarks_form"]["personality_dimensions"],

            "LEADER_P_S" : data["individual_remarks_table"]["LEADERSHIP POTENTIAL"]["SCORE"],
            "LEADER_P_I" : data["individual_remarks_table"]["LEADERSHIP POTENTIAL"]["INTERPRETATION"],
            "COMM_SKILLS_S" : data["individual_remarks_table"]["COMMUNICATION SKILLS"]["SCORE"],
            "COMM_SKILLS_I" : data["individual_remarks_table"]["COMMUNICATION SKILLS"]["INTERPRETATION"],
            "TEAM_COOR_S" : data["individual_remarks_table"]["TEAMWORK AND COOPERATION"]["SCORE"],
            "TEAM_COOR_I" : data["individual_remarks_table"]["TEAMWORK AND COOPERATION"]["INTERPRETATION"],
            "ADAPT_FLEX_S" : data["individual_remarks_table"]["ADAPTABILITY AND FLEXIBILITY"]["SCORE"],
            "ADAPT_FLEX_I" : data["individual_remarks_table"]["ADAPTABILITY AND FLEXIBILITY"]["INTERPRETATION"],
            "PRO_SOL_ABILITY_S" : data["individual_remarks_table"]["PROBLEM-SOLVING ABILITY"]["SCORE"],
            "PRO_SOL_ABILITY_I" : data["individual_remarks_table"]["PROBLEM-SOLVING ABILITY"]["INTERPRETATION"],

            "OVERALL_ASSESSMENT" : data["final_remarks"]["overall assessment"],

            
            "LEADER_DEV" : data["recommendations"]["leadership development"],
            "COMM_ENHANCE" : data["recommendations"]["communication enhancement"],
            "TEAM_ENHANCE" : data["recommendations"]["teamwork enhancement"],
            "ADAPT_TRAIN" : data["recommendations"]["adaptibility training"],
            "PRO_SOL_SKILLS_DEV" : data["recommendations"]["problem solving"],


            "CHECKED_BY_NAME" : data["final_remarks"]["checked by name"],
            "RANK" : data["final_remarks"]["rank"],
            }
    elif 'ATC' in Report_Data["CAMP_NAME"] or 'atc' in Report_Data["CAMP_NAME"]:
        doc = Document('report_card_ATC.docx')
        camp_report_data = {
                "ARCHERY_TTR" : data["activities"]["skill_activities"][0]["TIMES TO REPEAT"] ,
                "ARCHERY_TBI" : data["activities"]["skill_activities"][0]["TRAINED BY INS"] ,
                "LATHI_KATHI_TTR" : data["activities"]["skill_activities"][1]["TIMES TO REPEAT"] ,
                "LATHI_KATHI_TBI" : data["activities"]["skill_activities"][1]["TRAINED BY INS"] ,
                "RIFLE_SHOOTING_TTR" : data["activities"]["skill_activities"][2]["TIMES TO REPEAT"] ,
                "RIFLE_SHOOTING_TBI" : data["activities"]["skill_activities"][2]["TRAINED BY INS"] ,
                "MARTIAL_ARTS_TTR" : data["activities"]["skill_activities"][3]["TIMES TO REPEAT"] ,
                "MARTIAL_ARTS_TBI" : data["activities"]["skill_activities"][3]["TRAINED BY INS"] ,
                "HORSE_RIDING_TTR" : data["activities"]["skill_activities"][4]["TIMES TO REPEAT"] ,
                "HORSE_RIDING_TBI" : data ["activities"]["skill_activities"][4]["TRAINED BY INS"],

                "TREKKING_TTR" : data["activities"]["physical_activities"][0]["TIMES TO REPEAT"] ,
                "TREKKING_TBI" : data["activities"]["physical_activities"][0]["TRAINED BY INS"] ,
                "AEROBICS_TTR" : data["activities"]["physical_activities"][1]["TIMES TO REPEAT"] ,
                "AEROBICS_TBI" : data["activities"]["physical_activities"][1]["TRAINED BY INS"] ,
                "PT_EXERCISE_TTR" : data["activities"]["physical_activities"][2]["TIMES TO REPEAT"] ,
                "PT_EXERCISE_TBI" : data["activities"]["physical_activities"][2]["TRAINED BY INS"] ,

                "RAIN_DANCE_TTR" : data["activities"]["water_activities"][0]["TIMES TO REPEAT"] ,
                "RAIN_DANCE_TBI" : data["activities"]["water_activities"][0]["TRAINED BY INS"] ,

                "ROCK_CLIMBING_TTR" : data["activities"]["adventure_activities"][0]["TIMES TO REPEAT"] ,
                "ROCK_CLIMBING_TBI" : data["activities"]["adventure_activities"][0]["TRAINED BY INS"] ,
                "ZIP_LINE_TTR" : data["activities"]["adventure_activities"][1]["TIMES TO REPEAT"] ,
                "ZIP_LINE_TBI" : data["activities"]["adventure_activities"][1]["TRAINED BY INS"] ,

                "BURMA_BRIDGE_TTR" : data["activities"]["mcf_rope_course"][0]["TIMES TO REPEAT"] ,
                "BURMA_BRIDGE_TBI" : data["activities"]["mcf_rope_course"][0]["TRAINED BY INS"] ,
                "TAWA_WALK_TTR" : data["activities"]["mcf_rope_course"][1]["TIMES TO REPEAT"] ,
                "TAWA_WALK_TBI" : data["activities"]["mcf_rope_course"][1]["TRAINED BY INS"] ,
                "SINGLE_ROPE_WALK_TTR" : data["activities"]["mcf_rope_course"][2]["TIMES TO REPEAT"] ,
                "SINGLE_ROPE_WALK_TBI" : data["activities"]["mcf_rope_course"][2]["TRAINED BY INS"] ,
                "ZIGZAG_LADDER_WALK_TTR" : data["activities"]["mcf_rope_course"][3]["TIMES TO REPEAT"] ,
                "ZIGZAG_LADDER_WALK_TBI" : data["activities"]["mcf_rope_course"][3]["TRAINED BY INS"] ,
                "ONE_FEET_WALK_TTR" : data["activities"]["mcf_rope_course"][4]["TIMES TO REPEAT"] ,
                "ONE_FEET_WALK_TBI" : data["activities"]["mcf_rope_course"][4]["TRAINED BY INS"] ,

                "CAMP_FIRE_TTR" : data["activities"]["cultural_activities"][0]["TIMES TO REPEAT"] ,
                "CAMP_FIRE_TBI" : data["activities"]["cultural_activities"][0]["TRAINED BY INS"] ,


                "FUTURE_CAREER" : data["individual_remarks_form"]["future_career"],
                "TIME_MANAGEMENT" : data["individual_remarks_form"]["time_management"],
                "ACCOMMODATION" : data["individual_remarks_form"]["accommodation"],
                "FACILITIES" : data["individual_remarks_form"]["facilities"],
                "INS_TRAINING" : data["individual_remarks_form"]["ins_training"],
                "ATMOSPHERE" : data["individual_remarks_form"]["atmosphere"],
                "ACTIVITY" : data["individual_remarks_form"]["activity"],
                "UNIFORM" : data["individual_remarks_form"]["uniform"],
                "CERTIFICATION" : data["individual_remarks_form"]["certification"],
                "SUGGESTION" : data["individual_remarks_form"]["suggestion"],
                "BEST_ACHIEVEMENT" : data["individual_remarks_form"]["best_achievement"],
                "ACHIEVEMENT" : data["individual_remarks_form"]["achievement"],


        }
    elif 'PDC' in Report_Data["CAMP_NAME"] or 'pdc' in Report_Data["CAMP_NAME"] :
        doc = Document('report_card_PDC.docx')
        camp_report_data ={
            "ARCHERY_TTR" : data["activities"]["skill_activities"][0]["TIMES TO REPEAT"] ,
            "ARCHERY_TBI" : data["activities"]["skill_activities"][0]["TRAINED BY INS"] ,
            "LATHI_KATHI_TTR" : data["activities"]["skill_activities"][1]["TIMES TO REPEAT"] ,
            "LATHI_KATHI_TBI" : data["activities"]["skill_activities"][1]["TRAINED BY INS"] ,
            "RIFLE_SHOOTING_TTR" : data["activities"]["skill_activities"][2]["TIMES TO REPEAT"] ,
            "RIFLE_SHOOTING_TBI" : data["activities"]["skill_activities"][2]["TRAINED BY INS"] ,
            "MARTIAL_ARTS_TTR" : data["activities"]["skill_activities"][3]["TIMES TO REPEAT"] ,
            "MARTIAL_ARTS_TBI" : data["activities"]["skill_activities"][3]["TRAINED BY INS"] ,
            "HORSE_RIDING_TTR" : data["activities"]["skill_activities"][4]["TIMES TO REPEAT"] ,
            "HORSE_RIDING_TBI" : data ["activities"]["skill_activities"][4]["TRAINED BY INS"],
            "PISTOL_SHOOTING_TTR" : data["activities"]["skill_activities"][5]["TIMES TO REPEAT"],
            "PISTOL_SHOOTING_TBI" : data["activities"]["skill_activities"][5]["TRAINED BY INS"],
            
            "TREKKING_TTR" : data["activities"]["physical_activities"][0]["TIMES TO REPEAT"],
            "TREKKING_TBI" : data["activities"]["physical_activities"][0]["TRAINED BY INS"],
            "AEROBICS_TTR" : data["activities"]["physical_activities"][1]["TIMES TO REPEAT"],
            "AEROBICS_TBI" : data["activities"]["physical_activities"][1]["TRAINED BY INS"],
            "PT_EXERCISE_TTR" : data["activities"]["physical_activities"][2]["TIMES TO REPEAT"],
            "PT_EXERCISE_TBI" : data["activities"]["physical_activities"][2]["TRAINED BY INS"],
            "MARCH_PAST_TTR" : data["activities"]["physical_activities"][3]["TIMES TO REPEAT"],
            "MARCH_PAST_TBI" : data["activities"]["physical_activities"][3]["TRAINED BY INS"],
            "COMMANDO_ACTIVITIES_TTR" : data["activities"]["physical_activities"][4]["TIMES TO REPEAT"],
            "COMMANDO_ACTIVITIES_TBI" : data["activities"]["physical_activities"][4]["TRAINED BY INS"],
            
            "RAIN_DANCE_TTR" : data["activities"]["water_activities"][0]["TIMES TO REPEAT"],
            "RAIN_DANCE_TBI" : data["activities"]["water_activities"][0]["TRAINED BY INS"],
            "SWIMMING_TTR" : data["activities"]["water_activities"][1]["TIMES TO REPEAT"],
            "SWIMMING_TBI" : data["activities"]["water_activities"][1]["TRAINED BY INS"],

            "ROCK_CLIMBING_TTR" : data["activities"]["adventure_activities"][0]["TIMES TO REPEAT"],
            "ROCK_CLIMBING_TBI" : data["activities"]["adventure_activities"][0]["TRAINED BY INS"],
            "ZIP_LINE_TTR" : data["activities"]["adventure_activities"][1]["TIMES TO REPEAT"],
            "ZIP_LINE_TBI" : data["activities"]["adventure_activities"][1]["TRAINED BY INS"],
            "RAPPELLING_TTR" : data["activities"]["adventure_activities"][2]["TIMES TO REPEAT"],
            "RAPPELLING_TBI" : data["activities"]["adventure_activities"][2]["TRAINED BY INS"],
                
            "ROPE_BRIDGE_TTR" : data["activities"]["mcf_rope_course_activities"][0]["TIMES TO REPEAT"],
            "ROPE_BRIDGE_TBI" : data["activities"]["mcf_rope_course_activities"][0]["TRAINED BY INS"],
            "LADDER_WALKING_TTR" : data["activities"]["mcf_rope_course_activities"][1]["TIMES TO REPEAT"],
            "LADDER_WALKING_TBI" : data["activities"]["mcf_rope_course_activities"][1]["TRAINED BY INS"],
            "SINGLE_ROPE_WALK_TTR" : data["activities"]["mcf_rope_course_activities"][2]["TIMES TO REPEAT"],
            "SINGLE_ROPE_WALK_TBI" : data["activities"]["mcf_rope_course_activities"][2]["TRAINED BY INS"],
            "ZIGZAG_LADDER_WALK_TTR" : data["activities"]["mcf_rope_course_activities"][3]["TIMES TO REPEAT"],
            "ZIGZAG_LADDER_WALK_TBI" : data["activities"]["mcf_rope_course_activities"][3]["TRAINED BY INS"],
            "ONE_FEET_WALK_TTR" : data["activities"]["mcf_rope_course_activities"][4]["TIMES TO REPEAT"],
            "ONE_FEET_WALK_TBI" : data["activities"]["mcf_rope_course_activities"][4]["TRAINED BY INS"],
            "STRAIGHT_LINE_WALK_TTR" : data["activities"]["mcf_rope_course_activities"][5]["TIMES TO REPEAT"],
            "STRAIGHT_LINE_WALK_TBI" : data["activities"]["mcf_rope_course_activities"][5]["TRAINED BY INS"],
            
            "CAMP_FIRE_TTR" : data["activities"]["cultural_activities"][0]["TIMES TO REPEAT"],
            "CAMP_FIRE_TBI" : data["activities"]["cultural_activities"][0]["TRAINED BY INS"],
            "KARAOKE_TTR" : data["activities"]["cultural_activities"][1]["TIMES TO REPEAT"],
            "KARAOKE_TBI" : data["activities"]["cultural_activities"][1]["TRAINED BY INS"],

            "STRAIGHT_BALANCE_TTR" : data["activities"]["Obstacle_course_activities"][0]["TIMES TO REPEAT"], 
            "STRAIGHT_BALANCE_TBI" : data["activities"]["Obstacle_course_activities"][0]["TRAINED BY INS"], 
            "CLEAR_JUMP_TTR" : data["activities"]["Obstacle_course_activities"][1]["TIMES TO REPEAT"], 
            "CLEAR_JUMP_TBI" : data["activities"]["Obstacle_course_activities"][1]["TRAINED BY INS"], 
            "DOUBLE_WALT_TTR" : data["activities"]["Obstacle_course_activities"][2]["TIMES TO REPEAT"], 
            "DOUBLE_WALT_TBI" : data["activities"]["Obstacle_course_activities"][2]["TRAINED BY INS"], 
            "ZIGZAG_TTR" : data["activities"]["Obstacle_course_activities"][3]["TIMES TO REPEAT"], 
            "ZIGZAG_TBI" : data["activities"]["Obstacle_course_activities"][3]["TRAINED BY INS"], 
            "DOUBLE_JUMP_TTR" : data["activities"]["Obstacle_course_activities"][4]["TIMES TO REPEAT"], 
            "DOUBLE_JUMP_TBI" : data["activities"]["Obstacle_course_activities"][4]["TRAINED BY INS"], 
            "WALL_CLIMBING_TTR" : data["activities"]["Obstacle_course_activities"][5]["TIMES TO REPEAT"], 
            "WALL_CLIMBING_TBI" : data["activities"]["Obstacle_course_activities"][5]["TRAINED BY INS"], 
            "TYRE_JUMP_TTR" : data["activities"]["Obstacle_course_activities"][6]["TIMES TO REPEAT"], 
            "TYRE_JUMP_TBI" : data["activities"]["Obstacle_course_activities"][6]["TRAINED BY INS"], 
            "TARZAN_SWING_TTR" : data["activities"]["Obstacle_course_activities"][7]["TIMES TO REPEAT"], 
            "TARZAN_SWING_TBI" : data["activities"]["Obstacle_course_activities"][7]["TRAINED BY INS"], 

            "FIRST_AID_TTR" : data["activities"]["disaster_management_activities"][0]["TIMES TO REPEAT"], 
            "FIRST_AID_TBI" : data["activities"]["disaster_management_activities"][0]["TRAINED BY INS"], 
            "BANDAGE_TTR" : data["activities"]["disaster_management_activities"][1]["TIMES TO REPEAT"], 
            "BANDAGE_TBI" : data["activities"]["disaster_management_activities"][1]["TRAINED BY INS"], 
            "KNOTS_TTR" : data["activities"]["disaster_management_activities"][2]["TIMES TO REPEAT"], 
            "KNOTS_TBI" : data["activities"]["disaster_management_activities"][2]["TRAINED BY INS"], 

            "COMMANDO_NT_TTR" : data["activities"]["military_obstacle_activities"][0]["TIMES TO REPEAT"],
            "COMMANDO_NT_TBI" : data["activities"]["military_obstacle_activities"][0]["TRAINED BY INS"],
            "SPIDER_NET_TTR" : data["activities"]["military_obstacle_activities"][1]["TIMES TO REPEAT"],
            "SPIDER_NET_TBI" : data["activities"]["military_obstacle_activities"][1]["TRAINED BY INS"],
            "VERTICAL_NET_TTR" : data["activities"]["military_obstacle_activities"][2]["TIMES TO REPEAT"],
            "VERTICAL_NET_TBI" : data["activities"]["military_obstacle_activities"][2]["TRAINED BY INS"],

            "GROUP_ACTIVITIES_TTR" : data["activities"]["Team_building_activities"][0]["TIMES TO REPEAT"] , 
            "GROUP_ACTIVITIES_TBI" : data["activities"]["Team_building_activities"][0]["TRAINED BY INS"], 
            "SPORTS_ACTIVITIES_TTR" : data["activities"]["Team_building_activities"][1]["TIMES TO REPEAT"] , 
            "SPORTS_ACTIVITIES_TBI" : data["activities"]["Team_building_activities"][1]["TRAINED BY INS"] , 

            "FUTURE_CAREER" : data["individual_remarks_form"]["future_career"],

            "INTRO_INSTRUCTORS" : data["training"]["INTRODUCTION OF INSTRUCTORS, CADETS AND AGENDA SETTING"] ,
            "PERSONALITY_DEV" : data["training"]["CONCEPT OF PERSONALITY AND ITS DEVELOPMENT"] ,
            "SELFASSESSMENT_TEST" : data["training"]["INITIAL SELFASSESSMENT AND TEST"] ,
            "MORAL_VALUES__ETHICS" : data["training"]["MORALS, VALUES & ETHICS"] ,
            "SELF_EMOTIONAL_AWARNESS_" : data["training"]["SELF-AWARNESS & EMOTIONAL AWARENESS"] ,
            "LOGIC_REASONING" : data["training"]["LOGIC & REASONING"] ,
            "COMMUNICATION_SKILLS" : data["training"]["COMMUNICATION SKILLS"] ,
            "LEADERSHIP_SKLLS" : data["training"]["LEADERSHIP SKILLS"] ,
            "STAMINA_ENDURANCE" : data["training"]["STAMINA & ENDURANCE"] ,
            "TIME_MANAGEMENT" : data["training"]["TIME MANAGEMENT"] ,
            "PLANNING_ORG_EXERCISE" : data["training"]["PLANNING & ORGANISING EXERCISE"] ,
            "INTRO_TOOLS_DEV_PERSONALITY" : data["training"]["INTROSPECTION, TOOLS OF DEVELOPING PERSONALITY"] ,
            "ADVENTURE_ACTIVITY_PER_DEV" : data["training"]["ADVENTURE ACTIVITY & PERSONALITY DEVELOPMENT"] ,
            "SELF_EVALUATION_EXERCISE" : data["training"]["Self Evaluation Exercise, individual Counselling"] ,
            "INDIVIDUAL_COUNCELLING" : data["training"]["Individual Counselling session continue"] ,
            "PER_RELATION_OFFICER" : data["training"]["Personality and its relation to becoming officer or Good Citizen"] ,
            "PERSONALITY_ASSESSMENT" : data["training"]["Concept of personality assessment, introduction to officer like qualities"] ,

            "DATE_OF_ASSESSMENT" : data["dates"]["date_of_assessment"] ,
            "PERSONALITY_DIMENTIONS" : data["individual_remarks_form"]["personality_dimensions"] ,


            "LEADER_P_S" : data["individual_remarks_table"]["LEADERSHIP POTENTIAL"]["SCORE"],
            "LEADER_P_I" : data["individual_remarks_table"]["LEADERSHIP POTENTIAL"]["INTERPRETATION"],
            "COMM_SKILLS_S" : data["individual_remarks_table"]["COMMUNICATION SKILLS"]["SCORE"],
            "COMM_SKILLS_I" : data["individual_remarks_table"]["COMMUNICATION SKILLS"]["INTERPRETATION"],
            "TEAM_COOR_S" : data["individual_remarks_table"]["TEAMWORK AND COOPERATION"]["SCORE"],
            "TEAM_COOR_I" : data["individual_remarks_table"]["TEAMWORK AND COOPERATION"]["INTERPRETATION"],
            "ADAPT_FLEX_S" : data["individual_remarks_table"]["ADAPTABILITY AND FLEXIBILITY"]["SCORE"],
            "ADAPT_FLEX_I" : data["individual_remarks_table"]["ADAPTABILITY AND FLEXIBILITY"]["INTERPRETATION"],
            "PRO_SOL_ABILITY_S" : data["individual_remarks_table"]["PROBLEM-SOLVING ABILITY"]["SCORE"],
            "PRO_SOL_ABILITY_I" : data["individual_remarks_table"]["PROBLEM-SOLVING ABILITY"]["INTERPRETATION"],

            "OVERALL_ASSESSMENT" : data["final_remarks"]["overall assessment"],

            
            "LEADER_DEV" : data["recommendations"]["leadership development"],
            "COMM_ENHANCE" : data["recommendations"]["communication enhancement"],
            "TEAM_ENHANCE" : data["recommendations"]["teamwork enhancement"],
            "ADAPT_TRAIN" : data["recommendations"]["adaptibility training"],
            "PRO_SOL_SKILLS_DEV" : data["recommendations"]["problem solving"],



            "CHECKED_BY_NAME" : data["final_remarks"]["checked by name"],
            "RANK" : data["final_remarks"]["rank"],


        }
    else:
        print("No report template found for camp: " + Report_Data["CAMP_NAME"])


    replace_image_in_cell(doc, table_index=0, row_index=0, column_index=0, image_path=image_path_photo,w=1.0,h=1.2)
    
    for key, value in Report_Data.items():
        find_and_replace_tables(doc.tables, f'{{MERGEFIELD {key}}}', str(value),('Times New Roman', 13))
    for key, value in camp_report_data.items():
        find_and_replace_tables(doc.tables, f'{{MERGEFIELD {key}}}', str(value),('Times New Roman', 13))


    doc.save(docx_path)
    convert_to_pdf(docx_path,output_path)


# sid= "ATC30D2BC2"
# data={
#     "details": {
#         "name": "NAVNEETH  IYER",
#         "email": "shwetha.sharma05@gnail.com",
#         "phone": "9483984121",
#         "address": "NO 43 4TH MAIN SBM COLONY ANAND NAGAR BLR-24",
#         "camp_name": "COMMANDO TRAINING CAMP -2024 (CTC)",
#         "pickup_point": "BELAPUR",
#         "cqy_name": "abc",
#         "incharge_name": "abc"
#     },
#     "activities": {
#         "skill_activities": [
#             {
#                 "SR.NO.": 1,
#                 "SKILL": "Archery",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             },
#             {
#                 "SR.NO.": 2,
#                 "SKILL": "Lathi-Kathi",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             },
#             {
#                 "SR.NO.": 3,
#                 "SKILL": "Rifle Shooting",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             },
#             {
#                 "SR.NO.": 4,
#                 "SKILL": "Martial Arts",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             },
#             {
#                 "SR.NO.": 5,
#                 "SKILL": "Horse Riding",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             },
#             {
#                 "SR.NO.": 6,
#                 "SKILL": "Pistol Shooting",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             }
#         ],
#         "physical_activities": [
#             {
#                 "SR.NO.": 7,
#                 "SKILL": "Trekking",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             },
#             {
#                 "SR.NO.": 8,
#                 "SKILL": "Aerobics/Yoga",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             },
#             {
#                 "SR.NO.": 9,
#                 "SKILL": "P.T. and Mass P.T.Execise",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             },
#             {
#                 "SR.NO.": 10,
#                 "SKILL": "March Past/Drill",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             },
#             {
#                 "SR.NO.": 11,
#                 "SKILL": "Commando Activities",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             }
#         ],
#         "water_activities": [
#             {
#                 "SR.NO.": 12,
#                 "SKILL": "Rain Dance",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             },
#             {
#                 "SR.NO.": 13,
#                 "SKILL": "Swimming",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             }
#         ],
#         "adventure_activities": [
#             {
#                 "SR.NO.": 15,
#                 "SKILL": "Rock Climbing",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             },
#             {
#                 "SR.NO.": 16,
#                 "SKILL": "Zip Line",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             },
#             {
#                 "SR.NO.": 17,
#                 "SKILL": "Rappelling",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             }
#         ],
#         "mcf_rope_course_activities": [
#             {
#                 "SR.NO.": 18,
#                 "SKILL": "Rope Bridge",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             },
#             {
#                 "SR.NO.": 19,
#                 "SKILL": "Ladder Walking",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             },
#             {
#                 "SR.NO.": 20,
#                 "SKILL": "Single Rope Walk",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             },
#             {
#                 "SR.NO.": 21,
#                 "SKILL": "Zig Zag Ladder Walk",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             },
#             {
#                 "SR.NO.": 22,
#                 "SKILL": "One Feet Walk",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             },
#             {
#                 "SR.NO.": 23,
#                 "SKILL": "Straight Line Walk",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             }
#         ],
#         "cultural_activities": [
#             {
#                 "SR.NO.": 24,
#                 "SKILL": "Camp Fire",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             },
#             {
#                 "SR.NO.": 25,
#                 "SKILL": "Karaoke",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             }
#         ],
#         "Obstacle_course_activities": [
#             {
#                 "SR.NO.": 26,
#                 "SKILL": "Straight Balance",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             },
#             {
#                 "SR.NO.": 27,
#                 "SKILL": "Clear Jump",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             },
#             {
#                 "SR.NO.": 28,
#                 "SKILL": "Double Walt ",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             },
#             {
#                 "SR.NO.": 29,
#                 "SKILL": "Zig Zag",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             },
#             {
#                 "SR.NO.": 30,
#                 "SKILL": "Double Jump",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             },
#             {
#                 "SR.NO.": 31,
#                 "SKILL": "Wall Climbing",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             },
#             {
#                 "SR.NO.": 32,
#                 "SKILL": "Tire Jump",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             },
#             {
#                 "SR.NO.": 33,
#                 "SKILL": "Tarzan Swing",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             }
#         ],
#         "disaster_management_activities": [
#             {
#                 "SR.NO.": 34,
#                 "SKILL": "First Aid",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             },
#             {
#                 "SR.NO.": 35,
#                 "SKILL": "Bandage",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             },
#             {
#                 "SR.NO.": 36,
#                 "SKILL": "Knots",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             }
#         ],
#         "military_obstacle_activities": [
#             {
#                 "SR.NO.": 37,
#                 "SKILL": "Commando Net",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             },
#             {
#                 "SR.NO.": 38,
#                 "SKILL": "Spider Net",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             },
#             {
#                 "SR.NO.": 39,
#                 "SKILL": "Verticle Net",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             }
#         ],
#         "Team_building_activities": [
#             {
#                 "SR.NO.": 40,
#                 "SKILL": "Group Activities",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             },
#             {
#                 "SR.NO.": 41,
#                 "SKILL": "Sports Activities",
#                 "TIMES TO REPEAT": 0,
#                 "TRAINED BY INS": ""
#             }
#         ]
#     },
#     "dates": {
#         "checkin_Date": "2024-04-25",
#         "pickup_Date": "2024-04-11",
#         "checkout_Date": "2024-04-18",
#         "drop_Date": "2024-04-17",
#         "last_closing_ceremony": "2024-04-09",
#         "last_closing_time": "17:44"
#     },
#     "individual_remarks_form": {
#         "future_career": "abc",
#         "time_management": "abc",
#         "accommodation": "abc",
#         "facilities": "abc",
#         "ins_training": "abc",
#         "atmosphere": "abc",
#         "activity": "abc",
#         "uniform": "abc",
#         "certification": "abc",
#         "suggestion": "abc",
#         "best_achievement": "abcabc",
#         "achievement": "abc",
#         "personality_dimensions": "abc"
#     },
#     "individual_remarks_table": {
#         "LEADERSHIP POTENTIAL": {
#             "SCORE": "abc",
#             "INTERPRETATION": "abc"
#         },
#         "COMMUNICATION SKILLS": {
#             "SCORE": "abc",
#             "INTERPRETATION": "abc"
#         },
#         "TEAMWORK AND COOPERATION": {
#             "SCORE": "abc",
#             "INTERPRETATION": "abc"
#         },
#         "ADAPTABILITY AND FLEXIBILITY": {
#             "SCORE": "abc",
#             "INTERPRETATION": "abc"
#         },
#         "PROBLEM-SOLVING ABILITY": {
#             "SCORE": "abc",
#             "INTERPRETATION": "abc"
#         }
#     },
#     "recommendations": {
#         "leadership development": "abc",
#         "communication enhancement": "abc",
#         "teamwork enhancement": "abc",
#         "adaptibility training": "abc",
#         "problem solving": "abc"
#     },
#     "final_remarks": {
#         "Parents presence": "abc",
#         "best activity": "abc",
#         "remarks": "abc",
#         "checked by name": "abc",
#         "rank": "abc",
#         "overall assessment": "abc"
#     }
# }

# generate_report()


